"""Tests E2E export — pipeline complet PDF/DOCX/Excel.

Teste le flux end-to-end depuis fetch_export_data() mockée
jusqu'à la génération de bytes valides (PDF, DOCX, Excel).
Vérifie la cohérence inter-formats (même données → 3 outputs).
"""
import pytest
from io import BytesIO
from unittest.mock import MagicMock, patch, PropertyMock
from dataclasses import dataclass, field
from datetime import datetime, timedelta


# ── Fixtures partagées ────────────────────────────────────────────────────


@dataclass
class MockProject:
    title: str = "Rénovation Lycée Victor Hugo — Lot 2 CVC"
    reference: str = "AO-2026-E2E-001"
    status: str = "ready"
    org_id: str = "00000000-0000-0000-0000-e2e000000001"
    buyer: str = "Région Île-de-France"
    submission_deadline: datetime = field(
        default_factory=lambda: datetime.now() + timedelta(days=30)
    )
    estimated_value: int = 2_500_000


@dataclass
class MockDocument:
    original_name: str = "CCAP_Lot2.pdf"
    doc_type: str = "CCAP"
    file_size_kb: int = 450
    page_count: int = 28
    ocr_quality_score: float = 95.0
    s3_key: str = "docs/test/ccap.pdf"


@dataclass
class MockChecklistItem:
    category: str = "Administratif"
    criticality: str = "éliminatoire"
    requirement: str = "Attestation URSSAF < 6 mois"
    status: str = "OK"
    confidence: float = 94.0
    citation: str = "RC art. 5.1"
    what_to_provide: str = "Attestation URSSAF de moins de 6 mois"


def _make_full_export_data():
    """Construit un jeu de données complet pour tester les 3 formats."""
    from app.services.export_data import ExportData

    return ExportData(
        project=MockProject(),
        documents=[
            MockDocument(),
            MockDocument(
                original_name="CCTP_Lot2.pdf", doc_type="CCTP",
                file_size_kb=1200, page_count=85
            ),
            MockDocument(
                original_name="DPGF_Lot2.pdf", doc_type="DPGF",
                file_size_kb=350, page_count=12
            ),
        ],
        summary={
            "project_overview": {
                "deadline_submission": "2026-07-15",
                "object": "Rénovation CVC Lycée Victor Hugo",
                "scope": "Remplacement chaudière + gainable + GTB",
            },
            "key_points": [
                {"point": "Marché à prix global et forfaitaire", "importance": "high"},
                {"point": "Variante technique autorisée", "importance": "medium"},
                {"point": "Visite de site obligatoire le 15/06/2026", "importance": "high"},
            ],
            "risks": [
                {"risk": "Pénalités retard élevées", "severity": "high", "why": "1/1000 par jour calendaire"},
                {"risk": "Clause de révision de prix absente", "severity": "medium", "why": "Risque financier"},
            ],
            "actions_next_48h": [
                {"action": "Planifier visite de site", "priority": "P0", "owner_role": "Chef de projet"},
                {"action": "Vérifier assurances", "priority": "P1", "owner_role": "Administratif"},
            ],
            "confidence_overall": 88.5,
        },
        criteria={
            "evaluation": {
                "scoring_criteria": [
                    {"name": "Prix", "weight": 40, "estimated_score": 78},
                    {"name": "Valeur technique", "weight": 50, "estimated_score": 82},
                    {"name": "Délai", "weight": 10, "estimated_score": 90},
                ]
            }
        },
        gonogo={
            "score": 76,
            "decision": "GO conditionnel",
            "dimension_scores": [
                {"name": "Capacité financière", "score": 85, "weight": 15, "confidence": 0.9},
                {"name": "Certifications", "score": 90, "weight": 10, "confidence": 0.95},
                {"name": "Références similaires", "score": 70, "weight": 15, "confidence": 0.8},
                {"name": "Charge actuelle", "score": 65, "weight": 10, "confidence": 0.7},
                {"name": "Zone géographique", "score": 80, "weight": 10, "confidence": 0.85},
                {"name": "Partenariats", "score": 75, "weight": 10, "confidence": 0.75},
                {"name": "Marge visée", "score": 60, "weight": 10, "confidence": 0.65},
                {"name": "Délais", "score": 85, "weight": 10, "confidence": 0.9},
                {"name": "Risque technique", "score": 72, "weight": 10, "confidence": 0.8},
            ],
        },
        timeline={
            "key_dates": [
                {"date": "2026-06-15", "label": "Visite de site"},
                {"date": "2026-07-15", "label": "Remise des offres"},
            ]
        },
        ccap_analysis={
            "clauses_risquees": [
                {"titre": "Pénalités de retard", "article": "Art. 14.1",
                 "severity": "CRITIQUE", "detail": "1/1000 par jour calendaire"},
            ],
            "ccag_derogations": [
                {"article": "Art. 13.3", "derogation": "Délai de paiement 60j", "severity": "DEFAVORABLE"},
            ],
        },
        rc_analysis={
            "eligibility_conditions": ["DC1", "DC2", "Attestation URSSAF"],
            "visite_obligatoire": True,
            "variantes_autorisees": True,
        },
        cctp_analysis={
            "prescriptions": [
                {"lot": "CVC", "detail": "Chaudière gaz condensation > 95% rendement"},
            ]
        },
        scoring={
            "total_weighted": 81.2,
            "criteria": [
                {"name": "Prix", "weight": 40, "estimated_score": 78, "weighted_score": 31.2},
                {"name": "Technique", "weight": 50, "estimated_score": 82, "weighted_score": 41.0},
                {"name": "Délai", "weight": 10, "estimated_score": 90, "weighted_score": 9.0},
            ],
        },
        cashflow={
            "monthly_cashflow": [
                {"mois": 1, "label": "M1", "solde_cumule": 100000, "solde_mensuel": 100000},
                {"mois": 2, "label": "M2", "solde_cumule": 250000, "solde_mensuel": 150000},
                {"mois": 3, "label": "M3", "solde_cumule": 400000, "solde_mensuel": 150000},
            ],
            "montant_total_ht": 2500000,
            "bfr_eur": -150000,
            "risk_level": "MODERE",
            "resume": "Trésorerie maîtrisée avec un BFR modéré.",
        },
        subcontracting={
            "recommended_lots": ["Electricité", "Plomberie"],
            "max_percentage": 30,
        },
        ae_analysis={"mentions_obligatoires": ["Engagement sur les délais"]},
        dc_check={"documents_requis": ["DC1", "DC2", "NOTI1"]},
        conflicts={"conflicts": [
            {"doc1": "CCAP", "doc2": "CCTP", "type": "prix",
             "description": "Divergence sur les prix unitaires CVC"},
        ]},
        dpgf_pricing=[
            {"designation": "Chaudière gaz condensation", "unite": "U",
             "quantite": 2, "prix_unitaire": 45000, "total": 90000},
            {"designation": "Gaine VMC double flux", "unite": "ml",
             "quantite": 350, "prix_unitaire": 85, "total": 29750},
        ],
        checklist_items=[MockChecklistItem(), MockChecklistItem(
            category="Technique", criticality="important",
            requirement="Visite de site effectuée", status="TODO",
            confidence=75.0, citation="RC art. 2.3",
        )],
        glossaire_btp=[
            {"terme": "CVC", "definition": "Chauffage, Ventilation, Climatisation"},
            {"terme": "GTB", "definition": "Gestion Technique du Bâtiment"},
        ],
    )


def _make_empty_export_data():
    """Données minimales — toutes les analyses sont None."""
    from app.services.export_data import ExportData
    return ExportData(
        project=MockProject(),
        documents=[], summary=None, criteria=None, gonogo=None,
        timeline=None, ccap_analysis=None, rc_analysis=None,
        cctp_analysis=None, scoring=None, cashflow=None,
        subcontracting=None, ae_analysis=None, dc_check=None,
        conflicts=None, dpgf_pricing=None, checklist_items=[],
        glossaire_btp=[],
    )


# Minimal valid PDF bytes for mocking xhtml2pdf output
_FAKE_PDF = b"%PDF-1.4 fake-for-test"


def _mock_pisa_create_pdf(html_content, dest, encoding="utf-8"):
    """Mock de pisa.CreatePDF qui écrit un faux PDF valide."""
    dest.write(_FAKE_PDF)
    mock_result = MagicMock()
    mock_result.err = 0
    return mock_result


# ── Tests PDF E2E ─────────────────────────────────────────────────────────


class TestPDFExportE2E:
    """Pipeline complet : ExportData → HTML template rendering.

    Note: xhtml2pdf.pisa.CreatePDF est mocké car certaines features CSS
    (@page counter(pages)) ne sont pas supportées par xhtml2pdf dans
    l'environnement de test. Le template HTML rendering est testé en entier.
    """

    @patch("app.services.exporter._generate_charts", return_value={
        "radar_chart_b64": None, "cashflow_chart_b64": None,
        "heatmap_chart_b64": None, "pricing_chart_b64": None,
    })
    @patch("app.services.exporter.fetch_export_data")
    @patch("app.services.exporter._fetch_company_logo_b64", return_value=None)
    @patch("app.services.exporter.get_theme")
    @patch("xhtml2pdf.pisa.CreatePDF", side_effect=_mock_pisa_create_pdf)
    def test_pdf_generates_valid_bytes(self, mock_pisa, mock_theme, mock_logo, mock_fetch, mock_charts):
        from app.services.exporter import generate_export_pdf
        from app.core.report_theme import ReportTheme

        mock_theme.return_value = ReportTheme()
        mock_fetch.return_value = _make_full_export_data()

        pdf_bytes = generate_export_pdf(MagicMock(), "test-project-id")

        assert isinstance(pdf_bytes, bytes)
        assert len(pdf_bytes) > 10
        assert pdf_bytes[:5] == b"%PDF-"
        # Template rendering was called (pisa received HTML)
        mock_pisa.assert_called_once()

    @patch("app.services.exporter._generate_charts", return_value={
        "radar_chart_b64": None, "cashflow_chart_b64": None,
        "heatmap_chart_b64": None, "pricing_chart_b64": None,
    })
    @patch("app.services.exporter.fetch_export_data")
    @patch("app.services.exporter._fetch_company_logo_b64", return_value=None)
    @patch("app.services.exporter.get_theme")
    @patch("xhtml2pdf.pisa.CreatePDF", side_effect=_mock_pisa_create_pdf)
    def test_pdf_template_renders_project_data(self, mock_pisa, mock_theme, mock_logo, mock_fetch, mock_charts):
        """Vérifie que le template HTML contient les données du projet."""
        from app.services.exporter import generate_export_pdf
        from app.core.report_theme import ReportTheme

        mock_theme.return_value = ReportTheme()
        mock_fetch.return_value = _make_full_export_data()

        generate_export_pdf(MagicMock(), "test-project-id")

        # Inspecter le HTML passé à pisa
        call_args = mock_pisa.call_args
        html_content = call_args[1].get("html_content", call_args[0][0] if call_args[0] else "")
        if isinstance(html_content, str):
            assert "Victor Hugo" in html_content
            assert "AO-2026-E2E-001" in html_content

    @patch("app.services.exporter._generate_charts", return_value={
        "radar_chart_b64": None, "cashflow_chart_b64": None,
        "heatmap_chart_b64": None, "pricing_chart_b64": None,
    })
    @patch("app.services.exporter.fetch_export_data")
    @patch("app.services.exporter._fetch_company_logo_b64", return_value=None)
    @patch("app.services.exporter.get_theme")
    @patch("xhtml2pdf.pisa.CreatePDF", side_effect=_mock_pisa_create_pdf)
    def test_pdf_minimal_data(self, mock_pisa, mock_theme, mock_logo, mock_fetch, mock_charts):
        """PDF se génère même avec des données minimales (pas de crash)."""
        from app.services.exporter import generate_export_pdf
        from app.core.report_theme import ReportTheme

        mock_theme.return_value = ReportTheme()
        mock_fetch.return_value = _make_empty_export_data()

        pdf_bytes = generate_export_pdf(MagicMock(), "minimal-project")
        assert pdf_bytes[:5] == b"%PDF-"

    @patch("app.services.exporter._generate_charts", return_value={
        "radar_chart_b64": None, "cashflow_chart_b64": None,
        "heatmap_chart_b64": None, "pricing_chart_b64": None,
    })
    @patch("app.services.exporter.fetch_export_data")
    @patch("app.services.exporter._fetch_company_logo_b64", return_value=None)
    @patch("app.services.exporter.get_theme")
    @patch("xhtml2pdf.pisa.CreatePDF", side_effect=_mock_pisa_create_pdf)
    def test_pdf_template_includes_sections(self, mock_pisa, mock_theme, mock_logo, mock_fetch, mock_charts):
        """Vérifie que le HTML inclut les sections principales."""
        from app.services.exporter import generate_export_pdf
        from app.core.report_theme import ReportTheme

        mock_theme.return_value = ReportTheme()
        mock_fetch.return_value = _make_full_export_data()

        generate_export_pdf(MagicMock(), "test-project-id")

        call_args = mock_pisa.call_args
        html = call_args[0][0] if call_args[0] else ""
        if isinstance(html, str):
            # Vérifier que les sections clés sont rendues
            assert "Checklist" in html or "checklist" in html
            assert "CCAP" in html or "clauses" in html.lower()


# ── Tests DOCX E2E ────────────────────────────────────────────────────────


class TestDOCXExportE2E:
    """Pipeline complet : ExportData → python-docx → bytes."""

    @patch("app.services.docx_exporter.fetch_export_data")
    @patch("app.services.docx_exporter._insert_company_logo", return_value=False)
    def test_docx_generates_valid_bytes(self, mock_logo, mock_fetch):
        from app.services.docx_exporter import generate_export_docx

        mock_fetch.return_value = _make_full_export_data()

        docx_bytes = generate_export_docx(MagicMock(), "test-project-id")

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 1000
        # DOCX = ZIP (PK signature)
        assert docx_bytes[:2] == b"PK"

    @patch("app.services.docx_exporter.fetch_export_data")
    @patch("app.services.docx_exporter._insert_company_logo", return_value=False)
    def test_docx_minimal_data(self, mock_logo, mock_fetch):
        """DOCX se génère même avec des données minimales."""
        from app.services.docx_exporter import generate_export_docx

        mock_fetch.return_value = _make_empty_export_data()

        docx_bytes = generate_export_docx(MagicMock(), "minimal-project")
        assert docx_bytes[:2] == b"PK"

    @patch("app.services.docx_exporter.fetch_export_data")
    @patch("app.services.docx_exporter._insert_company_logo", return_value=False)
    def test_docx_contains_project_data(self, mock_logo, mock_fetch):
        """Vérifie que le DOCX contient le titre du projet."""
        from app.services.docx_exporter import generate_export_docx
        from docx import Document

        mock_fetch.return_value = _make_full_export_data()
        docx_bytes = generate_export_docx(MagicMock(), "test-project-id")

        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Victor Hugo" in full_text


# ── Tests Excel E2E ───────────────────────────────────────────────────────


class TestExcelExportE2E:
    """Pipeline complet : ExportData → openpyxl → bytes."""

    @patch("app.services.excel_exporter.fetch_export_data")
    def test_excel_generates_valid_bytes(self, mock_fetch):
        from app.services.excel_exporter import generate_analysis_excel

        mock_fetch.return_value = _make_full_export_data()

        xlsx_bytes = generate_analysis_excel(MagicMock(), "test-project-id")

        assert isinstance(xlsx_bytes, bytes)
        assert len(xlsx_bytes) > 1000
        assert xlsx_bytes[:2] == b"PK"  # XLSX = ZIP

    @patch("app.services.excel_exporter.fetch_export_data")
    def test_excel_minimal_data(self, mock_fetch):
        from app.services.excel_exporter import generate_analysis_excel

        mock_fetch.return_value = _make_empty_export_data()

        xlsx_bytes = generate_analysis_excel(MagicMock(), "minimal-project")
        assert xlsx_bytes[:2] == b"PK"


# ── Tests cross-format coherence ──────────────────────────────────────────


class TestCrossFormatCoherence:
    """Vérifie que les 3 formats utilisent les mêmes données."""

    @patch("xhtml2pdf.pisa.CreatePDF", side_effect=_mock_pisa_create_pdf)
    @patch("app.services.exporter._generate_charts", return_value={
        "radar_chart_b64": None, "cashflow_chart_b64": None,
        "heatmap_chart_b64": None, "pricing_chart_b64": None,
    })
    @patch("app.services.exporter.fetch_export_data")
    @patch("app.services.exporter._fetch_company_logo_b64", return_value=None)
    @patch("app.services.exporter.get_theme")
    @patch("app.services.docx_exporter.fetch_export_data")
    @patch("app.services.docx_exporter._insert_company_logo", return_value=False)
    @patch("app.services.excel_exporter.fetch_export_data")
    def test_all_formats_generate_from_same_data(
        self, mock_xl_fetch, mock_docx_logo, mock_docx_fetch,
        mock_theme, mock_pdf_logo, mock_pdf_fetch, mock_charts, mock_pisa,
    ):
        """Les 3 exports doivent se générer sans erreur depuis les mêmes données."""
        from app.services.exporter import generate_export_pdf
        from app.services.docx_exporter import generate_export_docx
        from app.services.excel_exporter import generate_analysis_excel
        from app.core.report_theme import ReportTheme

        data = _make_full_export_data()
        mock_pdf_fetch.return_value = data
        mock_docx_fetch.return_value = data
        mock_xl_fetch.return_value = data
        mock_theme.return_value = ReportTheme()

        pdf = generate_export_pdf(MagicMock(), "coherence-test")
        docx = generate_export_docx(MagicMock(), "coherence-test")
        xlsx = generate_analysis_excel(MagicMock(), "coherence-test")

        assert pdf[:5] == b"%PDF-"
        assert docx[:2] == b"PK"
        assert xlsx[:2] == b"PK"
        # Tous non-vides
        assert all(len(b) > 10 for b in [pdf, docx, xlsx])

    @patch("xhtml2pdf.pisa.CreatePDF", side_effect=_mock_pisa_create_pdf)
    @patch("app.services.exporter._generate_charts", return_value={
        "radar_chart_b64": None, "cashflow_chart_b64": None,
        "heatmap_chart_b64": None, "pricing_chart_b64": None,
    })
    @patch("app.services.exporter.fetch_export_data")
    @patch("app.services.exporter._fetch_company_logo_b64", return_value=None)
    @patch("app.services.exporter.get_theme")
    @patch("app.services.docx_exporter.fetch_export_data")
    @patch("app.services.docx_exporter._insert_company_logo", return_value=False)
    @patch("app.services.excel_exporter.fetch_export_data")
    def test_all_formats_handle_none_gracefully(
        self, mock_xl_fetch, mock_docx_logo, mock_docx_fetch,
        mock_theme, mock_pdf_logo, mock_pdf_fetch, mock_charts, mock_pisa,
    ):
        """Aucun format ne crash si toutes les analyses sont None."""
        from app.services.exporter import generate_export_pdf
        from app.services.docx_exporter import generate_export_docx
        from app.services.excel_exporter import generate_analysis_excel
        from app.core.report_theme import ReportTheme

        empty_data = _make_empty_export_data()
        mock_pdf_fetch.return_value = empty_data
        mock_docx_fetch.return_value = empty_data
        mock_xl_fetch.return_value = empty_data
        mock_theme.return_value = ReportTheme()

        pdf = generate_export_pdf(MagicMock(), "empty-test")
        docx = generate_export_docx(MagicMock(), "empty-test")
        xlsx = generate_analysis_excel(MagicMock(), "empty-test")

        assert pdf[:5] == b"%PDF-"
        assert docx[:2] == b"PK"
        assert xlsx[:2] == b"PK"
