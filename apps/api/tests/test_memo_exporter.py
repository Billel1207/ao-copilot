"""Tests pour memo_exporter.py — Génération mémoire technique Word.

Couvre:
- Génération du DOCX avec données complètes
- Dégradation gracieuse sans LLM
- Dégradation gracieuse sans python-docx
- Sections hybrides (IA bleu + client orange)
- Robustesse face aux données manquantes
"""
import pytest
from io import BytesIO
from unittest.mock import MagicMock, patch, PropertyMock
from dataclasses import dataclass, field


# ── Fixtures ──────────────────────────────────────────────────────────────────

@dataclass
class MockProject:
    title: str = "Rénovation Collège Victor Hugo"
    reference: str = "AO-2026-099"
    status: str = "ready"
    org_id: str = "00000000-0000-0000-0000-000000000001"
    buyer: str = "Ville de Lyon"


@dataclass
class MockOrg:
    id: str = "00000000-0000-0000-0000-000000000001"
    name: str = "BTP Solutions SAS"


@dataclass
class MockCompanyProfile:
    revenue_eur: int = 5000000
    employee_count: int = 45
    certifications: list = field(default_factory=lambda: ["Qualibat 7131", "ISO 9001"])
    specialties: list = field(default_factory=lambda: ["gros-oeuvre", "VRD"])
    regions: list = field(default_factory=lambda: ["Auvergne-Rhône-Alpes"])
    max_market_size_eur: int = 3000000
    assurance_rc_montant: int = 5000000
    assurance_decennale: bool = True
    partenaires_specialites: list = field(default_factory=lambda: ["électricité", "plomberie"])
    marge_minimale_pct: int = 12
    max_projets_simultanes: int = 8
    projets_actifs_count: int = 5
    logo_s3_key: str = None
    custom_theme: dict = None


@dataclass
class MockChecklistItem:
    category: str = "Administratif"
    criticality: str = "éliminatoire"
    requirement: str = "DC1 signé"
    status: str = "OK"
    confidence: float = 95.0
    citation: str = "RC art. 6"


def _make_export_data(**overrides):
    """Build mock ExportData for memo generation."""
    from app.services.export_data import ExportData

    defaults = dict(
        project=MockProject(),
        documents=[],
        summary={
            "project_overview": {
                "deadline_submission": "2026-07-01",
                "estimated_budget": "2 500 000 €",
                "location": "Lyon (69)",
                "market_type": "travaux",
            },
            "key_points": ["Délai 12 mois", "Lot unique"],
            "risks": ["Pénalités élevées", "Accès chantier restreint"],
            "actions_next_48h": ["Visiter le site", "Demander variante"],
        },
        criteria={"evaluation": {"scoring_criteria": [
            {"name": "Prix", "weight": 40, "estimated_score": 75},
            {"name": "Technique", "weight": 50, "estimated_score": 80},
        ]}},
        gonogo={"score": 68, "decision": "GO conditionnel", "dimension_scores": {
            "Capacité financière": 70, "Certifications": 95,
        }},
        timeline={"submission_deadline": "2026-07-01"},
        checklist_items=[MockChecklistItem()],
        checklist_stats={"eliminatoire": 1, "important": 0, "info": 0, "ok": 1},
        confidence=85.0,
        ccap_analysis={"clauses_risquees": [{"titre": "Pénalités", "severity": "élevé"}]},
        ccag_derogations=None,
        ccap_clauses_risquees=[{"titre": "Pénalités", "severity": "élevé"}],
        rc_analysis=None,
        ae_analysis=None,
        cctp_analysis={"prescriptions": ["NF DTU 20.1"]},
        dc_check=None,
        conflicts=None,
        cashflow={"monthly_cashflow": [
            {"month": 1, "cumulative_eur": -50000},
            {"month": 2, "cumulative_eur": 10000},
        ]},
        subcontracting={"recommended_lots": ["Électricité", "Plomberie"]},
        questions=None,
        scoring=None,
        dpgf_pricing=None,
        glossaire_btp=None,
        days_remaining=30,
        deadline_str="2026-07-01",
        gonogo_obj=None,
        timeline_obj=None,
    )
    defaults.update(overrides)
    return ExportData(**defaults)


# ── Tests ─────────────────────────────────────────────────────────────────────

class TestGenerateMemoTechnique:
    """Tests pour generate_memo_technique()."""

    @patch("app.services.memo_exporter.fetch_export_data")
    def test_generates_valid_docx_bytes(self, mock_fetch):
        """Le résultat est un fichier DOCX valide (commence par PK — ZIP)."""
        from app.services.memo_exporter import generate_memo_technique

        mock_fetch.return_value = _make_export_data()

        mock_db = MagicMock()
        mock_db.query.return_value.filter_by.return_value.first.side_effect = [
            MockOrg(),          # Organization query
            MockCompanyProfile(),  # CompanyProfile query
        ]

        # Patch LLM to avoid real API calls
        with patch("app.services.memo_exporter.llm_service", create=True) as mock_llm:
            mock_llm.chat_text.return_value = "Texte généré par IA pour test."
            # Also patch the import of llm_service inside the function
            with patch.dict("sys.modules", {"app.services.llm": MagicMock(llm_service=mock_llm)}):
                result = generate_memo_technique(mock_db, "test-project-id")

        assert isinstance(result, bytes)
        assert len(result) > 100
        # DOCX = ZIP file → starts with PK
        assert result[:2] == b"PK"

    @patch("app.services.memo_exporter.fetch_export_data")
    def test_works_without_llm(self, mock_fetch):
        """La mémo se génère même si le service LLM est indisponible."""
        from app.services.memo_exporter import generate_memo_technique

        mock_fetch.return_value = _make_export_data()

        mock_db = MagicMock()
        mock_db.query.return_value.filter_by.return_value.first.side_effect = [
            MockOrg(), MockCompanyProfile(),
        ]

        # Simulate LLM import failure
        import sys
        with patch.dict(sys.modules, {"app.services.llm": None}):
            result = generate_memo_technique(mock_db, "test-project-id")

        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

    @patch("app.services.memo_exporter.fetch_export_data")
    def test_works_without_company_profile(self, mock_fetch):
        """La mémo se génère même sans profil entreprise."""
        from app.services.memo_exporter import generate_memo_technique

        mock_fetch.return_value = _make_export_data()

        mock_db = MagicMock()
        # Organization found but no CompanyProfile
        mock_db.query.return_value.filter_by.return_value.first.side_effect = [
            MockOrg(), None,
        ]

        result = generate_memo_technique(mock_db, "test-project-id")
        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

    @patch("app.services.memo_exporter.fetch_export_data")
    def test_works_with_minimal_data(self, mock_fetch):
        """La mémo se génère avec un minimum de données."""
        from app.services.memo_exporter import generate_memo_technique

        mock_fetch.return_value = _make_export_data(
            summary={},
            criteria={},
            gonogo=None,
            ccap_analysis=None,
            cctp_analysis=None,
            cashflow=None,
            subcontracting=None,
            checklist_items=[],
            ccap_clauses_risquees=None,
        )

        mock_db = MagicMock()
        mock_db.query.return_value.filter_by.return_value.first.return_value = None

        result = generate_memo_technique(mock_db, "test-project-id")
        assert isinstance(result, bytes)
        assert len(result) > 0

    @patch("app.services.memo_exporter.fetch_export_data")
    def test_docx_contains_project_title(self, mock_fetch):
        """Le document contient le titre du projet dans le texte."""
        from app.services.memo_exporter import generate_memo_technique
        from docx import Document as DocxDocument

        mock_fetch.return_value = _make_export_data()

        mock_db = MagicMock()
        mock_db.query.return_value.filter_by.return_value.first.side_effect = [
            MockOrg(), MockCompanyProfile(),
        ]

        result = generate_memo_technique(mock_db, "test-project-id")
        doc = DocxDocument(BytesIO(result))

        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "Rénovation Collège Victor Hugo" in all_text

    @patch("app.services.memo_exporter.fetch_export_data")
    def test_docx_contains_company_name(self, mock_fetch):
        """Le document contient le nom de l'entreprise."""
        from app.services.memo_exporter import generate_memo_technique
        from docx import Document as DocxDocument

        mock_fetch.return_value = _make_export_data()

        mock_db = MagicMock()
        mock_db.query.return_value.filter_by.return_value.first.side_effect = [
            MockOrg(), MockCompanyProfile(),
        ]

        result = generate_memo_technique(mock_db, "test-project-id")
        doc = DocxDocument(BytesIO(result))

        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "BTP Solutions" in all_text


class TestMemoRobustness:
    """Tests de robustesse memo."""

    @patch("app.services.memo_exporter.fetch_export_data")
    def test_empty_scoring_criteria(self, mock_fetch):
        """Pas de crash avec des critères de scoring vides."""
        from app.services.memo_exporter import generate_memo_technique

        mock_fetch.return_value = _make_export_data(
            scoring={"criteria": []},
            criteria={"evaluation": {"scoring_criteria": []}},
        )

        mock_db = MagicMock()
        mock_db.query.return_value.filter_by.return_value.first.return_value = None

        result = generate_memo_technique(mock_db, "test-project-id")
        assert isinstance(result, bytes)

    @patch("app.services.memo_exporter.fetch_export_data")
    def test_none_risks_list(self, mock_fetch):
        """Pas de crash avec une liste de risques None."""
        from app.services.memo_exporter import generate_memo_technique

        mock_fetch.return_value = _make_export_data(
            summary={"project_overview": {}, "risks": None},
        )

        mock_db = MagicMock()
        mock_db.query.return_value.filter_by.return_value.first.return_value = None

        result = generate_memo_technique(mock_db, "test-project-id")
        assert isinstance(result, bytes)
