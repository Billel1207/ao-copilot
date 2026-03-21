"""Script pour générer le template Word docxtpl de la mémoire technique.

Usage:
    python -m app.templates.create_memo_template

Génère memo_technique_template.docx dans le même répertoire.
Ce template est utilisable avec docxtpl pour remplir les données dynamiquement.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def _set_cell_shading(cell, hex_color: str):
    """Apply background color to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_template():
    doc = Document()

    # ── Styles ────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x0F, 0x1B, 0x4C)
        hs.font.name = "Calibri"

    # ── Couverture ────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MEMOIRE TECHNIQUE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x0F, 0x1B, 0x4C)

    doc.add_paragraph()

    # Jinja2 placeholders for docxtpl
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("{{ project_title }}")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    doc.add_paragraph()

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    labels = ["Référence", "Acheteur", "Date limite", "Score Go/No-Go"]
    values = ["{{ reference }}", "{{ buyer }}", "{{ deadline }}", "{{ go_nogo_score }}/100"]
    for i, (label, value) in enumerate(zip(labels, values)):
        _set_cell_shading(table.rows[i].cells[0], "E2E8F0")
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].text = value

    doc.add_page_break()

    # ── Section 1: Introduction ───────────────────────────────────────────
    doc.add_heading("1. Introduction et compréhension du besoin", level=1)
    doc.add_paragraph("{{ narrative_intro }}")

    # ── Section 2: Présentation entreprise ────────────────────────────────
    doc.add_heading("2. Présentation de notre entreprise", level=1)
    doc.add_paragraph("{{ narrative_positioning }}")

    # ── Section 3: Radar chart ────────────────────────────────────────────
    doc.add_heading("3. Analyse de compatibilité", level=1)
    doc.add_paragraph("{{ radar_chart_placeholder }}")

    # ── Section 4: Méthodologie ──────────────────────────────────────────
    doc.add_heading("4. Méthodologie et organisation", level=1)
    doc.add_paragraph(
        "Détaillez ici votre méthodologie de travail, les moyens humains et "
        "matériels mobilisés, et l'organisation proposée."
    )

    # ── Section 5: Planning ──────────────────────────────────────────────
    doc.add_heading("5. Planning prévisionnel", level=1)
    doc.add_paragraph("{{ cashflow_chart_placeholder }}")

    # ── Section 6: Risques identifiés ────────────────────────────────────
    doc.add_heading("6. Risques identifiés et mesures de mitigation", level=1)

    risk_table = doc.add_table(rows=1, cols=3)
    risk_table.style = "Table Grid"
    headers = ["Risque", "Sévérité", "Mesure de mitigation"]
    for i, h in enumerate(headers):
        _set_cell_shading(risk_table.rows[0].cells[i], "0F1B4C")
        run = risk_table.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Jinja2 loop placeholder comment
    doc.add_paragraph(
        "{% for risk in risks %}\n"
        "{{ risk.titre }} | {{ risk.severity }} | {{ risk.detail }}\n"
        "{% endfor %}"
    )

    # ── Section 7: Plan d'action ─────────────────────────────────────────
    doc.add_heading("7. Plan d'action 48h", level=1)
    doc.add_paragraph("{{ narrative_action_plan }}")

    # ── Footer ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("AO COPILOT — Document confidentiel")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # ── Sauvegarder ───────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(__file__), "memo_technique_template.docx")
    doc.save(output_path)
    print(f"Template sauvegardé : {output_path}")
    return output_path


if __name__ == "__main__":
    create_template()
