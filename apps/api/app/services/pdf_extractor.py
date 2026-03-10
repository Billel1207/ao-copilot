"""Extraction texte multi-format : PDF, images (JPEG/PNG/TIFF), Word (.docx).

Inclut le scoring de qualité OCR (confidence 0-100) pour chaque page,
et la détection automatique du type de document à partir du contenu.
"""
from dataclasses import dataclass, field
from pathlib import Path
import re
import fitz  # PyMuPDF
import pytesseract
import structlog
from PIL import Image
import io

logger = structlog.get_logger(__name__)

# Marqueur inséré dans le texte si l'OCR échoue — permet à l'analyse IA
# de détecter que cette page n'a pas été extraite correctement.
OCR_FAILURE_MARKER = "[TEXTE_NON_EXTRACTIBLE]"


@dataclass
class PageResult:
    page_num: int
    raw_text: str
    char_count: int
    section: str | None
    needs_ocr: bool
    ocr_confidence: float | None = None  # Score OCR 0-100 (None si pas d'OCR)


SECTION_PATTERNS = [
    (r"(?i)^\s*(article\s+\d+|chapitre\s+\d+|section\s+\d+)", "section"),
    (r"(?i)^(RC|règlement de consultation)", "RC"),
    (r"(?i)^(CCTP|cahier des clauses techniques)", "CCTP"),
    (r"(?i)^(CCAP|cahier des clauses administratives)", "CCAP"),
    (r"(?i)^(DPGF|décomposition du prix)", "DPGF"),
    (r"(?i)^(BPU|bordereau des prix)", "BPU"),
    (r"(?i)^(acte d'engagement|acte engagement)", "AE"),
]

OCR_THRESHOLD = 50  # chars/page sous lesquels on déclenche l'OCR


def detect_section(text: str) -> str | None:
    for pattern, label in SECTION_PATTERNS:
        if re.search(pattern, text[:200]):
            return label
    return None


def extract_pages(pdf_bytes: bytes) -> list[PageResult]:
    """Extrait le texte page par page depuis un PDF."""
    results: list[PageResult] = []

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        raise ValueError(f"Impossible d'ouvrir le PDF : {e}")

    for page_num in range(len(doc)):
        page = doc[page_num]
        raw_text = page.get_text("text")
        char_count = len(raw_text.strip())
        needs_ocr = char_count < OCR_THRESHOLD
        ocr_confidence = None

        if needs_ocr:
            raw_text, ocr_confidence = _ocr_page_with_confidence(page)
            char_count = len(raw_text.strip())

        section = detect_section(raw_text)

        results.append(PageResult(
            page_num=page_num + 1,
            raw_text=raw_text.strip(),
            char_count=char_count,
            section=section,
            needs_ocr=needs_ocr,
            ocr_confidence=ocr_confidence,
        ))

    doc.close()
    return results


def _ocr_page_with_confidence(page: fitz.Page) -> tuple[str, float]:
    """OCR via Tesseract avec scoring de confiance par mot.

    Utilise image_to_data() au lieu de image_to_string() pour obtenir
    la confiance de chaque mot détecté (0-100).

    Returns:
        Tuple (texte_extrait, confiance_moyenne_0_100).
        Retourne (OCR_FAILURE_MARKER, 0.0) si Tesseract échoue.
    """
    try:
        mat = fitz.Matrix(2.0, 2.0)  # zoom x2 pour meilleure qualité OCR
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        image = Image.open(io.BytesIO(img_bytes))

        # Utiliser image_to_data pour obtenir la confiance par mot
        data = pytesseract.image_to_data(
            image,
            lang="fra+eng",
            output_type=pytesseract.Output.DICT,
            config="--oem 3 --psm 6",
        )

        # Reconstruire le texte et calculer la confiance moyenne
        words = []
        confidences = []
        for i, word in enumerate(data["text"]):
            word = word.strip()
            if word:
                words.append(word)
                conf = int(data["conf"][i])
                if conf >= 0:  # -1 = pas de confiance
                    confidences.append(conf)

        text = " ".join(words) if words else OCR_FAILURE_MARKER
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

        if avg_confidence < 50:
            logger.warning(
                "ocr_low_confidence",
                page_num=getattr(page, "number", "?") + 1,
                confidence=round(avg_confidence, 1),
                word_count=len(words),
            )

        return text, round(avg_confidence, 1)

    except Exception as exc:
        logger.warning(
            "ocr_page_failed",
            page_num=getattr(page, "number", "?"),
            error=str(exc),
            hint="Vérifiez que Tesseract est installé et que tesseract-ocr-fra est disponible",
        )
        return OCR_FAILURE_MARKER, 0.0


def _ocr_page(page: fitz.Page) -> str:
    """OCR simple (compatibilité) — utilise _ocr_page_with_confidence en interne."""
    text, _ = _ocr_page_with_confidence(page)
    return text


def has_sufficient_text(pages: list[PageResult], threshold: float = 0.3) -> bool:
    """Vérifie si au moins 30% des pages ont du texte suffisant."""
    if not pages:
        return False
    ok = sum(1 for p in pages if p.char_count >= OCR_THRESHOLD)
    return (ok / len(pages)) >= threshold


def clean_text(text: str) -> str:
    """Nettoyage basique : suppression répétitions headers/footers."""
    lines = text.split("\n")
    # Supprimer lignes très courtes répétitives (headers/footers)
    cleaned = [line for line in lines if len(line.strip()) > 3]
    return "\n".join(cleaned)


# ── Qualité OCR document ─────────────────────────────────────────────────

def compute_document_ocr_quality(pages: list[PageResult]) -> dict:
    """Calcule la qualité OCR globale d'un document à partir des pages.

    Returns:
        Dict avec :
        - ocr_score: float (0-100) — score moyen des pages OCR
        - pages_ocr_count: int — nombre de pages ayant nécessité l'OCR
        - pages_low_quality: int — pages avec confiance < 70
        - warning: str | None — message d'alerte si qualité insuffisante
    """
    ocr_pages = [p for p in pages if p.needs_ocr and p.ocr_confidence is not None]

    if not ocr_pages:
        return {
            "ocr_score": None,
            "pages_ocr_count": 0,
            "pages_low_quality": 0,
            "warning": None,
        }

    scores = [p.ocr_confidence for p in ocr_pages if p.ocr_confidence is not None]
    avg_score = sum(scores) / len(scores) if scores else 0.0
    pages_low = sum(1 for s in scores if s < 70)

    warning = None
    if avg_score < 50:
        warning = (
            f"Qualité OCR très faible ({avg_score:.0f}/100). "
            f"Le document est probablement un scan de mauvaise qualité. "
            f"Les résultats d'analyse peuvent être imprécis."
        )
    elif avg_score < 70:
        warning = (
            f"Qualité OCR moyenne ({avg_score:.0f}/100). "
            f"{pages_low} page(s) avec OCR insuffisant. "
            f"Certaines informations peuvent être manquantes."
        )

    return {
        "ocr_score": round(avg_score, 1),
        "pages_ocr_count": len(ocr_pages),
        "pages_low_quality": pages_low,
        "warning": warning,
    }


# ── Détection type document par contenu ──────────────────────────────────

# Patterns structurels pour identifier le type de document
# dans les premiers 3000 caractères du texte
DOC_TYPE_PATTERNS: dict[str, list[str]] = {
    "RC": [
        r"(?i)r[eè]glement\s+de\s+consultation",
        r"(?i)objet\s+de\s+la\s+consultation",
        r"(?i)crit[eè]res?\s+d['']attribution",
        r"(?i)modalit[eé]s?\s+de\s+remise\s+des\s+offres",
        r"(?i)date\s+limite\s+de\s+(r[eé]ception|remise)",
    ],
    "CCTP": [
        r"(?i)cahier\s+des\s+clauses\s+techniques\s+particuli[eè]res",
        r"(?i)C\.?C\.?T\.?P",
        r"(?i)prescriptions\s+techniques",
        r"(?i)sp[eé]cifications?\s+techniques?",
        r"(?i)description\s+des\s+(travaux|prestations|ouvrages)",
    ],
    "CCAP": [
        r"(?i)cahier\s+des\s+clauses\s+administratives\s+particuli[eè]res",
        r"(?i)C\.?C\.?A\.?P",
        r"(?i)p[eé]nalit[eé]s?\s+de\s+retard",
        r"(?i)retenue\s+de\s+garantie",
        r"(?i)d[eé]lai\s+de\s+paiement",
    ],
    "DPGF": [
        r"(?i)d[eé]composition\s+du\s+prix\s+global",
        r"(?i)D\.?P\.?G\.?F",
        r"(?i)prix\s+forfaitaire",
        r"(?i)d[eé]signation\s+des?\s+(prestations?|ouvrages?)",
    ],
    "BPU": [
        r"(?i)bordereau\s+des?\s+prix\s+unitaires?",
        r"(?i)B\.?P\.?U",
        r"(?i)prix\s+unitaire",
        r"(?i)unit[eé]\s+de\s+mesure",
    ],
    "AE": [
        r"(?i)acte\s+d['']engagement",
        r"(?i)A\.?E\.?\s",
        r"(?i)ATTRI\s*1",
        r"(?i)le\s+candidat\s+s['']engage",
        r"(?i)montant\s+total\s+du\s+march[eé]",
    ],
}


def detect_doc_type_from_content(text: str) -> str | None:
    """Détecte le type de document DCE à partir de son contenu textuel.

    Analyse les 3000 premiers caractères du texte pour identifier le type
    de document via des patterns structurels.

    Args:
        text: Texte extrait du document.

    Returns:
        Type détecté ("RC", "CCTP", "CCAP", "DPGF", "BPU", "AE") ou None si non identifiable.
    """
    if not text or len(text) < 20:
        return None

    snippet = text[:3000]
    scores: dict[str, int] = {}

    for doc_type, patterns in DOC_TYPE_PATTERNS.items():
        count = 0
        for pattern in patterns:
            if re.search(pattern, snippet):
                count += 1
        if count > 0:
            scores[doc_type] = count

    if not scores:
        return None

    # Retourner le type avec le plus de matches
    best = max(scores, key=scores.get)  # type: ignore[arg-type]

    # Exiger au moins 2 matches pour éviter les faux positifs
    if scores[best] >= 2:
        logger.info(
            "doc_type_detected_from_content",
            detected_type=best,
            match_count=scores[best],
        )
        return best

    return None


# ── Extracteurs multi-format ─────────────────────────────────────────────

def extract_image_file(image_bytes: bytes, filename: str) -> list[PageResult]:
    """Extrait le texte d'une image (JPEG, PNG, TIFF) via OCR Tesseract avec confiance."""
    try:
        image = Image.open(io.BytesIO(image_bytes))

        # Utiliser image_to_data pour confiance
        data = pytesseract.image_to_data(
            image,
            lang="fra+eng",
            output_type=pytesseract.Output.DICT,
            config="--oem 3 --psm 6",
        )

        words = []
        confidences = []
        for i, word in enumerate(data["text"]):
            word = word.strip()
            if word:
                words.append(word)
                conf = int(data["conf"][i])
                if conf >= 0:
                    confidences.append(conf)

        text = " ".join(words) if words else OCR_FAILURE_MARKER
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

        return [PageResult(
            page_num=1,
            raw_text=text.strip(),
            char_count=len(text.strip()),
            section=detect_section(text),
            needs_ocr=True,
            ocr_confidence=round(avg_confidence, 1),
        )]
    except Exception as exc:
        logger.warning("image_extraction_failed", filename=filename, error=str(exc))
        return [PageResult(
            page_num=1,
            raw_text=OCR_FAILURE_MARKER,
            char_count=0,
            section=None,
            needs_ocr=True,
            ocr_confidence=0.0,
        )]


def extract_docx_file(docx_bytes: bytes) -> list[PageResult]:
    """Extrait le texte d'un document Word (.docx) via python-docx."""
    try:
        from docx import Document  # python-docx
        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Inclure aussi les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        full_text = "\n".join(paragraphs) or OCR_FAILURE_MARKER
        return [PageResult(
            page_num=1,
            raw_text=full_text,
            char_count=len(full_text),
            section=detect_section(full_text),
            needs_ocr=False,
            ocr_confidence=None,  # Pas d'OCR pour DOCX
        )]
    except Exception as exc:
        logger.warning("docx_extraction_failed", error=str(exc))
        return [PageResult(
            page_num=1,
            raw_text=OCR_FAILURE_MARKER,
            char_count=0,
            section=None,
            needs_ocr=False,
            ocr_confidence=None,
        )]


IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tiff", ".tif"}


def extract_document(file_bytes: bytes, filename: str) -> list[PageResult]:
    """Dispatch vers le bon extracteur selon l'extension du fichier."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_pages(file_bytes)
    elif ext in IMAGE_EXTENSIONS:
        return extract_image_file(file_bytes, filename)
    elif ext == ".docx":
        return extract_docx_file(file_bytes)
    else:
        raise ValueError(f"Format non supporté : {ext}")
