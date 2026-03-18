"""Génération PDF d'export avec xhtml2pdf + Word avec python-docx."""
import uuid
import structlog
from datetime import datetime
from io import BytesIO
from jinja2 import Environment, BaseLoader
from sqlalchemy.orm import Session

from app.core.report_theme import get_theme

logger = structlog.get_logger(__name__)

from app.models.project import AoProject
from app.models.document import AoDocument
from app.models.analysis import ExtractionResult, ChecklistItem, CriteriaItem

EXPORT_TEMPLATE = """
<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<style>
  /* ══════════════ TYPOGRAPHIE & BASE ══════════════ */
  @page {
    margin: 12mm 14mm 20mm 14mm;
    size: A4;
    @bottom-left {
      content: "AO Copilot — Confidentiel";
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-size: 7.5px; color: #94A3B8;
    }
    @bottom-center {
      content: counter(page) " / " counter(pages);
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-size: 7.5px; color: #64748B; font-weight: 600;
    }
    @bottom-right {
      content: "Généré le {{ generation_date }}";
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      font-size: 7.5px; color: #94A3B8;
    }
  }
  @page :first {
    @bottom-left { content: ""; }
    @bottom-center { content: ""; }
    @bottom-right { content: ""; }
  }
  body {
    font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
    font-size: 11px; color: #1a1a2e; line-height: 1.55;
    margin: 0; padding: 0;
  }
  .page { padding: 8px 0; }

  /* ══════════════ TITRES ══════════════ */
  h1 {
    font-size: 20px; color: #0f1b4c; font-weight: 700;
    border-bottom: 2px solid #2563eb; padding-bottom: 6px;
    margin-top: 0; margin-bottom: 14px; letter-spacing: 0.3px;
  }
  h2 {
    font-size: 14px; color: #0f1b4c; font-weight: 700;
    margin-top: 22px; margin-bottom: 8px;
    border-bottom: 1px solid #cbd5e1; padding-bottom: 3px;
  }
  h3 {
    font-size: 12px; color: #334155; font-weight: 600;
    margin-top: 16px; margin-bottom: 6px;
  }

  /* ══════════════ BADGES ══════════════ */
  .badge {
    display: inline-block; padding: 2px 6px; font-size: 8px;
    font-weight: 700; letter-spacing: 0.2px;
  }
  .badge-red { background: #fde8e8; color: #b91c1c; border: 1px solid #f5c6c6; }
  .badge-yellow { background: #fef3c7; color: #b45309; border: 1px solid #fde68a; }
  .badge-green { background: #d1fae5; color: #047857; border: 1px solid #a7f3d0; }
  .badge-blue { background: #dbeafe; color: #1d4ed8; border: 1px solid #bfdbfe; }
  .badge-gray { background: #f1f5f9; color: #475569; border: 1px solid #e2e8f0; }

  /* ══════════════ TABLEAUX ══════════════ */
  table {
    width: 100%; border-collapse: collapse; margin-top: 10px; margin-bottom: 10px;
  }
  th {
    background: #0f1b4c; color: #ffffff; font-weight: 600;
    text-align: left; padding: 5px 6px; font-size: 9px;
    letter-spacing: 0.3px; text-transform: uppercase;
    border: 1px solid #1e3a8a;
    word-wrap: break-word; overflow: hidden;
  }
  td {
    padding: 4px 6px; border: 1px solid #e2e8f0;
    vertical-align: top; font-size: 9px; line-height: 1.4;
    word-wrap: break-word; overflow: hidden;
  }
  tr:nth-child(even) td { background: #f8fafc; }
  .citation { font-style: italic; color: #64748b; font-size: 7px; margin-top: 2px; }
  .risk-high td { background: #fef2f2; }
  .risk-medium td { background: #fffbeb; }

  /* ══════════════ ENCADRÉS ══════════════ */
  .summary-box {
    background: #eff6ff; border-left: 5px solid #2563eb;
    padding: 12px 16px; margin: 12px 0; font-size: 11px; line-height: 1.6;
  }
  .financial-box {
    background: #ecfdf5; border-left: 5px solid #10b981;
    padding: 12px 16px; margin: 12px 0; font-size: 11px;
  }
  .warning-box {
    background: #fef2f2; border-left: 5px solid #ef4444;
    padding: 10px 14px; font-size: 10px; color: #991b1b; margin: 10px 0;
  }
  .info-box {
    background: #eff6ff; border-left: 5px solid #3b82f6;
    padding: 10px 14px; font-size: 10px; color: #1e3a8a; margin: 10px 0;
  }

  /* ══════════════ GO/NO-GO BOX ══════════════ */
  .gonogo-box { padding: 18px; margin: 14px 0; text-align: center; }
  .gonogo-go { background: #d1fae5; border: 3px solid #059669; }
  .gonogo-attention { background: #fef3c7; border: 3px solid #d97706; }
  .gonogo-nogo { background: #fde8e8; border: 3px solid #dc2626; }
  .gonogo-score { font-size: 36px; font-weight: 800; letter-spacing: -0.5px; }
  .gonogo-label { font-size: 16px; font-weight: 700; margin-top: 4px; letter-spacing: 1px; }

  /* ══════════════ KPI CARDS ══════════════ */
  .stat-grid { display: table; width: 100%; margin: 14px 0; }
  .stat-cell {
    display: table-cell; text-align: center;
    padding: 14px 8px; width: 25%;
    border: 1px solid #e2e8f0; background: #f8fafc;
  }
  .stat-number { font-size: 24px; font-weight: 800; color: #0f1b4c; }
  .stat-label { font-size: 9px; color: #64748b; margin-top: 4px; text-transform: uppercase; letter-spacing: 0.5px; }

  /* ══════════════ BARRE DE CONFIANCE ══════════════ */
  .confidence-bar { background: #e2e8f0; height: 10px; margin-top: 6px; }
  .confidence-fill { height: 10px; }

  /* ══════════════ PAGE DE COUVERTURE ══════════════ */
  .cover-page { text-align: center; padding-top: 60px; }
  .cover-brand {
    font-size: 13px; color: #2563eb; letter-spacing: 4px;
    text-transform: uppercase; font-weight: 700; margin-bottom: 6px;
  }
  .cover-line {
    width: 80px; height: 3px; background: #2563eb;
    margin: 12px auto;
  }
  .cover-title {
    font-size: 30px; color: #0f1b4c; font-weight: 800;
    margin-bottom: 6px; line-height: 1.25; letter-spacing: -0.3px;
  }
  .cover-subtitle { font-size: 16px; color: #475569; margin-bottom: 30px; font-weight: 400; }
  .cover-info { font-size: 12px; color: #64748b; margin: 5px 0; line-height: 1.5; }
  .cover-badge {
    display: inline-block; padding: 8px 28px;
    font-size: 14px; font-weight: 700; margin-top: 16px;
    letter-spacing: 0.5px;
  }
  .cover-meta {
    margin-top: 50px; padding-top: 16px;
    border-top: 1px solid #cbd5e1;
    color: #94a3b8; font-size: 10px;
  }

  /* ══════════════ SOMMAIRE ══════════════ */
  .toc { margin: 16px 0; }
  .toc-item {
    padding: 6px 0 6px 8px; border-bottom: 1px dotted #cbd5e1;
    font-size: 11px; line-height: 1.5;
  }
  .toc-num {
    display: inline-block; width: 28px; font-weight: 700;
    color: #2563eb; font-size: 11px;
  }
  .toc-text { color: #334155; }

  /* ══════════════ DISCLAIMER & FOOTER ══════════════ */
  .disclaimer {
    background: #fffbeb; border: 1px solid #fbbf24;
    padding: 12px 16px; font-size: 9px; color: #78350f;
    margin-top: 24px;
  }
  .footer {
    text-align: center; color: #94a3b8; font-size: 8px;
    margin-top: 16px; padding-top: 8px;
    border-top: 1px solid #e2e8f0;
  }
  .page-break { page-break-before: always; }

  /* ══════════════ SECTION HEADER (numérotation élégante) ══════════════ */
  .section-num {
    display: inline-block; background: #0f1b4c; color: #ffffff;
    padding: 2px 8px; font-size: 11px; font-weight: 700;
    margin-right: 6px;
  }

  /* ══════════════ PARTIE SEPARATOR ══════════════ */
  .part-header {
    background: #0f1b4c; color: #ffffff; padding: 12px 18px;
    font-size: 14px; font-weight: 800; letter-spacing: 1.5px;
    text-transform: uppercase; margin: 0; page-break-before: always;
  }
  .part-header-sub {
    font-size: 10px; font-weight: 400; letter-spacing: 0.5px;
    color: #94a3b8; margin-top: 4px;
  }

  /* ══════════════ RISK GAUGE BAR ══════════════ */
  .risk-gauge {
    background: #e2e8f0; height: 14px; margin: 8px 0; position: relative;
  }
  .risk-gauge-fill {
    height: 14px; display: inline-block;
  }
  .risk-gauge-label {
    font-size: 8px; font-weight: 700; color: #ffffff;
    padding: 0 6px; line-height: 14px;
  }

  /* ══════════════ STATUS BADGE (DPGF) ══════════════ */
  .status-sous { background: #fde8e8; color: #b91c1c; border: 1px solid #f5c6c6; font-weight: 700; padding: 2px 6px; font-size: 8px; }
  .status-sur { background: #fef3c7; color: #b45309; border: 1px solid #fde68a; font-weight: 700; padding: 2px 6px; font-size: 8px; }
  .status-normal { background: #d1fae5; color: #047857; border: 1px solid #a7f3d0; font-weight: 700; padding: 2px 6px; font-size: 8px; }
  .status-inconnu { background: #f1f5f9; color: #475569; border: 1px solid #e2e8f0; font-weight: 700; padding: 2px 6px; font-size: 8px; }
</style>
</head>
<body>

<!-- ═══════════ PAGE DE COUVERTURE ═══════════ -->
<div class="page cover-page">
  <div class="cover-brand">AO COPILOT</div>
  <div class="cover-line"></div>
  <div style="font-size: 12px; color: #64748b; letter-spacing: 3px; text-transform: uppercase; margin-bottom: 30px;">Rapport d'analyse DCE</div>

  <div class="cover-title">{{ project.title }}</div>
  <div class="cover-subtitle">{{ project.buyer or 'Acheteur public' }}</div>

  <!-- Infos projet -->
  <div style="margin: 30px 0;">
    <div class="cover-info"><strong>Reference :</strong> {{ project.reference or 'N/A' }}</div>
    {% if summary and summary.project_overview.deadline_submission %}
    <div class="cover-info"><strong>Date limite :</strong> {{ summary.project_overview.deadline_submission|datefr }}</div>
    {% endif %}
    {% if summary and summary.project_overview.estimated_budget %}
    <div class="cover-info"><strong>Budget estime :</strong> {{ summary.project_overview.estimated_budget }}</div>
    {% endif %}
    <div class="cover-info"><strong>Lieu :</strong> {{ summary.project_overview.location if summary else 'N/A' }}</div>
  </div>

  {% if days_remaining is not none %}
  <div style="margin-top: 24px;">
    <div style="display: inline-block; padding: 12px 32px; font-size: 20px; font-weight: 800; letter-spacing: 1px;
      {% if days_remaining <= 3 %}background: #fde8e8; color: #b91c1c; border: 3px solid #dc2626;
      {% elif days_remaining <= 7 %}background: #fef3c7; color: #b45309; border: 3px solid #d97706;
      {% elif days_remaining <= 14 %}background: #dbeafe; color: #1d4ed8; border: 3px solid #2563eb;
      {% else %}background: #d1fae5; color: #047857; border: 3px solid #059669;{% endif %}">
      {% if days_remaining < 0 %}EXPIRE (J{{ days_remaining }})
      {% elif days_remaining == 0 %}DERNIER JOUR
      {% else %}J-{{ days_remaining }}{% endif %}
    </div>
    {% if days_remaining >= 0 %}
    <div class="cover-info" style="margin-top: 8px;">{{ days_remaining }} jour{{ 's' if days_remaining > 1 else '' }} restant{{ 's' if days_remaining > 1 else '' }}</div>
    {% else %}
    <div class="cover-info" style="margin-top: 8px; color: #dc2626; font-weight: 600;">Date limite depassee depuis {{ -days_remaining }} jour{{ 's' if -days_remaining > 1 else '' }}</div>
    {% endif %}
  </div>
  {% endif %}

  {% if gonogo %}
  <div style="margin-top: 24px;">
    {% set reco = gonogo.recommendation|upper if gonogo.recommendation else 'ATTENTION' %}
    <div class="cover-badge {% if reco == 'GO' %}badge-green{% elif reco == 'NO-GO' %}badge-red{% else %}badge-yellow{% endif %}">
      RECOMMANDATION : {{ reco }}
    </div>
    <div class="cover-info" style="margin-top: 8px; font-size: 13px;">Score Go/No-Go : <strong style="color: #0f1b4c;">{{ gonogo.score }}/100</strong></div>
  </div>
  {% endif %}

  <div class="cover-meta">
    Genere le {{ generated_at }} | Confiance IA : {{ "%.0f"|format(confidence * 100) if confidence else 'N/A' }}%
  </div>
</div>

<!-- ═══════════ SOMMAIRE ═══════════ -->
<div class="page page-break">
  <h1>Sommaire</h1>
  <div class="toc">
    <!-- PARTIE I : DÉCISION STRATÉGIQUE -->
    <div class="toc-item" style="background: #f0f4ff; font-weight: 700; padding-left: 4px; border-bottom: 2px solid #2563eb;">PARTIE I — DECISION STRATEGIQUE</div>
    <div class="toc-item"><span class="toc-num">1.</span> <span class="toc-text">Synthese decisionnelle — Recommandation Go/No-Go et indicateurs cles</span></div>
    {% if rc_analysis %}
    <div class="toc-item"><span class="toc-num">2.</span> <span class="toc-text">Fiche signaletique du marche — Procedure, allotissement, groupement, CCAG</span></div>
    {% endif %}
    {% if summary %}
    <div class="toc-item"><span class="toc-num">3.</span> <span class="toc-text">Resume executif — Objet du marche, points cles extraits</span></div>
    {% endif %}

    <!-- PARTIE II : CONFORMITÉ JURIDIQUE -->
    <div class="toc-item" style="background: #f0f4ff; font-weight: 700; padding-left: 4px; border-bottom: 2px solid #2563eb;">PARTIE II — CONFORMITE JURIDIQUE</div>
    {% if dc_check %}
    <div class="toc-item"><span class="toc-num">4.</span> <span class="toc-text">Verification administrative DC — Attestations, certifications, formulaires</span></div>
    {% endif %}
    {% if checklist_items %}
    <div class="toc-item"><span class="toc-num">5.</span> <span class="toc-text">Checklist de conformite — {{ checklist_items|length }} exigences ({{ checklist_stats.eliminatoire }} eliminatoires)</span></div>
    {% endif %}
    {% if ccag_derogations %}
    <div class="toc-item"><span class="toc-num">6.</span> <span class="toc-text">Derogations CCAG-Travaux 2021 — {{ ccag_derogations|length }} derogations detectees</span></div>
    {% endif %}
    {% if ccap_clauses_risquees %}
    <div class="toc-item"><span class="toc-num">7.</span> <span class="toc-text">Clauses risquees CCAP — {{ ccap_clauses_risquees|length }} clauses a risque</span></div>
    {% endif %}

    <!-- PARTIE III : ANALYSE FINANCIÈRE -->
    <div class="toc-item" style="background: #f0f4ff; font-weight: 700; padding-left: 4px; border-bottom: 2px solid #2563eb;">PARTIE III — ANALYSE FINANCIERE</div>
    {% if ae_analysis %}
    <div class="toc-item"><span class="toc-num">8.</span> <span class="toc-text">Analyse Acte d'Engagement — Prix, penalites, garanties, risques contractuels</span></div>
    {% endif %}
    <div class="toc-item"><span class="toc-num">9.</span> <span class="toc-text">Synthese financiere — Montants, avance, penalites, revision des prix</span></div>
    {% if dpgf_pricing %}
    <div class="toc-item"><span class="toc-num">10.</span> <span class="toc-text">Benchmark tarifaire DPGF — {{ dpgf_pricing|length }} postes analyses vs referentiel</span></div>
    {% endif %}
    {% if cashflow_data %}
    <div class="toc-item"><span class="toc-num">11.</span> <span class="toc-text">Simulation tresorerie et BFR — Impact financier previsionnel</span></div>
    {% endif %}

    <!-- PARTIE IV : ANALYSE TECHNIQUE -->
    <div class="toc-item" style="background: #f0f4ff; font-weight: 700; padding-left: 4px; border-bottom: 2px solid #2563eb;">PARTIE IV — ANALYSE TECHNIQUE</div>
    {% if cctp_analysis %}
    <div class="toc-item"><span class="toc-num">12.</span> <span class="toc-text">Analyse technique CCTP — Exigences techniques, normes DTU, contradictions</span></div>
    {% endif %}
    {% if conflicts %}
    <div class="toc-item"><span class="toc-num">13.</span> <span class="toc-text">Contradictions inter-documents — {{ (conflicts.nb_total or (conflicts.conflicts or [])|length) }} incoherences DCE</span></div>
    {% endif %}
    {% if subcontracting %}
    <div class="toc-item"><span class="toc-num">14.</span> <span class="toc-text">Analyse sous-traitance — Obligations, risques, conformite loi 1975</span></div>
    {% endif %}

    <!-- PARTIE V : SCORING & STRATÉGIE -->
    <div class="toc-item" style="background: #f0f4ff; font-weight: 700; padding-left: 4px; border-bottom: 2px solid #2563eb;">PARTIE V — SCORING ET STRATEGIE</div>
    {% if criteria %}
    <div class="toc-item"><span class="toc-num">15.</span> <span class="toc-text">Criteres d'attribution — Conditions d'eligibilite et grille de notation</span></div>
    {% endif %}
    {% if scoring_simulation %}
    <div class="toc-item"><span class="toc-num">16.</span> <span class="toc-text">Simulation note acheteur — Note estimee {{ "%.1f"|format(scoring_simulation.total_score or scoring_simulation.note_globale_estimee or 0) }}/20 et leviers</span></div>
    {% endif %}
    {% if gonogo %}
    <div class="toc-item"><span class="toc-num">17.</span> <span class="toc-text">Go/No-Go detaille — Analyse 9 dimensions avec gaps et forces profil</span></div>
    {% endif %}

    <!-- PARTIE VI : PLAN D'ACTION -->
    <div class="toc-item" style="background: #f0f4ff; font-weight: 700; padding-left: 4px; border-bottom: 2px solid #2563eb;">PARTIE VI — PLAN D'ACTION</div>
    {% if summary %}
    <div class="toc-item"><span class="toc-num">18.</span> <span class="toc-text">Analyse des risques — {{ (summary.risks or [])|length }} risques identifies et plan d'actions 48h</span></div>
    {% endif %}
    {% if questions_list %}
    <div class="toc-item"><span class="toc-num">19.</span> <span class="toc-text">Questions prioritaires pour l'acheteur — {{ questions_list|length }} questions a poser</span></div>
    {% endif %}
    {% if timeline %}
    <div class="toc-item"><span class="toc-num">20.</span> <span class="toc-text">Calendrier et dates cles — Echeances de soumission et d'execution</span></div>
    {% endif %}

    <!-- ANNEXES -->
    <div class="toc-item" style="background: #f0f4ff; font-weight: 700; padding-left: 4px; border-bottom: 2px solid #2563eb;">ANNEXES</div>
    {% if documents_inventory %}
    <div class="toc-item"><span class="toc-num">A1.</span> <span class="toc-text">Inventaire des documents — {{ documents_inventory|length }} pieces DCE analysees</span></div>
    {% endif %}
    {% if glossaire_btp %}
    <div class="toc-item"><span class="toc-num">A2.</span> <span class="toc-text">Glossaire BTP — {{ glossaire_btp|length }} termes cles</span></div>
    {% endif %}
    <div class="toc-item"><span class="toc-num">A3.</span> <span class="toc-text">Avertissement IA et mentions legales</span></div>
  </div>

  <div class="info-box" style="margin-top: 20px;">
    <strong>Guide de lecture rapide :</strong><br>
    <span style="color: #2563eb; font-weight: 700;">5 min</span> — Couverture + Synthese decisionnelle (pages 1-2)<br>
    <span style="color: #2563eb; font-weight: 700;">15 min</span> — + Resume executif + Risques + Calendrier<br>
    <span style="color: #2563eb; font-weight: 700;">30 min</span> — Rapport complet avec checklist et criteres detailles
  </div>
</div>

<!-- ═══════════════════════════════════════════════════════ -->
<!-- PARTIE I : DECISION STRATEGIQUE -->
<!-- ═══════════════════════════════════════════════════════ -->
<div class="part-header">
  PARTIE I — Decision strategique
  <div class="part-header-sub">Recommandation Go/No-Go, fiche marche et resume executif</div>
</div>

<!-- ═══════════ SYNTHÈSE DÉCISIONNELLE (1 page) ═══════════ -->
<div class="page page-break">
  <h1><span class="section-num">1</span> Synthese decisionnelle</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Vue d'ensemble en 1 page pour les decideurs (DG, Directeur commercial, Responsable AO)</p>

  <!-- Score Go/No-Go -->
  {% if gonogo %}
  {% set reco = gonogo.recommendation|upper if gonogo.recommendation else 'ATTENTION' %}
  <div class="gonogo-box {% if reco == 'GO' %}gonogo-go{% elif reco == 'NO-GO' %}gonogo-nogo{% else %}gonogo-attention{% endif %}">
    <div class="gonogo-score">{{ gonogo.score }}/100</div>
    <div class="gonogo-label">{{ reco }}</div>
    <div style="font-size: 10px; margin-top: 4px;">{{ gonogo.summary }}</div>
  </div>
  {% endif %}

  <!-- Indicateurs clés -->
  <div class="stat-grid">
    <div class="stat-cell">
      <div class="stat-number">{{ checklist_stats.eliminatoire }}</div>
      <div class="stat-label">Éliminatoires</div>
    </div>
    <div class="stat-cell">
      <div class="stat-number">{{ checklist_stats.important }}</div>
      <div class="stat-label">Importants</div>
    </div>
    <div class="stat-cell">
      <div class="stat-number">{{ (summary.risks or [])|length if summary else 0 }}</div>
      <div class="stat-label">Risques identifiés</div>
    </div>
    <div class="stat-cell">
      <div class="stat-number">{{ (summary.actions_next_48h or [])|length if summary else 0 }}</div>
      <div class="stat-label">Actions à mener</div>
    </div>
  </div>

  <!-- Top 3 risques -->
  {% if summary and summary.risks %}
  <h3>Top 3 risques critiques</h3>
  <table>
    <tr><th style="width:25%">Risque</th><th style="width:12%">Severite</th><th style="width:63%">Impact</th></tr>
    {% for r in (summary.risks or [])[:3] %}
    <tr class="risk-{{ r.severity }}">
      <td><strong>{{ r.risk|trunc(50) }}</strong></td>
      <td><span class="badge {% if r.severity == 'high' %}badge-red{% elif r.severity == 'medium' %}badge-yellow{% else %}badge-gray{% endif %}">{{ (r.severity or '')|upper }}</span></td>
      <td style="font-size: 8px;">{{ r.why|trunc(120) }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  <!-- Forces et faiblesses Go/No-Go -->
  {% if gonogo and (gonogo.strengths or gonogo.risks) %}
  <div style="display: table; width: 100%; margin-top: 12px;">
    <div style="display: table-cell; width: 48%; vertical-align: top;">
      <h3 style="color: #16A34A;">Forces</h3>
      {% for s in (gonogo.strengths or [])[:3] %}<div style="margin: 4px 0; font-size: 10px;">+ {{ s }}</div>{% endfor %}
    </div>
    <div style="display: table-cell; width: 4%;"></div>
    <div style="display: table-cell; width: 48%; vertical-align: top;">
      <h3 style="color: #DC2626;">Points de vigilance</h3>
      {% for r in (gonogo.risks or [])[:3] %}<div style="margin: 4px 0; font-size: 10px;">- {{ r }}</div>{% endfor %}
    </div>
  </div>
  {% endif %}

  <!-- Radar chart Go/No-Go -->
  {% if radar_chart_b64 %}
  <div style="text-align:center; margin: 14px 0;">
    <img src="data:image/png;base64,{{ radar_chart_b64 }}"
         style="width:100%; max-width:460px; display:inline-block;" alt="Radar Go/No-Go"/>
  </div>
  {% endif %}

  <!-- Dimensions Go/No-Go (9 axes) -->
  {% if gonogo and (gonogo.breakdown or gonogo.dimension_scores) %}
  <h3>Scores par dimension ({{ 9 if gonogo.dimension_scores else 4 }} axes)</h3>
  <table>
    <tr><th style="width:40%">Dimension</th><th style="width:15%">Score</th><th style="width:15%">Niveau</th><th style="width:30%">Detail</th></tr>
    {% set dim_labels = {'financial_capacity': 'Capacite financiere', 'market_size': 'Taille du marche', 'certifications': 'Certifications', 'geographic_zone': 'Zone geographique', 'insurance_adequacy': 'Assurances', 'margin_viability': 'Viabilite marge', 'capacity_charge': 'Charge de travail', 'subcontracting_coverage': 'Sous-traitance', 'historical_success': 'Taux de succes', 'technical_fit': 'Adequation technique', 'financial_capacity': 'Capacite financiere', 'timeline_feasibility': 'Faisabilite planning', 'competitive_position': 'Position concurrentielle'} %}
    {% if gonogo.dimension_scores %}
      {% for dim_key, dim_val in [('financial_capacity', gonogo.dimension_scores.financial_capacity), ('market_size', gonogo.dimension_scores.market_size), ('certifications', gonogo.dimension_scores.certifications), ('geographic_zone', gonogo.dimension_scores.geographic_zone), ('insurance_adequacy', gonogo.dimension_scores.insurance_adequacy), ('margin_viability', gonogo.dimension_scores.margin_viability), ('capacity_charge', gonogo.dimension_scores.capacity_charge), ('subcontracting_coverage', gonogo.dimension_scores.subcontracting_coverage), ('historical_success', gonogo.dimension_scores.historical_success)] %}
      {% if dim_val is not none %}
      <tr>
        <td>{{ dim_labels[dim_key] or dim_key }}</td>
        <td><strong>{{ dim_val }}/100</strong></td>
        <td><span class="badge {% if dim_val >= 70 %}badge-green{% elif dim_val >= 50 %}badge-yellow{% else %}badge-red{% endif %}">{% if dim_val >= 70 %}BON{% elif dim_val >= 50 %}MOYEN{% else %}FAIBLE{% endif %}</span></td>
        <td style="font-size: 8px;"></td>
      </tr>
      {% endif %}
      {% endfor %}
    {% elif gonogo.breakdown %}
      {% for dim_name, dim_score in [('Adequation technique', gonogo.breakdown.technical_fit), ('Capacite financiere', gonogo.breakdown.financial_capacity), ('Faisabilite planning', gonogo.breakdown.timeline_feasibility), ('Position concurrentielle', gonogo.breakdown.competitive_position)] %}
      {% if dim_score is not none %}
      <tr>
        <td>{{ dim_name }}</td>
        <td><strong>{{ dim_score }}/100</strong></td>
        <td><span class="badge {% if dim_score >= 70 %}badge-green{% elif dim_score >= 50 %}badge-yellow{% else %}badge-red{% endif %}">{% if dim_score >= 70 %}BON{% elif dim_score >= 50 %}MOYEN{% else %}FAIBLE{% endif %}</span></td>
        <td style="font-size: 8px;"></td>
      </tr>
      {% endif %}
      {% endfor %}
    {% endif %}
  </table>

  <!-- Gaps et forces du profil -->
  {% if gonogo.profile_gaps %}
  <h3 style="color: #DC2626;">Lacunes du profil entreprise</h3>
  <div class="warning-box">
    {% for gap in (gonogo.profile_gaps or [])[:6] %}
    <div style="margin: 3px 0; font-size: 10px;">• {{ gap }}</div>
    {% endfor %}
  </div>
  {% endif %}
  {% if gonogo.profile_strengths %}
  <h3 style="color: #16A34A;">Points forts du profil</h3>
  <div class="financial-box">
    {% for s in (gonogo.profile_strengths or [])[:6] %}
    <div style="margin: 3px 0; font-size: 10px;">• {{ s }}</div>
    {% endfor %}
  </div>
  {% endif %}
  {% endif %}

  <!-- Actions P0 -->
  {% if summary and summary.actions_next_48h %}
  <h3>Actions prioritaires P0</h3>
  <table>
    <tr><th style="width:70%">Action</th><th style="width:30%">Responsable</th></tr>
    {% for a in (summary.actions_next_48h or []) if a.priority == 'P0' %}
    <tr>
      <td style="font-size: 8px;">{{ a.action|trunc(100) }}</td>
      <td style="font-size: 8px;">{{ a.owner_role|trunc(30) }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  <!-- Confiance IA -->
  {% if confidence %}
  <div style="margin-top: 12px;">
    <span style="font-size: 10px; color: #64748b;">Indice de confiance IA : <strong>{{ "%.0f"|format(confidence * 100) }}%</strong></span>
    <div class="confidence-bar">
      <div class="confidence-fill" style="width: {{ "%.0f"|format(confidence * 100) }}%; background: {% if confidence >= 0.8 %}#059669{% elif confidence >= 0.6 %}#d97706{% else %}#dc2626{% endif %};"></div>
    </div>
  </div>
  {% endif %}
</div>

<!-- ═══════════ FICHE SIGNALÉTIQUE DU MARCHÉ ═══════════ -->
{% if rc_analysis %}
<div class="page page-break">
  <h1><span class="section-num">2</span> Fiche signaletique du marche</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Donnees extraites du Reglement de Consultation (RC) et des pieces contractuelles</p>

  <table>
    <tr><th style="width:35%">Élément</th><th style="width:65%">Valeur</th></tr>
    {% if rc_analysis.procedure_type %}
    <tr><td><strong>Procédure</strong></td><td>{{ rc_analysis.procedure_type }}</td></tr>
    {% endif %}
    {% if rc_analysis.allotissement %}
    <tr><td><strong>Allotissement</strong></td><td>{{ rc_analysis.allotissement }}</td></tr>
    {% endif %}
    {% if rc_analysis.groupement %}
    <tr><td><strong>Groupement</strong></td><td>{{ rc_analysis.groupement }}</td></tr>
    {% endif %}
    {% if rc_analysis.variantes %}
    <tr><td><strong>Variantes</strong></td><td>{{ rc_analysis.variantes }}</td></tr>
    {% endif %}
    {% if rc_analysis.subcontracting_allowed is not none %}
    <tr><td><strong>Sous-traitance</strong></td><td>{% if rc_analysis.subcontracting_allowed %}Autorisée{% else %}Non autorisée{% endif %}</td></tr>
    {% endif %}
    {% if rc_analysis.ccag_reference %}
    <tr><td><strong>CCAG de référence</strong></td><td>{{ rc_analysis.ccag_reference }}</td></tr>
    {% endif %}
    {% if rc_analysis.visite_obligatoire is not none %}
    <tr><td><strong>Visite de site</strong></td><td>{% if rc_analysis.visite_obligatoire %}<span class="badge badge-red">OBLIGATOIRE</span>{% else %}Facultative{% endif %}</td></tr>
    {% endif %}
    {% if rc_analysis.dume_required is not none %}
    <tr><td><strong>DUME</strong></td><td>{% if rc_analysis.dume_required %}Requis{% else %}Non requis{% endif %}</td></tr>
    {% endif %}
  </table>

  {% if rc_analysis.lots %}
  <h3>Decomposition en lots</h3>
  <table>
    <tr><th style="width:10%">N°</th><th style="width:50%">Intitule</th><th style="width:40%">Montant estime</th></tr>
    {% for lot in (rc_analysis.lots or []) %}
    <tr>
      <td>{{ lot.number }}</td>
      <td>{{ lot.title }}</td>
      <td>{{ lot.estimated_amount or 'Non précisé' }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  <!-- Conditions d'éligibilité -->
  {% if rc_analysis.who_can_apply %}
  <h3>Conditions d'eligibilite</h3>
  <table>
    <tr><th style="width:65%">Condition</th><th style="width:15%">Type</th><th style="width:20%">Detail</th></tr>
    {% for cond in (rc_analysis.who_can_apply or []) %}
    <tr>
      <td style="font-size: 8px;">{{ (cond.condition or '')|trunc(120) }}</td>
      <td><span class="badge {% if cond.type == 'hard' %}badge-red{% else %}badge-yellow{% endif %}">{{ (cond.type or 'soft')|upper }}</span></td>
      <td style="font-size: 8px;">{{ (cond.details or '')|trunc(50) }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  <!-- Informations complémentaires -->
  {% if rc_analysis.variantes_autorisees is not none or rc_analysis.duree_validite_offres_jours %}
  <h3>Informations complementaires</h3>
  <table>
    <tr><th style="width:40%">Element</th><th style="width:60%">Valeur</th></tr>
    {% if rc_analysis.variantes_autorisees is not none %}
    <tr><td><strong>Variantes autorisees</strong></td><td>{% if rc_analysis.variantes_autorisees %}Oui{% else %}Non{% endif %}</td></tr>
    {% endif %}
    {% if rc_analysis.duree_validite_offres_jours %}
    <tr><td><strong>Duree validite offres</strong></td><td>{{ rc_analysis.duree_validite_offres_jours }} jours</td></tr>
    {% endif %}
    {% if rc_analysis.langue_offre %}
    <tr><td><strong>Langue offre</strong></td><td>{{ rc_analysis.langue_offre }}</td></tr>
    {% endif %}
    {% if rc_analysis.devise_offre %}
    <tr><td><strong>Devise</strong></td><td>{{ rc_analysis.devise_offre }}</td></tr>
    {% endif %}
  </table>
  {% endif %}
</div>
{% endif %}
{% if summary %}
<!-- ═══════════ RÉSUMÉ EXÉCUTIF ═══════════ -->
<div class="page page-break">
  <h1><span class="section-num">3</span> Resume executif</h1>

  <div class="summary-box">
    <strong>Objet :</strong> {{ summary.project_overview.scope }}<br>
    <strong>Acheteur :</strong> {{ summary.project_overview.buyer }}<br>
    <strong>Lieu :</strong> {{ summary.project_overview.location }}<br>
    <strong>Date limite :</strong> {{ summary.project_overview.deadline_submission|datefr }}<br>
    {% if summary.project_overview.estimated_budget %}<strong>Budget estimé :</strong> {{ summary.project_overview.estimated_budget }}<br>{% endif %}
    {% if summary.project_overview.market_type %}<strong>Type de marché :</strong> {{ summary.project_overview.market_type }}<br>{% endif %}
    {% if summary.project_overview.procedure %}<strong>Procédure :</strong> {{ summary.project_overview.procedure }}<br>{% endif %}
    {% if summary.project_overview.allotissement %}<strong>Allotissement :</strong> {{ summary.project_overview.allotissement }}<br>{% endif %}
    {% if summary.project_overview.duree_marche %}<strong>Durée du marché :</strong> {{ summary.project_overview.duree_marche }}<br>{% endif %}
    {% if summary.project_overview.ccag_reference %}<strong>CCAG de référence :</strong> {{ summary.project_overview.ccag_reference }}{% endif %}
  </div>

  <h3>Points clés extraits du DCE ({{ (summary.key_points or [])|length }} points)</h3>
  <table>
    <tr><th style="width:12%">Categorie</th><th style="width:8%">Priorite</th><th style="width:80%">Point cle</th></tr>
    {% for kp in (summary.key_points or []) %}
    <tr>
      <td style="font-size: 7px;"><span class="badge badge-blue" style="padding:1px 4px; font-size:7px;">{{ (kp.label or '')|trunc(12,'') }}</span></td>
      <td><span class="badge {% if kp.importance == 'high' %}badge-red{% elif kp.importance == 'medium' %}badge-yellow{% else %}badge-gray{% endif %}" style="padding:1px 4px; font-size:7px;">{% if kp.importance == 'high' %}HAUT{% elif kp.importance == 'medium' %}MOY.{% else %}BAS{% endif %}</span></td>
      <td style="font-size: 9px;">{{ (kp.point or kp.value or '')|trunc(250) }}
        {% for c in (kp.citations or [])[:1] %}<div class="citation">{% if c.doc %}{{ c.doc|trunc(20,'') }}{% endif %}{% if c.page %} p.{{ c.page }}{% endif %}</div>{% endfor %}
      </td>
    </tr>
    {% endfor %}
  </table>
</div>

{% endif %}

<!-- ═══════════════════════════════════════════════════════ -->
<!-- PARTIE II : CONFORMITE JURIDIQUE -->
<!-- ═══════════════════════════════════════════════════════ -->
<div class="part-header">
  PARTIE II — Conformite juridique
  <div class="part-header-sub">Verification administrative, checklist, derogations CCAG, clauses risquees</div>
</div>

<!-- ═══════════ VÉRIFICATION ADMINISTRATIVE DC ═══════════ -->
{% if dc_check %}
<div class="page page-break">
  <h1><span class="section-num">4</span> Verification administrative (DC1/DC2)</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Conformite des pieces administratives requises — attestations, certifications, formulaires</p>

  {% if dc_check.resume %}
  <div class="summary-box">{{ dc_check.resume }}</div>
  {% endif %}

  {% if dc_check.documents_requis %}
  <h3>Documents et attestations requis</h3>
  <table>
    <tr><th style="width:5%">#</th><th style="width:35%">Document</th><th style="width:12%">Obligatoire</th><th style="width:15%">Statut</th><th style="width:33%">Details</th></tr>
    {% for doc in (dc_check.documents_requis or []) %}
    <tr>
      <td>{{ loop.index }}</td>
      <td style="font-size: 8px;"><strong>{{ (doc.label or doc.document or '')|trunc(60) }}</strong></td>
      <td><span class="badge {% if doc.obligatoire %}badge-red{% else %}badge-gray{% endif %}">{% if doc.obligatoire %}OUI{% else %}NON{% endif %}</span></td>
      <td><span class="badge {% if doc.statut == 'FOURNI' or doc.statut == 'OK' %}badge-green{% elif doc.statut == 'MANQUANT' %}badge-red{% elif doc.statut == 'EXPIRE' %}badge-yellow{% else %}badge-gray{% endif %}">{{ doc.statut or 'INCONNU' }}</span></td>
      <td style="font-size: 8px;">{{ (doc.details or doc.description or '')|trunc(80) }}{% if doc.date_validite %} <em style="color:#64748b;">Validité : {{ doc.date_validite|datefr }}</em>{% endif %}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  {% if dc_check.certifications_requises %}
  <h3>Certifications requises</h3>
  <div class="info-box">
    {% for cert in (dc_check.certifications_requises or []) %}
    <div style="margin: 3px 0; font-size: 10px;">• {{ cert }}</div>
    {% endfor %}
  </div>
  {% endif %}

  {% if dc_check.alertes %}
  <h3>Alertes</h3>
  {% for alerte in (dc_check.alertes or [])[:5] %}
  <div class="warning-box">{{ alerte }}</div>
  {% endfor %}
  {% endif %}
</div>
{% endif %}
<!-- ═══════════ CHECKLIST ═══════════ -->
{% if checklist_items %}
<div class="page page-break">
  <h1><span class="section-num">5</span> Checklist de conformite ({{ checklist_items|length }} exigences)</h1>

  <div class="info-box">
    <strong>Note :</strong> Le statut "MANQUANT" indique les documents/justificatifs à préparer pour la soumission.
    Utilisez cette checklist comme liste de contrôle avant le dépôt de votre offre.
  </div>

  <!-- Statistiques checklist -->
  <div class="stat-grid">
    <div class="stat-cell" style="border-left: 3px solid #DC2626;">
      <div class="stat-number" style="color: #DC2626;">{{ checklist_stats.eliminatoire }}</div>
      <div class="stat-label">Éliminatoires</div>
    </div>
    <div class="stat-cell" style="border-left: 3px solid #D97706;">
      <div class="stat-number" style="color: #D97706;">{{ checklist_stats.important }}</div>
      <div class="stat-label">Importants</div>
    </div>
    <div class="stat-cell" style="border-left: 3px solid #64748B;">
      <div class="stat-number" style="color: #64748B;">{{ checklist_stats.info }}</div>
      <div class="stat-label">Informatifs</div>
    </div>
    <div class="stat-cell" style="border-left: 3px solid #16A34A;">
      <div class="stat-number" style="color: #16A34A;">{{ checklist_stats.ok }}</div>
      <div class="stat-label">Conformes</div>
    </div>
  </div>

  <table>
    <tr><th style="width:4%">#</th><th style="width:29%">Exigence</th><th style="width:13%">Categorie</th><th style="width:15%">Criticite</th><th style="width:10%">Statut</th><th style="width:29%">A fournir</th></tr>
    {% for item in checklist_items %}
    <tr>
      <td>{{ loop.index }}</td>
      <td style="font-size: 9px;">{{ (item.requirement or '')|trunc(100) }}
        {% for c in (item.citations or [])[:2] %}<div class="citation">{% if c is string %}{{ c|trunc(30) }}{% elif c.doc or c.page %}{% if c.doc %}{{ c.doc|trunc(20,'') }}{% endif %}{% if c.page %} p.{{ c.page }}{% endif %}{% endif %}</div>{% endfor %}
      </td>
      <td style="font-size: 8px;">{{ (item.category or '-')|trunc(15,'') }}</td>
      <td style="font-size: 7px;"><span class="badge {% if item.criticality and 'liminato' in item.criticality|lower %}badge-red{% elif item.criticality == 'Important' %}badge-yellow{% else %}badge-gray{% endif %}" style="padding:1px 4px; font-size:7px;">{% if item.criticality and 'liminato' in item.criticality|lower %}ELIM.{% elif item.criticality == 'Important' %}IMP.{% else %}INFO{% endif %}</span></td>
      <td style="font-size: 7px;"><span class="badge {% if item.status == 'OK' %}badge-green{% elif item.status == 'MANQUANT' %}badge-red{% else %}badge-yellow{% endif %}" style="padding:1px 4px; font-size:7px;">{{ item.status or '?' }}</span></td>
      <td style="font-size: 8px;">{{ (item.what_to_provide or '-')|trunc(90) }}</td>
    </tr>
    {% endfor %}
  </table>

  <!-- Documents prioritaires à préparer (éliminatoires MANQUANTS) -->
  {% set elim_manquants = [] %}
  {% for item in checklist_items %}
    {% if item.criticality and 'liminatoire' in item.criticality|lower and item.status == 'MANQUANT' %}
      {% if elim_manquants.append(item) %}{% endif %}
    {% endif %}
  {% endfor %}
  {% if elim_manquants %}
  <div class="page-break"></div>
  <h2 style="color: #DC2626;">Documents prioritaires à préparer ({{ elim_manquants|length }} éliminatoires)</h2>
  <div class="warning-box">
    <strong>Attention :</strong> Ces {{ elim_manquants|length }} documents sont <strong>éliminatoires</strong>.
    Leur absence entraînera le rejet automatique de votre candidature. Préparez-les en priorité absolue.
  </div>
  <table>
    <tr><th style="width:5%">#</th><th style="width:45%">Document / Justificatif requis</th><th style="width:50%">Detail a fournir</th></tr>
    {% for item in elim_manquants %}
    <tr class="risk-high">
      <td>{{ loop.index }}</td>
      <td style="font-size: 8px;"><strong>{{ (item.requirement or '')|trunc(80) }}</strong></td>
      <td style="font-size: 8px;">{{ (item.what_to_provide or 'Voir exigences du RC/CCAP')|trunc(100) }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}
</div>
{% endif %}
<!-- ═══════════ DÉROGATIONS CCAG ═══════════ -->
{% if ccag_derogations %}
<div class="page page-break">
  <h1><span class="section-num">6</span> Derogations au CCAG-Travaux 2021</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Clauses du CCAP/CCTP qui derogent au CCAG-Travaux 2021 — a verifier pour le chiffrage</p>

  <div class="warning-box">
    <strong>Attention :</strong> {{ ccag_derogations|length }} dérogation{{ 's' if ccag_derogations|length > 1 else '' }}
    au CCAG-Travaux 2021 détectée{{ 's' if ccag_derogations|length > 1 else '' }}.
    Certaines peuvent avoir un impact significatif sur votre prix et vos risques contractuels.
  </div>

  <table>
    <tr><th style="width:20%">Article CCAG</th><th style="width:35%">Derogation CCAP</th><th style="width:10%">Gravité</th><th style="width:35%">Evaluation</th></tr>
    {% for d in ccag_derogations %}
    <tr>
      <td><strong>{{ (d.article_ccag or d.article or '—')|trunc(40) }}</strong></td>
      <td style="font-size: 8px;">{{ (d.derogation or '—')|trunc(120) }}</td>
      <td><span class="badge {% if (d.severity or d.impact) == 'high' or (d.severity or d.impact) == 'fort' %}badge-red{% elif (d.severity or d.impact) == 'medium' or (d.severity or d.impact) == 'moyen' %}badge-yellow{% else %}badge-gray{% endif %}">{% if (d.severity or d.impact) == 'high' or (d.severity or d.impact) == 'fort' %}FORT{% elif (d.severity or d.impact) == 'medium' or (d.severity or d.impact) == 'moyen' %}MOY.{% else %}BAS{% endif %}</span></td>
      <td style="font-size: 8px;">{{ (d.evaluation or d.risk_comment or '')|trunc(120) }}</td>
    </tr>
    {% endfor %}
  </table>

  <div class="info-box">
    <strong>Conseil :</strong> Intégrez ces dérogations dans votre chiffrage. En particulier, vérifiez les pénalités
    (art. 20), les délais de paiement (art. 11), la retenue de garantie (art. 32) et les conditions de résiliation (art. 46).
  </div>
</div>
{% endif %}
<!-- ═══════════ CLAUSES RISQUÉES CCAP ═══════════ -->
{% if ccap_clauses_risquees %}
<div class="page page-break">
  <h1><span class="section-num">7</span> Clauses risquees du CCAP</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Clauses contractuelles a risque detectees dans le CCAP — impact sur prix et execution</p>

  <div class="warning-box">
    <strong>Attention :</strong> {{ ccap_clauses_risquees|length }} clause{{ 's' if ccap_clauses_risquees|length > 1 else '' }}
    a risque detectee{{ 's' if ccap_clauses_risquees|length > 1 else '' }} dans le CCAP.
    {% set nb_crit = [] %}{% for c in ccap_clauses_risquees %}{% if c.risk_level == 'CRITIQUE' or c.risk_level == 'critique' %}{% if nb_crit.append(1) %}{% endif %}{% endif %}{% endfor %}
    {% if nb_crit %}Dont {{ nb_crit|length }} de niveau CRITIQUE.{% endif %}
  </div>

  <table>
    <tr><th style="width:30%">Clause</th><th style="width:12%">Type</th><th style="width:10%">Risque</th><th style="width:48%">Conseil</th></tr>
    {% for c in ccap_clauses_risquees %}
    <tr class="{% if c.risk_level == 'CRITIQUE' or c.risk_level == 'critique' or c.risk_level == 'HAUT' or c.risk_level == 'haut' %}risk-high{% elif c.risk_level == 'MOYEN' or c.risk_level == 'moyen' %}risk-medium{% endif %}">
      <td style="font-size: 8px;"><strong>{{ (c.clause_text or c.description or '')|trunc(80) }}</strong>
        {% if c.article_reference %}<br><span class="citation">Art. {{ c.article_reference }}</span>{% endif %}
        {% if c.citation %}<br><span class="citation">{{ c.citation|trunc(60) }}</span>{% endif %}
      </td>
      <td style="font-size: 8px;">{{ (c.risk_type or c.type_risque or '')|trunc(20) }}</td>
      <td><span class="badge {% if c.risk_level == 'CRITIQUE' or c.risk_level == 'critique' %}badge-red{% elif c.risk_level == 'HAUT' or c.risk_level == 'haut' %}badge-red{% elif c.risk_level == 'MOYEN' or c.risk_level == 'moyen' %}badge-yellow{% else %}badge-gray{% endif %}">{{ (c.risk_level or 'INFO')|upper }}</span></td>
      <td style="font-size: 8px;">{{ (c.conseil or c.recommendation or '')|trunc(120) }}</td>
    </tr>
    {% endfor %}
  </table>
</div>
{% endif %}

<!-- ═══════════════════════════════════════════════════════ -->
<!-- PARTIE III : ANALYSE FINANCIERE -->
<!-- ═══════════════════════════════════════════════════════ -->
<div class="part-header">
  PARTIE III — Analyse financiere
  <div class="part-header-sub">Acte d'engagement, synthese financiere, benchmark DPGF, tresorerie</div>
</div>

<!-- ═══════════ ANALYSE ACTE D'ENGAGEMENT ═══════════ -->
{% if ae_analysis %}
<div class="page page-break">
  <h1><span class="section-num">8</span> Analyse de l'Acte d'Engagement</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Engagements contractuels, conditions financieres et clauses a risque extraits de l'AE</p>

  {% if ae_analysis.resume %}
  <div class="summary-box">{{ ae_analysis.resume }}</div>
  {% endif %}

  <!-- Score risque AE -->
  {% if ae_analysis.score_risque_global is not none %}
  <div style="margin: 12px 0; padding: 10px; background: {% if ae_analysis.score_risque_global >= 70 %}#fef2f2{% elif ae_analysis.score_risque_global >= 40 %}#fffbeb{% else %}#ecfdf5{% endif %}; border-left: 5px solid {% if ae_analysis.score_risque_global >= 70 %}#ef4444{% elif ae_analysis.score_risque_global >= 40 %}#f59e0b{% else %}#10b981{% endif %};">
    <strong>Score de risque contractuel :</strong> {{ ae_analysis.score_risque_global }}/100
    — <span style="font-weight: 700; color: {% if ae_analysis.score_risque_global >= 70 %}#b91c1c{% elif ae_analysis.score_risque_global >= 40 %}#b45309{% else %}#047857{% endif %};">{% if ae_analysis.score_risque_global >= 70 %}ÉLEVÉ{% elif ae_analysis.score_risque_global >= 40 %}MODÉRÉ{% else %}FAIBLE{% endif %}</span>
    <div class="confidence-bar" style="margin-top: 6px;">
      <div class="confidence-fill" style="width: {{ ae_analysis.score_risque_global }}%; background: {% if ae_analysis.score_risque_global >= 70 %}#ef4444{% elif ae_analysis.score_risque_global >= 40 %}#f59e0b{% else %}#10b981{% endif %};"></div>
    </div>
  </div>
  {% endif %}

  <!-- Données structurées AE -->
  <h3>Conditions financieres et contractuelles</h3>
  <table>
    <tr><th style="width:35%">Element</th><th style="width:65%">Valeur</th></tr>
    {% if ae_analysis.prix_forme %}
    <tr><td><strong>Forme du prix</strong></td><td>{{ ae_analysis.prix_forme }}</td></tr>
    {% endif %}
    {% if ae_analysis.montant_total_ht %}
    <tr><td><strong>Montant total HT</strong></td><td>{{ ae_analysis.montant_total_ht }}</td></tr>
    {% endif %}
    {% if ae_analysis.duree_marche %}
    <tr><td><strong>Duree du marche</strong></td><td>{{ ae_analysis.duree_marche }}</td></tr>
    {% endif %}
    {% if ae_analysis.prix_revision is not none %}
    <tr><td><strong>Revision des prix</strong></td><td>{% if ae_analysis.prix_revision %}Oui{% if ae_analysis.prix_revision_details %} — {{ ae_analysis.prix_revision_details|trunc(120) }}{% endif %}{% else %}Non{% endif %}</td></tr>
    {% endif %}
    {% if ae_analysis.reconduction is not none %}
    <tr><td><strong>Reconduction</strong></td><td>{% if ae_analysis.reconduction %}Oui{% if ae_analysis.reconduction_details %} — {{ ae_analysis.reconduction_details|trunc(100) }}{% endif %}{% else %}Non{% endif %}</td></tr>
    {% endif %}
    {% if ae_analysis.penalites_retard %}
    <tr class="risk-medium"><td><strong>Penalites de retard</strong></td><td>{{ ae_analysis.penalites_retard|trunc(150) }}</td></tr>
    {% endif %}
    {% if ae_analysis.retenue_garantie_pct is not none %}
    <tr><td><strong>Retenue de garantie</strong></td><td>{{ ae_analysis.retenue_garantie_pct }}%</td></tr>
    {% endif %}
    {% if ae_analysis.avance_pct is not none %}
    <tr><td><strong>Avance forfaitaire</strong></td><td>{{ ae_analysis.avance_pct }}%</td></tr>
    {% endif %}
    {% if ae_analysis.delai_paiement_jours is not none %}
    <tr {% if ae_analysis.delai_paiement_jours > 30 %}class="risk-high"{% endif %}><td><strong>Delai de paiement</strong></td><td>{{ ae_analysis.delai_paiement_jours }} jours{% if ae_analysis.delai_paiement_jours > 30 %} <span class="badge badge-red">ILLEGAL > 30j</span>{% endif %}</td></tr>
    {% endif %}
  </table>

  <!-- Clauses risquées AE -->
  {% if ae_analysis.clauses_risquees %}
  <h3>Clauses a risque de l'Acte d'Engagement ({{ ae_analysis.clauses_risquees|length }})</h3>
  <table>
    <tr><th style="width:12%">Type</th><th style="width:35%">Clause</th><th style="width:10%">Risque</th><th style="width:43%">Conseil</th></tr>
    {% for c in (ae_analysis.clauses_risquees or []) %}
    <tr class="{% if c.risk_level == 'CRITIQUE' or c.risk_level == 'HAUT' or c.risk_level == 'high' %}risk-high{% elif c.risk_level == 'MOYEN' or c.risk_level == 'medium' %}risk-medium{% endif %}">
      <td style="font-size: 8px;">{{ (c.clause_type or c.type or '')|trunc(20) }}</td>
      <td style="font-size: 8px;">{{ (c.description or '')|trunc(100) }}{% if c.citation %}<br><span class="citation">{{ c.citation|trunc(60) }}</span>{% endif %}</td>
      <td><span class="badge {% if c.risk_level in ['CRITIQUE', 'HAUT', 'high'] %}badge-red{% elif c.risk_level in ['MOYEN', 'medium'] %}badge-yellow{% else %}badge-gray{% endif %}">{{ (c.risk_level or 'INFO')|upper }}</span></td>
      <td style="font-size: 8px;">{{ (c.conseil or c.recommendation or '')|trunc(120) }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}
</div>
{% endif %}
<!-- ═══════════ SYNTHÈSE FINANCIÈRE ═══════════ -->
{% if summary %}
<div class="page page-break">
  <h1><span class="section-num">9</span> Synthese financiere</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Elements financiers cles extraits du DCE pour l'aide a la decision</p>

  <div class="financial-box">
    <strong>Budget global estimé :</strong> {{ summary.project_overview.estimated_budget or 'Non précisé dans le DCE' }}<br>
    <strong>Type de prix :</strong> {{ summary.project_overview.market_type or 'Non précisé' }}
  </div>

  <!-- Extraction des données financières depuis key_points -->
  <h3>Éléments financiers extraits</h3>
  <table>
    <tr><th style="width:100%">Point financier</th></tr>
    {% set has_financial = [] %}
    {% for kp in (summary.key_points or []) %}
    {% set kp_point_low = (kp.point or '')|lower %}
    {% if 'prix' in kp_point_low or 'avance' in kp_point_low or 'retenue' in kp_point_low or 'paiement' in kp_point_low or 'pénalité' in kp_point_low or 'révision' in kp_point_low or 'financ' in kp_point_low or 'budget' in kp_point_low or 'garantie' in kp_point_low or 'caution' in kp_point_low %}
    {% if has_financial.append(1) %}{% endif %}
    <tr>
      <td style="font-size: 9px;">{{ kp.point|trunc(200) }}</td>
    </tr>
    {% endif %}
    {% endfor %}
    {% if not has_financial %}
    <tr><td style="font-size: 9px; color: #64748b;">Aucun élément financier spécifique extrait des points clés.</td></tr>
    {% endif %}
  </table>

  {% if summary.risks %}
  <h3>Risques financiers identifiés</h3>
  <table>
    <tr><th style="width:30%">Risque</th><th style="width:10%">Severite</th><th style="width:60%">Impact financier</th></tr>
    {% for r in (summary.risks or []) %}
    {% set r_why_low = (r.why or '')|lower %}{% set r_risk_low = (r.risk or '')|lower %}
    {% if 'financ' in r_why_low or 'prix' in r_risk_low or 'pénalité' in r_risk_low or 'coût' in r_why_low or 'trésorerie' in r_why_low or 'paiement' in r_risk_low %}
    <tr class="risk-{{ r.severity }}">
      <td><strong>{{ r.risk|trunc(50) }}</strong></td>
      <td><span class="badge {% if r.severity == 'high' %}badge-red{% elif r.severity == 'medium' %}badge-yellow{% else %}badge-gray {% endif %}">{{ (r.severity or '')|upper }}</span></td>
      <td style="font-size: 8px;">{{ r.why|trunc(150) }}</td>
    </tr>
    {% endif %}
    {% endfor %}
  </table>
  {% endif %}

  <!-- Enrichissement: exposition pénalités AE + impact avance/retenue cashflow -->
  {% if ae_analysis %}
  <h3>Exposition financiere contractuelle (donnees AE)</h3>
  <table>
    <tr><th style="width:40%">Element</th><th style="width:60%">Valeur / Impact</th></tr>
    {% if ae_analysis.penalites_retard %}
    <tr class="risk-medium"><td><strong>Penalites de retard</strong></td><td>{{ ae_analysis.penalites_retard|trunc(150) }}</td></tr>
    {% endif %}
    {% if ae_analysis.retenue_garantie_pct is not none %}
    <tr><td><strong>Retenue de garantie</strong></td><td>{{ ae_analysis.retenue_garantie_pct }}%{% if ae_analysis.montant_total_ht %} — soit ~{{ "{:,.0f}".format((ae_analysis.retenue_garantie_pct or 0) / 100.0 * (ae_analysis.montant_total_ht|float or 0)).replace(",", " ") }} € HT{% endif %}</td></tr>
    {% endif %}
    {% if ae_analysis.avance_pct is not none %}
    <tr><td><strong>Avance forfaitaire</strong></td><td>{{ ae_analysis.avance_pct }}%{% if ae_analysis.montant_total_ht %} — soit ~{{ "{:,.0f}".format((ae_analysis.avance_pct or 0) / 100.0 * (ae_analysis.montant_total_ht|float or 0)).replace(",", " ") }} € HT{% endif %}</td></tr>
    {% endif %}
    {% if ae_analysis.delai_paiement_jours is not none %}
    <tr {% if ae_analysis.delai_paiement_jours > 30 %}class="risk-high"{% endif %}><td><strong>Delai de paiement</strong></td><td>{{ ae_analysis.delai_paiement_jours }} jours{% if ae_analysis.delai_paiement_jours > 30 %} <span class="badge badge-red">EXCESSIF</span>{% endif %}</td></tr>
    {% endif %}
  </table>
  {% endif %}

  {% if cashflow_data and cashflow_data.bfr_peak %}
  <h3>Impact tresorerie previsionnel</h3>
  <div class="financial-box">
    <strong>BFR maximal :</strong> {{ "{:,.0f}".format((cashflow_data.bfr_peak or cashflow_data.bfr_eur or 0)|abs).replace(",", " ") }} €
    {% if cashflow_data.avance_impact_eur %}<br><strong>Impact avance :</strong> +{{ "{:,.0f}".format(cashflow_data.avance_impact_eur).replace(",", " ") }} € (tresorerie){% endif %}
    {% if cashflow_data.retenue_impact_eur %}<br><strong>Impact retenue :</strong> -{{ "{:,.0f}".format(cashflow_data.retenue_impact_eur).replace(",", " ") }} € (immobilise){% endif %}
  </div>
  {% endif %}

  <div class="warning-box">
    <strong>Recommandation :</strong> Avant de chiffrer votre offre, vérifiez les éléments suivants :
    formule de révision des prix, montant de l'avance, conditions de paiement, pénalités de retard,
    retenue de garantie, et tout risque de surcoût identifié (pollution, aléas géotechniques, etc.).
  </div>
</div>
{% endif %}

<!-- ═══════════ BENCHMARK TARIFAIRE DPGF ═══════════ -->
{% if dpgf_pricing %}
<div class="page page-break">
  <h1><span class="section-num">10</span> Benchmark tarifaire DPGF</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Comparaison des prix unitaires DPGF avec le referentiel marche BTP France 2026</p>

  {% if pricing_chart_b64 %}
  <div style="text-align:center; margin: 10px 0 14px 0;">
    <img src="data:image/png;base64,{{ pricing_chart_b64 }}"
         style="width:100%; max-width:520px; display:inline-block;" alt="Benchmark prix DPGF"/>
  </div>
  {% endif %}

  {% set nb_sous = [] %}{% set nb_sur = [] %}{% set nb_normal = [] %}{% set nb_inconnu = [] %}
  {% for line in dpgf_pricing %}
    {% if line.status == 'SOUS_EVALUE' %}{% if nb_sous.append(1) %}{% endif %}
    {% elif line.status == 'SUR_EVALUE' %}{% if nb_sur.append(1) %}{% endif %}
    {% elif line.status == 'NORMAL' %}{% if nb_normal.append(1) %}{% endif %}
    {% else %}{% if nb_inconnu.append(1) %}{% endif %}
    {% endif %}
  {% endfor %}

  <div class="stat-grid">
    <div class="stat-cell" style="border-left: 3px solid #0f1b4c;">
      <div class="stat-number">{{ dpgf_pricing|length }}</div>
      <div class="stat-label">Postes analyses</div>
    </div>
    <div class="stat-cell" style="border-left: 3px solid #DC2626;">
      <div class="stat-number" style="color: #DC2626;">{{ nb_sous|length }}</div>
      <div class="stat-label">Sous-evalues</div>
    </div>
    <div class="stat-cell" style="border-left: 3px solid #D97706;">
      <div class="stat-number" style="color: #D97706;">{{ nb_sur|length }}</div>
      <div class="stat-label">Sur-evalues</div>
    </div>
    <div class="stat-cell" style="border-left: 3px solid #16A34A;">
      <div class="stat-number" style="color: #16A34A;">{{ nb_normal|length }}</div>
      <div class="stat-label">Normaux</div>
    </div>
  </div>

  <table>
    <tr><th style="width:25%">Designation</th><th style="width:10%">Prix unit.</th><th style="width:20%">Reference marche</th><th style="width:10%">Ratio</th><th style="width:12%">Statut</th><th style="width:23%">Alerte</th></tr>
    {% for line in dpgf_pricing[:30] %}
    <tr class="{% if line.status == 'SOUS_EVALUE' %}risk-high{% elif line.status == 'SUR_EVALUE' %}risk-medium{% endif %}">
      <td style="font-size: 8px;"><strong>{{ (line.designation or '')|trunc(40) }}</strong></td>
      <td style="font-size: 8px;">{{ "%.2f"|format(line.prix_unitaire or 0) }} €</td>
      <td style="font-size: 7px;">{% if line.reference_match %}{{ line.reference_match|trunc(25) }}<br>{{ "%.0f"|format(line.reference_prix_min or 0) }}–{{ "%.0f"|format(line.reference_prix_max or 0) }} € (moy. {{ "%.0f"|format(line.reference_prix_moyen or 0) }}){% else %}—{% endif %}</td>
      <td style="font-size: 8px;">{% if line.ratio_vs_moyen %}{{ "%.2f"|format(line.ratio_vs_moyen) }}x{% else %}—{% endif %}</td>
      <td><span class="badge {% if line.status == 'SOUS_EVALUE' %}badge-red{% elif line.status == 'SUR_EVALUE' %}badge-yellow{% elif line.status == 'NORMAL' %}badge-green{% else %}badge-gray{% endif %}">{{ line.status or 'INCONNU' }}</span></td>
      <td style="font-size: 7px;">{{ (line.alerte or '')|trunc(60) }}</td>
    </tr>
    {% endfor %}
  </table>

  <div class="info-box">
    <strong>Avertissement :</strong> Les prix de reference sont des fourchettes indicatives 2026 (France metropolitaine).
    Les ecarts regionaux, la conjoncture materiaux et la complexite du chantier peuvent faire varier les prix de ±30%.
    Coefficient geographique applique : {{ dpgf_pricing[0].geo_coefficient if dpgf_pricing else '1.00' }}.
  </div>
</div>
{% endif %}

<!-- ═══════════ SIMULATION TRÉSORERIE / BFR ═══════════ -->
{% if cashflow_data %}
<div class="page page-break">
  <h1><span class="section-num">11</span> Simulation tresorerie et BFR</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Impact previsionnel sur la tresorerie de l'entreprise — aide a la decision financiere</p>

  {% set bfr_val = cashflow_data.bfr_peak or cashflow_data.bfr_eur %}
  {% if bfr_val %}
  <div class="financial-box">
    <strong>BFR maximal estimé :</strong> {{ "{:,.0f}".format(bfr_val|abs).replace(",", " ") }} €
    {% if cashflow_data.bfr_month %} — atteint au mois {{ cashflow_data.bfr_month }}{% endif %}<br>
    {% set margin_val = cashflow_data.margin_estimate or cashflow_data.marge_brute_pct %}
    {% if margin_val %}
    <strong>Marge estimée :</strong> {{ "%.1f"|format(margin_val) }}%
    {% endif %}
    {% if cashflow_data.montant_total_ht %}
    <br><strong>Montant total HT :</strong> {{ "{:,.0f}".format(cashflow_data.montant_total_ht).replace(",", " ") }} €
    {% endif %}
  </div>
  {% endif %}

  {% if cashflow_data.resume %}
  <div class="summary-box">{{ cashflow_data.resume }}</div>
  {% endif %}

  {% if cashflow_data.monthly_cashflow %}
  <h3>Trésorerie mensuelle prévisionnelle</h3>
  {% if cashflow_chart_b64 %}
  <div style="text-align:center; margin: 10px 0 14px 0;">
    <img src="data:image/png;base64,{{ cashflow_chart_b64 }}"
         style="width:100%; max-width:520px; display:inline-block;" alt="Courbe trésorerie"/>
  </div>
  {% endif %}
  <table>
    <tr><th style="width:15%">Mois</th><th style="width:22%">Depenses HT</th><th style="width:22%">Encaissements HT</th><th style="width:23%">Solde cumule</th></tr>
    {% for m in (cashflow_data.monthly_cashflow or [])[:15] %}
    <tr>
      <td>{{ m.label or ('M' ~ (m.month or m.mois or '?')) }}</td>
      <td style="color: #DC2626; font-size: 8px;">-{{ "{:,.0f}".format(m.expenses or m.depenses_ht or 0).replace(",", " ") }} EUR</td>
      <td style="color: #16A34A; font-size: 8px;">+{{ "{:,.0f}".format(m.income or m.encaissement_ht or 0).replace(",", " ") }} EUR</td>
      {% set cum_val = m.cumulative or m.solde_cumule or 0 %}
      <td style="font-size: 8px;"><strong style="color: {% if cum_val < 0 %}#DC2626{% else %}#16A34A{% endif %};">{{ "{:,.0f}".format(cum_val).replace(",", " ") }} EUR</strong></td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  {% if cashflow_data.risk_level %}
  <div class="warning-box">
    <strong>Niveau de risque trésorerie :</strong> {{ cashflow_data.risk_level }}
    {% if cashflow_data.tension_months %} — {{ cashflow_data.tension_months|length }} mois en tension{% endif %}
  </div>
  {% endif %}

  {% if cashflow_data.warnings %}
  <h3>Alertes trésorerie</h3>
  {% for w in (cashflow_data.warnings or [])[:5] %}
  <div class="warning-box">{{ w }}</div>
  {% endfor %}
  {% endif %}

  <div class="info-box">
    <strong>Hypothèses :</strong> Délai de paiement acheteur 30j
    {% if cashflow_data.avance_impact_eur %}, avance {{ "{:,.0f}".format(cashflow_data.avance_impact_eur).replace(",", " ") }} €{% endif %}
    {% if cashflow_data.retenue_impact_eur %}, retenue garantie {{ "{:,.0f}".format(cashflow_data.retenue_impact_eur).replace(",", " ") }} €{% endif %}
    {% if cashflow_data.duree_mois %}, durée {{ cashflow_data.duree_mois }} mois{% endif %}.
    Ajustez selon les termes du CCAP.
  </div>
</div>
{% endif %}

<!-- ═══════════════════════════════════════════════════════ -->
<!-- PARTIE IV : ANALYSE TECHNIQUE -->
<!-- ═══════════════════════════════════════════════════════ -->
<div class="part-header">
  PARTIE IV — Analyse technique
  <div class="part-header-sub">CCTP, contradictions inter-documents, sous-traitance</div>
</div>

<!-- ═══════════ ANALYSE CCTP ═══════════ -->
{% if cctp_analysis %}
<div class="page page-break">
  <h1><span class="section-num">12</span> Analyse technique CCTP</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Synthese des exigences techniques extraites du Cahier des Clauses Techniques Particulieres</p>

  {% if cctp_analysis.technical_summary %}
  <div class="summary-box">
    <strong>Synthèse technique :</strong> {{ cctp_analysis.technical_summary }}
  </div>
  {% endif %}

  {% if cctp_analysis.categories %}
  {% for cat in (cctp_analysis.categories or []) %}
  <h3>{{ cat.name or 'Catégorie' }} {% if cat.risk_level %}<span class="badge {% if cat.risk_level == 'high' %}badge-red{% elif cat.risk_level == 'medium' %}badge-yellow{% else %}badge-green{% endif %}">{{ cat.risk_level|upper }}</span>{% endif %}</h3>
  {% if cat.items %}
  <table>
    <tr><th style="width:30%">Exigence</th><th style="width:38%">Detail</th><th style="width:20%">Norme/DTU</th><th style="width:12%">Risque</th></tr>
    {% for item in (cat.items or [])[:8] %}
    <tr>
      <td style="font-size: 8px;"><strong>{{ (item.requirement or item.label or '')[:60] }}{% if (item.requirement or item.label or '')|length > 60 %}...{% endif %}</strong></td>
      <td style="font-size: 8px;">{{ (item.detail or item.value or '')[:100] }}{% if (item.detail or item.value or '')|length > 100 %}...{% endif %}</td>
      <td style="font-size: 8px;">{{ (item.norm or item.standard or '')[:30] }}</td>
      <td><span class="badge {% if item.risk == 'high' %}badge-red{% elif item.risk == 'medium' %}badge-yellow{% else %}badge-gray{% endif %}">{{ item.risk|upper if item.risk else 'INFO' }}</span></td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}
  {% endfor %}
  {% endif %}

  {% if cctp_analysis.contradictions %}
  <h3 style="color: #DC2626;">Contradictions détectées dans le CCTP</h3>
  <div class="warning-box">
    <strong>Attention :</strong> {{ cctp_analysis.contradictions|length }} contradiction{{ 's' if cctp_analysis.contradictions|length > 1 else '' }}
    interne{{ 's' if cctp_analysis.contradictions|length > 1 else '' }} détectée{{ 's' if cctp_analysis.contradictions|length > 1 else '' }} dans le CCTP.
    Posez la question à l'acheteur avant la date limite.
  </div>
  <table>
    <tr><th style="width:45%">Contradiction</th><th style="width:25%">References</th><th style="width:30%">Recommandation</th></tr>
    {% for c in (cctp_analysis.contradictions or [])[:5] %}
    <tr class="risk-high">
      <td style="font-size: 8px;">{{ (c.description or c.issue or '')[:120] }}{% if (c.description or c.issue or '')|length > 120 %}...{% endif %}</td>
      <td style="font-size: 8px;">{{ (c.references or c.source or '')[:60] }}</td>
      <td style="font-size: 8px;">{{ (c.recommendation or 'Demander clarification')[:80] }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  {% if cctp_analysis.essais_controles %}
  <h3>Essais et contrôles requis</h3>
  <table>
    <tr><th style="width:35%">Essai / Contrôle</th><th style="width:15%">Phase</th><th style="width:15%">Obligatoire</th><th style="width:35%">Référence</th></tr>
    {% for essai in (cctp_analysis.essais_controles or [])[:10] %}
    <tr>
      <td style="font-size: 8px;"><strong>{{ (essai.label or essai.name or essai)|trunc(60) }}</strong></td>
      <td style="font-size: 8px;">{{ (essai.phase or '')|trunc(20) }}</td>
      <td><span class="badge {% if essai.obligatoire or essai.mandatory %}badge-red{% else %}badge-gray{% endif %}">{% if essai.obligatoire or essai.mandatory %}OUI{% else %}NON{% endif %}</span></td>
      <td style="font-size: 8px;">{{ (essai.reference or essai.norme or '')|trunc(40) }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  {% if cctp_analysis.documents_execution %}
  <h3>Documents d'exécution requis</h3>
  <div class="info-box">
    {% for doc_exec in (cctp_analysis.documents_execution or [])[:8] %}
    <div style="margin: 4px 0; font-size: 10px;">• {% if doc_exec is string %}{{ doc_exec }}{% else %}{{ doc_exec.label or doc_exec.name or doc_exec }}{% if doc_exec.deadline %} — <em>{{ doc_exec.deadline }}</em>{% endif %}{% endif %}</div>
    {% endfor %}
  </div>
  {% endif %}

  {% if cctp_analysis.score_complexite_technique is not none %}
  <div style="margin: 12px 0; padding: 10px; background: {% if cctp_analysis.score_complexite_technique >= 70 %}#fef2f2{% elif cctp_analysis.score_complexite_technique >= 40 %}#fffbeb{% else %}#ecfdf5{% endif %}; border-left: 5px solid {% if cctp_analysis.score_complexite_technique >= 70 %}#ef4444{% elif cctp_analysis.score_complexite_technique >= 40 %}#f59e0b{% else %}#10b981{% endif %};">
    <strong>Score de complexité technique :</strong> {{ cctp_analysis.score_complexite_technique }}/100
    — <span style="font-weight: 700; color: {% if cctp_analysis.score_complexite_technique >= 70 %}#b91c1c{% elif cctp_analysis.score_complexite_technique >= 40 %}#b45309{% else %}#047857{% endif %};">{% if cctp_analysis.score_complexite_technique >= 70 %}COMPLEXE{% elif cctp_analysis.score_complexite_technique >= 40 %}MODÉRÉ{% else %}SIMPLE{% endif %}</span>
    <div class="confidence-bar" style="margin-top: 6px;">
      <div class="confidence-fill" style="width: {{ cctp_analysis.score_complexite_technique }}%; background: {% if cctp_analysis.score_complexite_technique >= 70 %}#ef4444{% elif cctp_analysis.score_complexite_technique >= 40 %}#f59e0b{% else %}#10b981{% endif %};"></div>
    </div>
  </div>
  {% endif %}

  {% if cctp_analysis.environmental_requirements %}
  <h3>Exigences environnementales</h3>
  <div class="info-box">
    {% for req in (cctp_analysis.environmental_requirements or [])[:5] %}
    <div style="margin: 4px 0; font-size: 10px;">• {{ req }}</div>
    {% endfor %}
  </div>
  {% endif %}
</div>
{% endif %}
<!-- ═══════════ CONTRADICTIONS INTER-DOCUMENTS ═══════════ -->
{% if conflicts %}
<div class="page page-break">
  <h1><span class="section-num">13</span> Contradictions inter-documents</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Incoherences detectees entre les pieces du DCE (RC, CCAP, CCTP, DPGF, AE)</p>

  <div class="warning-box">
    <strong>Attention :</strong> {{ conflicts.nb_total or (conflicts.conflicts or [])|length }} contradiction{{ 's' if (conflicts.nb_total or (conflicts.conflicts or [])|length) > 1 else '' }} detectee{{ 's' if (conflicts.nb_total or (conflicts.conflicts or [])|length) > 1 else '' }}
    {% if conflicts.nb_critiques %} dont <strong>{{ conflicts.nb_critiques }} critique{{ 's' if conflicts.nb_critiques > 1 else '' }}</strong>{% endif %}.
    Posez ces questions a l'acheteur avant la date limite.
  </div>

  {% if conflicts.resume %}
  <div class="summary-box">{{ conflicts.resume }}</div>
  {% endif %}

  {% if heatmap_chart_b64 %}
  <div style="text-align:center; margin: 10px 0 14px 0;">
    <img src="data:image/png;base64,{{ heatmap_chart_b64 }}"
         style="width:100%; max-width:500px; display:inline-block;" alt="Heatmap conflits DCE"/>
  </div>
  {% endif %}

  {% if conflicts.conflicts %}
  <table>
    <tr><th style="width:12%">Type</th><th style="width:10%">Gravite</th><th style="width:10%">Doc A</th><th style="width:10%">Doc B</th><th style="width:33%">Description</th><th style="width:25%">Recommandation</th></tr>
    {% for c in (conflicts.conflicts or []) %}
    <tr class="{% if c.severity == 'CRITIQUE' or c.severity == 'critique' %}risk-high{% elif c.severity == 'HAUT' or c.severity == 'haut' %}risk-high{% elif c.severity == 'MOYEN' or c.severity == 'moyen' %}risk-medium{% endif %}">
      <td style="font-size: 7px;"><span class="badge {% if c.conflict_type == 'clause_illegale' %}badge-red{% elif c.conflict_type == 'montant' %}badge-yellow{% elif c.conflict_type == 'delai' %}badge-yellow{% else %}badge-gray{% endif %}">{{ (c.conflict_type or '')|trunc(15) }}</span></td>
      <td><span class="badge {% if c.severity in ['CRITIQUE', 'critique'] %}badge-red{% elif c.severity in ['HAUT', 'haut'] %}badge-red{% elif c.severity in ['MOYEN', 'moyen'] %}badge-yellow{% else %}badge-gray{% endif %}">{{ (c.severity or 'INFO')|upper }}</span></td>
      <td style="font-size: 7px;">{{ (c.doc_a or '')|trunc(12) }}</td>
      <td style="font-size: 7px;">{{ (c.doc_b or '')|trunc(12) }}</td>
      <td style="font-size: 8px;">{{ (c.description or '')|trunc(100) }}
        {% if c.citation_a %}<br><span class="citation">A: {{ c.citation_a|trunc(50) }}</span>{% endif %}
        {% if c.citation_b %}<br><span class="citation">B: {{ c.citation_b|trunc(50) }}</span>{% endif %}
      </td>
      <td style="font-size: 8px;">{{ (c.recommendation or '')|trunc(80) }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  <div class="info-box">
    <strong>Conseil :</strong> Les contradictions de type "clause_illegale" et "montant" sont les plus critiques.
    Posez systematiquement la question a l'acheteur via la plateforme de dematerialisation.
    Gardez une trace ecrite des reponses pour securiser votre offre.
  </div>
</div>
{% endif %}
<!-- ═══════════ ANALYSE SOUS-TRAITANCE ═══════════ -->
{% if subcontracting %}
<div class="page page-break">
  <h1><span class="section-num">14</span> Analyse sous-traitance</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Obligations et risques lies a la sous-traitance — conformite loi du 31/12/1975</p>

  {% if subcontracting.summary %}
  <div class="summary-box">{{ subcontracting.summary }}</div>
  {% endif %}

  {% if subcontracting.conditions %}
  <h3>Conditions contractuelles</h3>
  <table>
    <tr><th style="width:30%">Condition</th><th style="width:55%">Detail</th><th style="width:15%">Risque</th></tr>
    {% for c in (subcontracting.conditions or []) %}
    <tr>
      <td style="font-size: 8px;"><strong>{{ (c.label or c.condition or '')[:50] }}</strong></td>
      <td style="font-size: 8px;">{{ (c.detail or c.value or '')[:120] }}{% if (c.detail or c.value or '')|length > 120 %}...{% endif %}</td>
      <td><span class="badge {% if c.risk == 'high' %}badge-red{% elif c.risk == 'medium' %}badge-yellow{% else %}badge-green{% endif %}">{{ c.risk|upper if c.risk else 'INFO' }}</span></td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  {% if subcontracting.risks %}
  <h3>Risques identifiés</h3>
  {% for r in (subcontracting.risks or [])[:5] %}
  <div class="warning-box">
    <strong>{{ r.risk or r.title or '' }} :</strong> {{ r.detail or r.description or '' }}
    {% if r.mitigation %}<br><em style="color: #1E40AF;">Mitigation : {{ r.mitigation }}</em>{% endif %}
  </div>
  {% endfor %}
  {% endif %}

  {% if subcontracting.legal_obligations %}
  <h3>Obligations légales</h3>
  <div class="info-box">
    {% for o in (subcontracting.legal_obligations or [])[:6] %}
    <div style="margin: 4px 0; font-size: 10px;">• {{ o }}</div>
    {% endfor %}
  </div>
  {% endif %}

  <!-- Analyse par lot (compétences internes vs sous-traitance) -->
  {% if subcontracting.lots_analysis %}
  <h3>Analyse par lot — Competences et recommandations</h3>
  <table>
    <tr><th style="width:15%">Lot</th><th style="width:25%">Competence requise</th><th style="width:12%">Interne</th><th style="width:15%">Sous-traiter</th><th style="width:12%">Risque</th><th style="width:21%">Justification</th></tr>
    {% for lot in (subcontracting.lots_analysis or []) %}
    <tr>
      <td style="font-size: 8px;">{{ (lot.lot or '')|trunc(20) }}</td>
      <td style="font-size: 8px;">{{ (lot.competence_requise or '')|trunc(40) }}</td>
      <td><span class="badge {% if lot.competence_interne %}badge-green{% else %}badge-red{% endif %}">{% if lot.competence_interne %}OUI{% else %}NON{% endif %}</span></td>
      <td><span class="badge {% if lot.sous_traitance_recommandee %}badge-yellow{% else %}badge-green{% endif %}">{% if lot.sous_traitance_recommandee %}OUI{% else %}NON{% endif %}</span></td>
      <td><span class="badge {% if lot.risque == 'eleve' or lot.risque == 'élevé' %}badge-red{% elif lot.risque == 'modere' or lot.risque == 'modéré' %}badge-yellow{% else %}badge-green{% endif %}">{{ (lot.risque or 'faible')|upper }}</span></td>
      <td style="font-size: 7px;">{{ (lot.justification or '')|trunc(60) }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  <!-- Conflits sous-traitance -->
  {% if subcontracting.conflits %}
  <h3 style="color: #DC2626;">Conflits detectes</h3>
  {% for c in (subcontracting.conflits or [])[:5] %}
  <div class="warning-box">
    <strong>{{ (c.type or 'Conflit')|trunc(30) }} :</strong> {{ (c.description or '')|trunc(150) }}
  </div>
  {% endfor %}
  {% endif %}

  <!-- Paiement direct -->
  {% if subcontracting.paiement_direct_applicable is not none %}
  <div class="info-box">
    <strong>Paiement direct (art. 133 CMP) :</strong>
    {% if subcontracting.paiement_direct_applicable %}Applicable{% else %}Non applicable{% endif %}
    {% if subcontracting.seuil_paiement_direct_eur %} — Seuil : {{ "{:,.0f}".format(subcontracting.seuil_paiement_direct_eur).replace(",", " ") }} € TTC{% endif %}
  </div>
  {% endif %}

  {% if subcontracting.recommandations %}
  <h3>Recommandations strategiques</h3>
  {% for r in (subcontracting.recommandations or [])[:5] %}
  <div style="margin: 4px 0; font-size: 10px; padding: 4px 8px; background: #f8fafc; border-left: 3px solid #2563eb;">{{ r }}</div>
  {% endfor %}
  {% endif %}
</div>
{% endif %}

<!-- ═══════════════════════════════════════════════════════ -->
<!-- PARTIE V : SCORING ET STRATEGIE -->
<!-- ═══════════════════════════════════════════════════════ -->
<div class="part-header">
  PARTIE V — Scoring et strategie
  <div class="part-header-sub">Criteres d'attribution, simulation note, Go/No-Go detaille</div>
</div>

<!-- ═══════════ CRITÈRES D'ATTRIBUTION ═══════════ -->
{% if criteria %}
<div class="page page-break">
  <h1><span class="section-num">15</span> Criteres d'attribution</h1>

  {% if criteria.evaluation.eligibility_conditions %}
  <h3>Conditions d'éligibilité ({{ criteria.evaluation.eligibility_conditions|length }})</h3>
  <table>
    <tr><th style="width:80%">Condition</th><th style="width:20%">Type</th></tr>
    {% for c in (criteria.evaluation.eligibility_conditions or []) %}
    <tr>
      <td>{{ (c.condition or '')|trunc(120) }}</td>
      <td><span class="badge {% if c.type == 'hard' %}badge-red{% else %}badge-yellow{% endif %}">{{ c.type|upper if c.type else '—' }}</span></td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  {% if criteria.evaluation.scoring_criteria %}
  <h3>Grille de notation</h3>
  <table>
    <tr><th style="width:40%">Critère</th><th style="width:15%">Pondération</th><th style="width:45%">Notes et recommandations</th></tr>
    {% for c in (criteria.evaluation.scoring_criteria or []) %}
    <tr>
      <td><strong>{{ c.criterion }}</strong></td>
      <td><span class="badge badge-blue">{% if c.weight_percent %}{{ c.weight_percent }}%{% else %}N/S{% endif %}</span></td>
      <td>{{ c.notes or '-' }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}
</div>
{% endif %}
<!-- ═══════════ SIMULATION NOTE ACHETEUR ═══════════ -->
{% if scoring_simulation %}
<div class="page page-break">
  <h1><span class="section-num">16</span> Simulation de la note acheteur</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Estimation de votre note finale basee sur les criteres de notation du DCE et votre profil entreprise</p>

  <div class="gonogo-box" style="background: #EFF6FF; border: 2px solid #1E40AF;">
    <div class="gonogo-score" style="color: #1E40AF;">{{ "%.1f"|format(scoring_simulation.total_score or scoring_simulation.note_globale_estimee or 0) }}/20</div>
    <div class="gonogo-label" style="color: #1E40AF;">Note estimée</div>
    {% if scoring_simulation.rank_estimate or scoring_simulation.classement_probable %}
    <div style="font-size: 10px; margin-top: 4px; color: #64748B;">Position estimée : {{ scoring_simulation.rank_estimate or scoring_simulation.classement_probable }}</div>
    {% endif %}
  </div>

  {% set criteria_list = scoring_simulation.criteria_scores or scoring_simulation.dimensions or [] %}
  {% if criteria_list %}
  <h3>Detail par critere</h3>
  <table>
    <tr><th style="width:22%">Critere</th><th style="width:8%">Poids</th><th style="width:10%">Note</th><th style="width:30%">Justification</th><th style="width:30%">Conseils amelioration</th></tr>
    {% for cs in criteria_list %}
    <tr>
      <td><strong>{{ (cs.criterion or '')|trunc(35) }}</strong></td>
      <td>{{ cs.weight or cs.weight_pct or '' }}%</td>
      <td><span class="badge {% if (cs.score or cs.estimated_score or 0) >= 14 %}badge-green{% elif (cs.score or cs.estimated_score or 0) >= 10 %}badge-yellow{% else %}badge-red{% endif %}">{{ "%.1f"|format(cs.score or cs.estimated_score or 0) }}/20</span></td>
      <td style="font-size: 8px;">{{ (cs.justification or cs.recommendation or '-')|trunc(90) }}</td>
      <td style="font-size: 8px;">{% if cs.tips_to_improve %}{% for tip in (cs.tips_to_improve or [])[:2] %}• {{ tip|trunc(50) }}<br>{% endfor %}{% else %}{{ (cs.recommendation or '-')|trunc(80) }}{% endif %}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  {% set levers = scoring_simulation.improvement_levers or [] %}
  {% set axes = scoring_simulation.axes_amelioration or [] %}
  {% if levers %}
  <h3>Leviers d'amélioration (gain potentiel)</h3>
  <table>
    <tr><th style="width:35%">Action</th><th style="width:12%">Gain</th><th style="width:53%">Detail</th></tr>
    {% for lever in levers[:5] %}
    <tr>
      <td><strong>{{ lever.action|trunc(50) }}</strong></td>
      <td><span class="badge badge-green">+{{ "%.1f"|format(lever.gain or 0) }}</span></td>
      <td style="font-size: 8px;">{{ (lever.detail or '')|trunc(130) }}</td>
    </tr>
    {% endfor %}
  </table>
  {% elif axes %}
  <h3>Axes d'amélioration prioritaires</h3>
  <table>
    <tr><th style="width:5%">#</th><th style="width:95%">Recommandation</th></tr>
    {% for axe in axes[:5] %}
    <tr>
      <td>{{ loop.index }}</td>
      <td style="font-size: 8px;">{{ axe|trunc(200) }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}

  <div class="info-box">
    <strong>Note :</strong> Cette simulation est indicative. La note réelle dépend de la qualité de votre offre technique,
    du nombre de concurrents et de l'appréciation subjective de la commission d'attribution.
  </div>
</div>
{% endif %}

<!-- ═══════════ GO/NO-GO DÉTAILLÉ (9 dimensions) ═══════════ -->
{% if gonogo and (gonogo.dimension_scores or gonogo.breakdown) %}
<div class="page page-break">
  <h1><span class="section-num">17</span> Go/No-Go detaille — 9 dimensions</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Analyse detaillee de l'adequation du profil entreprise avec les exigences du marche</p>

  {% set reco = gonogo.recommendation|upper if gonogo.recommendation else 'ATTENTION' %}
  <div class="gonogo-box {% if reco == 'GO' %}gonogo-go{% elif reco == 'NO-GO' %}gonogo-nogo{% else %}gonogo-attention{% endif %}">
    <div class="gonogo-score">{{ gonogo.score }}/100</div>
    <div class="gonogo-label">{{ reco }}</div>
    {% if gonogo.summary %}<div style="font-size: 10px; margin-top: 4px;">{{ gonogo.summary }}</div>{% endif %}
  </div>

  <h3>Scores par dimension</h3>
  <table>
    <tr><th style="width:30%">Dimension</th><th style="width:12%">Score</th><th style="width:12%">Niveau</th><th style="width:46%">Explication</th></tr>
    {% set dim_labels = {'financial_capacity': 'Capacite financiere', 'market_size': 'Taille du marche', 'certifications': 'Certifications requises', 'geographic_zone': 'Zone geographique', 'insurance_adequacy': 'Couverture assurances', 'margin_viability': 'Viabilite de la marge', 'capacity_charge': 'Charge de travail', 'subcontracting_coverage': 'Couverture sous-traitance', 'historical_success': 'Historique succes', 'technical_fit': 'Adequation technique', 'timeline_feasibility': 'Faisabilite planning', 'competitive_position': 'Position concurrentielle'} %}
    {% if gonogo.dimension_scores %}
      {% for dim_key, dim_val in [('financial_capacity', gonogo.dimension_scores.financial_capacity), ('market_size', gonogo.dimension_scores.market_size), ('certifications', gonogo.dimension_scores.certifications), ('geographic_zone', gonogo.dimension_scores.geographic_zone), ('insurance_adequacy', gonogo.dimension_scores.insurance_adequacy), ('margin_viability', gonogo.dimension_scores.margin_viability), ('capacity_charge', gonogo.dimension_scores.capacity_charge), ('subcontracting_coverage', gonogo.dimension_scores.subcontracting_coverage), ('historical_success', gonogo.dimension_scores.historical_success)] %}
      {% if dim_val is not none %}
      <tr>
        <td><strong>{{ dim_labels[dim_key] or dim_key }}</strong></td>
        <td>
          <strong>{{ dim_val }}/100</strong>
          <div class="confidence-bar" style="margin-top: 4px;"><div class="confidence-fill" style="width: {{ dim_val }}%; background: {% if dim_val >= 70 %}#059669{% elif dim_val >= 50 %}#d97706{% else %}#dc2626{% endif %};"></div></div>
        </td>
        <td><span class="badge {% if dim_val >= 70 %}badge-green{% elif dim_val >= 50 %}badge-yellow{% else %}badge-red{% endif %}">{% if dim_val >= 70 %}BON{% elif dim_val >= 50 %}MOYEN{% else %}FAIBLE{% endif %}</span></td>
        <td style="font-size: 8px;"></td>
      </tr>
      {% endif %}
      {% endfor %}
    {% elif gonogo.breakdown %}
      {% for dim_name, dim_score in [('Adequation technique', gonogo.breakdown.technical_fit), ('Capacite financiere', gonogo.breakdown.financial_capacity), ('Faisabilite planning', gonogo.breakdown.timeline_feasibility), ('Position concurrentielle', gonogo.breakdown.competitive_position)] %}
      {% if dim_score is not none %}
      <tr>
        <td><strong>{{ dim_name }}</strong></td>
        <td>
          <strong>{{ dim_score }}/100</strong>
          <div class="confidence-bar" style="margin-top: 4px;"><div class="confidence-fill" style="width: {{ dim_score }}%; background: {% if dim_score >= 70 %}#059669{% elif dim_score >= 50 %}#d97706{% else %}#dc2626{% endif %};"></div></div>
        </td>
        <td><span class="badge {% if dim_score >= 70 %}badge-green{% elif dim_score >= 50 %}badge-yellow{% else %}badge-red{% endif %}">{% if dim_score >= 70 %}BON{% elif dim_score >= 50 %}MOYEN{% else %}FAIBLE{% endif %}</span></td>
        <td style="font-size: 8px;"></td>
      </tr>
      {% endif %}
      {% endfor %}
    {% endif %}
  </table>

  {% if gonogo.profile_gaps %}
  <h3 style="color: #DC2626;">Lacunes identifiees du profil</h3>
  <div class="warning-box">
    {% for gap in (gonogo.profile_gaps or []) %}
    <div style="margin: 3px 0; font-size: 10px;">• {{ gap }}</div>
    {% endfor %}
  </div>
  {% endif %}

  {% if gonogo.profile_strengths %}
  <h3 style="color: #16A34A;">Points forts identifies</h3>
  <div class="financial-box">
    {% for s in (gonogo.profile_strengths or []) %}
    <div style="margin: 3px 0; font-size: 10px;">• {{ s }}</div>
    {% endfor %}
  </div>
  {% endif %}

  <div class="info-box">
    <strong>Methodologie :</strong> Le score Go/No-Go est calcule a partir de 9 dimensions ponderees
    (capacite financiere, certifications, zone geographique, assurances, marge, charge, sous-traitance,
    taille marche, historique). Le profil entreprise est croise avec les exigences du DCE.
  </div>
</div>
{% endif %}

<!-- ═══════════════════════════════════════════════════════ -->
<!-- PARTIE VI : PLAN D'ACTION -->
<!-- ═══════════════════════════════════════════════════════ -->
<div class="part-header">
  PARTIE VI — Plan d'action
  <div class="part-header-sub">Risques, questions pour l'acheteur, calendrier</div>
</div>

{% if summary %}
<!-- ═══════════ RISQUES & ACTIONS ═══════════ -->
<div class="page page-break">
  <h1><span class="section-num">18</span> Analyse des risques</h1>

  <h3>Risques identifiés ({{ (summary.risks or [])|length }})</h3>
  <table>
    <tr><th style="width:20%">Risque</th><th style="width:8%">Severite</th><th style="width:40%">Analyse</th><th style="width:32%">Mitigation</th></tr>
    {% for r in (summary.risks or []) %}
    <tr class="risk-{{ r.severity }}">
      <td><strong>{{ r.risk|trunc(50) }}</strong></td>
      <td><span class="badge {% if r.severity == 'high' %}badge-red{% elif r.severity == 'medium' %}badge-yellow{% else %}badge-gray{% endif %}">{{ (r.severity or '')|upper }}</span></td>
      <td style="font-size: 8px;">{{ r.why|trunc(140) }}</td>
      <td style="font-size: 8px;">{{ (r.mitigation or '')|trunc(100) }}</td>
    </tr>
    {% endfor %}
  </table>

  <h3>Plan d'actions sous 48h ({{ (summary.actions_next_48h or [])|length }} actions)</h3>
  <table>
    <tr><th style="width:45%">Action</th><th style="width:25%">Responsable</th><th style="width:15%">Priorite</th><th style="width:15%">Delai</th></tr>
    {% for a in (summary.actions_next_48h or []) %}
    <tr>
      <td style="font-size: 8px;">{{ a.action|trunc(110) }}</td>
      <td style="font-size: 8px;">{{ a.owner_role|trunc(30) }}</td>
      <td><span class="badge {% if a.priority == 'P0' %}badge-red{% elif a.priority == 'P1' %}badge-yellow{% else %}badge-gray{% endif %}">{{ a.priority }}</span></td>
      <td style="font-size: 7px;">{{ (a.deadline_relative or '')|trunc(20) }}</td>
    </tr>
    {% endfor %}
  </table>
</div>

{% endif %}
<!-- ═══════════ QUESTIONS POUR L'ACHETEUR ═══════════ -->
{% if questions_list %}
<div class="page page-break">
  <h1><span class="section-num">19</span> Questions prioritaires pour l'acheteur</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Questions a poser a l'acheteur avant la date limite — classees par priorite</p>

  <div class="info-box">
    <strong>Conseil :</strong> Posez ces questions via la plateforme de dématérialisation
    avant la date limite de questions. Les réponses seront diffusées à tous les candidats.
  </div>

  <table>
    <tr><th style="width:5%">#</th><th style="width:10%">Priorite</th><th style="width:50%">Question</th><th style="width:35%">Justification</th></tr>
    {% for q in questions_list[:15] %}
    <tr>
      <td>{{ loop.index }}</td>
      <td><span class="badge {% if q.priority == 'high' %}badge-red{% elif q.priority == 'medium' %}badge-yellow{% else %}badge-gray{% endif %}">{{ q.priority|upper if q.priority else 'INFO' }}</span></td>
      <td style="font-size: 8px;">{{ q.question|trunc(140) }}</td>
      <td style="font-size: 8px;">{{ (q.justification or '')[:100] }}{% if q.justification and q.justification|length > 100 %}...{% endif %}</td>
    </tr>
    {% endfor %}
  </table>
</div>
{% endif %}

<!-- ═══════════ DATES CLÉS ═══════════ -->
{% if timeline %}
<div class="page page-break">
  <h1><span class="section-num">20</span> Calendrier et dates cles</h1>
  <table>
    <tr><th style="width:60%">Echeance</th><th style="width:40%">Date / Duree</th></tr>
    {% if timeline.submission_deadline %}
    <tr class="risk-high">
      <td><strong>Date limite de remise</strong> <span style="font-size:7px; color:#b91c1c;">(imperatif)</span></td>
      <td><span class="badge badge-red">{{ timeline.submission_deadline|datefr }}</span></td>
    </tr>
    {% endif %}
    {% if timeline.execution_duration_months %}
    <tr>
      <td style="font-size:9px;">Duree d'execution</td>
      <td style="font-size:9px;">{{ timeline.execution_duration_months }} mois</td>
    </tr>
    {% endif %}
    {% for kd in (timeline.key_dates or []) %}
    {% set kd_label_low = (kd.label or '')|lower %}
    {% if 'remise des offres' not in kd_label_low %}
    <tr>
      <td style="font-size:9px;">{{ (kd.label or '')|trunc(55,'') }}{% if kd.mandatory %} <span style="font-size:7px; color:#b91c1c;">(oblig.)</span>{% endif %}</td>
      <td style="font-size:9px;">{{ kd.date|datefr }}</td>
    </tr>
    {% endif %}
    {% endfor %}
  </table>
</div>
{% endif %}

<!-- ═══════════════════════════════════════════════════════ -->
<!-- PARTIE VII : ANNEXES -->
<!-- ═══════════════════════════════════════════════════════ -->
<div class="part-header">
  ANNEXES
  <div class="part-header-sub">Inventaire des documents, glossaire BTP, avertissement IA</div>
</div>

<!-- ═══════════ INVENTAIRE DOCUMENTS ANALYSÉS ═══════════ -->
{% if documents_inventory %}
<div class="page page-break">
  <h1><span class="section-num">A1</span> Inventaire des documents analyses</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Liste des pieces du DCE prises en compte dans cette analyse</p>

  <table>
    <tr><th style="width:5%">#</th><th style="width:35%">Document</th><th style="width:12%">Type</th><th style="width:10%">Pages</th><th style="width:15%">Taille</th><th style="width:23%">Qualite OCR</th></tr>
    {% for doc in documents_inventory %}
    <tr>
      <td>{{ loop.index }}</td>
      <td style="font-size: 8px;">{{ doc.name|trunc(45) }}</td>
      <td><span class="badge badge-blue">{{ doc.doc_type or 'AUTRES' }}</span></td>
      <td>{{ doc.pages or '—' }}</td>
      <td>{{ doc.size_display }}</td>
      <td>
        {% if doc.ocr_quality %}
        <span class="badge {% if doc.ocr_quality >= 70 %}badge-green{% elif doc.ocr_quality >= 40 %}badge-yellow{% else %}badge-red{% endif %}">{{ "%.0f"|format(doc.ocr_quality) }}%</span>
        {% else %}—{% endif %}
      </td>
    </tr>
    {% endfor %}
  </table>

  <div style="margin-top: 8px; font-size: 9px; color: #64748B;">
    Total : {{ documents_inventory|length }} document{{ 's' if documents_inventory|length > 1 else '' }}
    {% set total_pages = documents_inventory|sum(attribute='pages') %}
    {% if total_pages %} — {{ total_pages }} pages analysées{% endif %}
  </div>
</div>
{% endif %}

<!-- ═══════════ GLOSSAIRE BTP ═══════════ -->
{% if glossaire_btp %}
<div class="page page-break">
  <h1><span class="section-num">A2</span> Glossaire BTP</h1>
  <p style="color: #64748b; font-size: 10px; margin-top: -8px;">Definitions des termes techniques et acronymes utilises dans ce rapport</p>

  <table>
    <tr><th style="width:18%">Terme</th><th style="width:82%">Definition</th></tr>
    {% for term, definition in glossaire_btp %}
    <tr>
      <td style="font-size: 9px;"><strong>{{ term }}</strong></td>
      <td style="font-size: 8px;">{{ definition|trunc(250) }}</td>
    </tr>
    {% endfor %}
  </table>
</div>
{% endif %}

<!-- ═══════════ FOOTER & DISCLAIMER ═══════════ -->
<div class="page page-break">
  <h1><span class="section-num">A3</span> Avertissement IA et mentions legales</h1>

  <div class="disclaimer">
    <strong>Avertissement :</strong> Ce rapport est genere par intelligence artificielle (Claude, Anthropic) a partir des documents du DCE fournis.
    Il constitue une <strong>aide a la decision</strong> et ne se substitue pas a l'analyse humaine d'un expert marches publics.
    Les informations extraites doivent etre systematiquement verifiees avant toute soumission d'offre.<br><br>
    <strong>Confiance globale de l'analyse :</strong> {{ "%.0f"|format(confidence * 100) if confidence else 'N/A' }}%
    {% if confidence %}
    <div class="confidence-bar" style="margin-top: 8px;">
      <div class="confidence-fill" style="width: {{ "%.0f"|format(confidence * 100) }}%; background: {% if confidence >= 0.8 %}#059669{% elif confidence >= 0.6 %}#d97706{% else %}#dc2626{% endif %};"></div>
    </div>
    {% endif %}
  </div>

  <div class="info-box" style="margin-top: 16px;">
    <strong>Donnees utilisees :</strong> Seuls les documents uploades dans le projet ont ete analyses.
    L'IA ne dispose d'aucune information externe (pas d'acces internet, pas de base de donnees externe).
    La qualite de l'analyse depend directement de la qualite des documents fournis (lisibilite OCR, exhaustivite du DCE).
  </div>

  <div class="footer" style="margin-top: 40px;">
    <div style="margin-bottom: 4px;"><strong style="color: #2563eb; letter-spacing: 2px;">AO COPILOT</strong></div>
    aocopilot.fr — {{ generated_at }} — Rapport confidentiel — Reproduction interdite
  </div>
</div>

</body>
</html>
"""


def generate_export_pdf(db: Session, project_id: str) -> bytes:
    project = db.query(AoProject).filter_by(id=uuid.UUID(project_id)).first()
    if not project:
        raise ValueError("Projet introuvable")

    pid = uuid.UUID(project_id)

    summary_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="summary"
    ).order_by(ExtractionResult.version.desc()).first()

    criteria_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="criteria"
    ).order_by(ExtractionResult.version.desc()).first()

    gonogo_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="gonogo"
    ).order_by(ExtractionResult.version.desc()).first()

    timeline_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="timeline"
    ).order_by(ExtractionResult.version.desc()).first()

    checklist_items = db.query(ChecklistItem).filter_by(
        project_id=pid
    ).order_by(ChecklistItem.criticality, ChecklistItem.category).all()

    # ── Compute checklist statistics ──
    checklist_stats = {"eliminatoire": 0, "important": 0, "info": 0, "ok": 0}
    for item in checklist_items:
        crit = (item.criticality or "").lower()
        status = (item.status or "").upper()
        if "liminatoire" in crit:
            checklist_stats["eliminatoire"] += 1
        elif "important" in crit:
            checklist_stats["important"] += 1
        else:
            checklist_stats["info"] += 1
        if status == "OK":
            checklist_stats["ok"] += 1

    # ── Extract confidence from summary ──
    confidence = None
    if summary_result and summary_result.payload:
        confidence = summary_result.payload.get("confidence_overall") or summary_result.payload.get("confidence")

    # ── Prepare gonogo & timeline payloads ──
    gonogo = gonogo_result.payload if gonogo_result else None
    timeline = timeline_result.payload if timeline_result else None

    # Convert dicts to object-like access for Jinja2 dot notation
    class _DictObj:
        """Allow dict.key access in Jinja2 templates.
        Returns None for missing scalar attrs, [] for missing list attrs."""
        def __init__(self, d):
            self._keys = set()
            for k, v in (d or {}).items():
                self._keys.add(k)
                if isinstance(v, list):
                    setattr(self, k, [_DictObj(i) if isinstance(i, dict) else i for i in v])
                elif isinstance(v, dict):
                    setattr(self, k, _DictObj(v))
                else:
                    setattr(self, k, v)
        def __bool__(self):
            return True
        def __getattr__(self, name):
            if name.startswith('_'):
                raise AttributeError(name)
            return None  # Return None for missing attributes instead of raising
        def __iter__(self):
            return iter([])
        def __len__(self):
            return 0

    gonogo_obj = _DictObj(gonogo) if gonogo else None
    timeline_obj = _DictObj(timeline) if timeline else None

    # ── Compute days remaining until submission deadline ──
    days_remaining = None
    deadline_str = None
    if summary_result and summary_result.payload:
        po = summary_result.payload.get("project_overview", {})
        deadline_str = po.get("deadline_submission")
    if not deadline_str and timeline:
        deadline_str = timeline.get("submission_deadline")
    if deadline_str:
        try:
            dl = deadline_str
            if 'T' in str(dl):
                deadline_dt = datetime.fromisoformat(str(dl).replace('Z', '+00:00')).replace(tzinfo=None)
            else:
                deadline_dt = datetime.strptime(str(dl)[:10], '%Y-%m-%d')
            days_remaining = (deadline_dt - datetime.now()).days
        except (ValueError, TypeError):
            days_remaining = None

    # ── Extract scoring simulation (if available) ──
    scoring_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="scoring"
    ).order_by(ExtractionResult.version.desc()).first()
    scoring_simulation = scoring_result.payload if scoring_result else None

    # ── Extract CCAP analysis: derogations + clauses risquées (if available) ──
    # Try new naming first, fallback to old naming for legacy projects
    ccap_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="ccap_risks"
    ).order_by(ExtractionResult.version.desc()).first()
    if not ccap_result:
        ccap_result = db.query(ExtractionResult).filter_by(
            project_id=pid, result_type="ccap"
        ).order_by(ExtractionResult.version.desc()).first()
    ccag_derogations = None
    ccap_clauses_risquees = None
    if ccap_result and ccap_result.payload:
        ccag_derogations = ccap_result.payload.get("ccag_derogations") or ccap_result.payload.get("derogations")
        ccap_clauses_risquees = ccap_result.payload.get("clauses_risquees")

    # ── Extract RC analysis for fiche signalétique (if available) ──
    rc_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="rc_analysis"
    ).order_by(ExtractionResult.version.desc()).first()
    if not rc_result:
        rc_result = db.query(ExtractionResult).filter_by(
            project_id=pid, result_type="rc"
        ).order_by(ExtractionResult.version.desc()).first()
    rc_analysis = rc_result.payload if rc_result else None

    # ── Extract questions for buyer (if available) ──
    questions_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="questions"
    ).order_by(ExtractionResult.version.desc()).first()
    questions_list = None
    if questions_result and questions_result.payload:
        questions_list = questions_result.payload.get("questions") or questions_result.payload.get("priority_questions")

    # ── Extract CCTP analysis (if available) ──
    cctp_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="cctp_analysis"
    ).order_by(ExtractionResult.version.desc()).first()
    if not cctp_result:
        cctp_result = db.query(ExtractionResult).filter_by(
            project_id=pid, result_type="cctp"
        ).order_by(ExtractionResult.version.desc()).first()
    cctp_analysis = cctp_result.payload if cctp_result else None

    # ── Extract AE (Acte d'Engagement) analysis (if available) ──
    ae_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="ae_analysis"
    ).order_by(ExtractionResult.version.desc()).first()
    ae_analysis = ae_result.payload if ae_result else None

    # ── Extract DC Check (administrative documents) (if available) ──
    dc_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="dc_check"
    ).order_by(ExtractionResult.version.desc()).first()
    dc_check = dc_result.payload if dc_result else None

    # ── Extract Conflicts (cross-document contradictions) (if available) ──
    conflicts_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="conflicts"
    ).order_by(ExtractionResult.version.desc()).first()
    conflicts = conflicts_result.payload if conflicts_result else None

    # ── Extract Cashflow / BFR simulation (if available) ──
    cashflow_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="cashflow_simulation"
    ).order_by(ExtractionResult.version.desc()).first()
    if not cashflow_result:
        cashflow_result = db.query(ExtractionResult).filter_by(
            project_id=pid, result_type="cashflow"
        ).order_by(ExtractionResult.version.desc()).first()
    cashflow_data = cashflow_result.payload if cashflow_result else None

    # ── Extract Subcontracting analysis (if available) ──
    subcontracting_result = db.query(ExtractionResult).filter_by(
        project_id=pid, result_type="subcontracting"
    ).order_by(ExtractionResult.version.desc()).first()
    subcontracting = subcontracting_result.payload if subcontracting_result else None

    # ── DPGF Pricing benchmark (if DPGF/BPU documents exist) ──
    dpgf_pricing = None
    try:
        dpgf_result = db.query(ExtractionResult).filter_by(
            project_id=pid, result_type="dpgf_pricing"
        ).order_by(ExtractionResult.version.desc()).first()
        if dpgf_result and dpgf_result.payload:
            dpgf_pricing = dpgf_result.payload.get("lines") or dpgf_result.payload.get("pricing_lines")
        if not dpgf_pricing:
            # Try on-the-fly from DPGF extraction data
            dpgf_extract = db.query(ExtractionResult).filter_by(
                project_id=pid, result_type="dpgf_extraction"
            ).order_by(ExtractionResult.version.desc()).first()
            if dpgf_extract and dpgf_extract.payload:
                rows = dpgf_extract.payload.get("rows") or dpgf_extract.payload.get("lines") or []
                if rows:
                    from app.services.btp_pricing import check_dpgf_pricing
                    # Detect region from project location
                    region = "france"
                    if summary_result and summary_result.payload:
                        loc = summary_result.payload.get("project_overview", {}).get("location", "")
                        if loc:
                            region = loc.lower().replace(" ", "-")
                    dpgf_pricing = check_dpgf_pricing(rows, region=region)
    except Exception as exc:
        logger.warning(f"DPGF pricing benchmark non disponible: {exc}")
        dpgf_pricing = None

    # ── Build glossaire BTP (filtered to terms used in the report) ──
    glossaire_btp = None
    try:
        from app.services.btp_knowledge import BTP_GLOSSARY
        # Select key terms most likely referenced in an AO report
        priority_terms = [
            "RC", "CCTP", "CCAP", "DPGF", "BPU", "AE", "DCE", "CCAG", "CCAG-Travaux",
            "DC1", "DC2", "DC4", "DOE", "DIUO", "OPR", "GPA",
            "Retenue de garantie", "Caution bancaire", "Avance forfaitaire",
            "Pénalités de retard", "Révision de prix", "Intérêts moratoires",
            "Sous-traitance", "Mandataire", "Cotraitant",
            "Garantie décennale", "Garantie biennale", "RC Pro",
            "Allotissement", "Variante", "MAPA", "NF DTU",
            "PPSPS", "VRD", "BIM",
        ]
        glossaire_btp = [(t, BTP_GLOSSARY[t]) for t in priority_terms if t in BTP_GLOSSARY]
    except Exception as exc:
        logger.warning(f"Glossaire BTP non disponible: {exc}")
        glossaire_btp = None

    # ── Build documents inventory ──
    docs = db.query(AoDocument).filter_by(
        project_id=pid
    ).order_by(AoDocument.doc_type, AoDocument.original_name).all()
    documents_inventory = []
    for doc in docs:
        size_kb = doc.file_size_kb or 0
        size_display = f"{size_kb} Ko" if size_kb < 1024 else f"{size_kb / 1024:.1f} Mo"
        documents_inventory.append({
            "name": doc.original_name,
            "doc_type": doc.doc_type,
            "pages": doc.page_count or 0,
            "size_display": size_display,
            "ocr_quality": doc.ocr_quality_score,
        })

    env = Environment(loader=BaseLoader(), autoescape=True)

    # Custom Jinja2 filter to format ISO dates nicely
    def format_date_fr(value):
        """Convert ISO date string to French format: 15/04/2026 à 12h00."""
        if not value or not isinstance(value, str):
            return value or ''
        try:
            # Handle ISO datetime: 2026-04-15T12:00:00
            if 'T' in str(value):
                dt = datetime.fromisoformat(str(value).replace('Z', '+00:00'))
                return dt.strftime('%d/%m/%Y à %Hh%M')
            # Handle ISO date: 2026-04-15
            dt = datetime.strptime(str(value)[:10], '%Y-%m-%d')
            return dt.strftime('%d/%m/%Y')
        except (ValueError, TypeError):
            return str(value)

    env.filters['datefr'] = format_date_fr

    def safe_truncate(value, length=100, suffix='...'):
        """Safely truncate a value, handling None."""
        if not value:
            return ''
        s = str(value)
        if len(s) <= length:
            return s
        return s[:length] + suffix

    env.filters['trunc'] = safe_truncate

    template = env.from_string(EXPORT_TEMPLATE)

    # ── Génération des graphiques (graceful degradation si matplotlib absent) ──
    radar_chart_b64: str | None = None
    cashflow_chart_b64: str | None = None
    heatmap_chart_b64: str | None = None
    pricing_chart_b64: str | None = None
    try:
        from app.services.chart_generator import (
            generate_gonogo_radar,
            generate_cashflow_chart,
            generate_risk_heatmap,
            generate_pricing_benchmark_bars,
            chart_to_base64,
        )
        if gonogo:
            _dims = gonogo.get("dimension_scores") or gonogo.get("breakdown") or {}
            _score = gonogo.get("score")
            _title = project.title or "AO"
            radar_chart_b64 = chart_to_base64(
                generate_gonogo_radar(_dims if isinstance(_dims, dict) else {}, _score, _title)
            )
        if cashflow_data:
            cashflow_chart_b64 = chart_to_base64(
                generate_cashflow_chart(cashflow_data, project.title or "AO")
            )
        if conflicts:
            _conflict_list = conflicts.get("conflicts") or conflicts.get("items") or []
            heatmap_chart_b64 = chart_to_base64(
                generate_risk_heatmap(_conflict_list, project.title or "AO")
            )
        if dpgf_pricing:
            pricing_chart_b64 = chart_to_base64(
                generate_pricing_benchmark_bars(dpgf_pricing, project.title or "AO")
            )
    except Exception as _chart_err:
        logger.warning("chart_generation_skipped", error=str(_chart_err))

    try:
        html_content = template.render(
            project=project,
            summary=_DictObj(summary_result.payload) if summary_result and summary_result.payload else None,
            checklist_items=checklist_items,
            criteria=_DictObj(criteria_result.payload) if criteria_result and criteria_result.payload else None,
            gonogo=gonogo_obj,
            timeline=timeline_obj,
            checklist_stats=checklist_stats,
            confidence=confidence,
            days_remaining=days_remaining,
            scoring_simulation=_DictObj(scoring_simulation) if scoring_simulation else None,
            ccag_derogations=[_DictObj(d) for d in ccag_derogations] if ccag_derogations else None,
            ccap_clauses_risquees=[_DictObj(c) for c in ccap_clauses_risquees] if ccap_clauses_risquees else None,
            rc_analysis=_DictObj(rc_analysis) if rc_analysis else None,
            questions_list=[_DictObj(q) if isinstance(q, dict) else q for q in questions_list] if questions_list else None,
            documents_inventory=[_DictObj(d) for d in documents_inventory] if documents_inventory else None,
            cctp_analysis=_DictObj(cctp_analysis) if cctp_analysis else None,
            cashflow_data=_DictObj(cashflow_data) if cashflow_data else None,
            subcontracting=_DictObj(subcontracting) if subcontracting else None,
            ae_analysis=_DictObj(ae_analysis) if ae_analysis else None,
            dc_check=_DictObj(dc_check) if dc_check else None,
            conflicts=_DictObj(conflicts) if conflicts else None,
            dpgf_pricing=[_DictObj(line) for line in dpgf_pricing] if dpgf_pricing else None,
            glossaire_btp=glossaire_btp,
            generated_at=datetime.now().strftime("%d/%m/%Y %H:%M"),
            generation_date=datetime.now().strftime("%d/%m/%Y"),
            theme=get_theme(),
            radar_chart_b64=radar_chart_b64,
            cashflow_chart_b64=cashflow_chart_b64,
            heatmap_chart_b64=heatmap_chart_b64,
            pricing_chart_b64=pricing_chart_b64,
        )
    except Exception as exc:
        logger.error(f"Erreur rendu template Jinja2 pour project {project_id}: {exc}")
        raise RuntimeError(f"Erreur de génération du template PDF: {exc}") from exc

    try:
        from xhtml2pdf import pisa
        output = BytesIO()
        result = pisa.CreatePDF(html_content, dest=output, encoding="utf-8")
        if result.err:
            raise RuntimeError(f"xhtml2pdf errors: {result.err}")
        pdf_bytes = output.getvalue()
    except Exception as exc:
        logger.error(f"Erreur xhtml2pdf pour project {project_id}: {exc}")
        raise RuntimeError(f"Erreur de génération PDF: {exc}") from exc

    return pdf_bytes


def generate_export_docx(db: Session, project_id: str) -> bytes:
    """Génère un rapport Word (.docx) complet — design identique au PDF.
    Headers bleu foncé, encadrés colorés, badges, Go/No-Go box, disclaimer IA.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from io import BytesIO
    except ImportError as e:
        raise RuntimeError("python-docx non installé. Ajoutez python-docx==1.1.0.") from e

    pid = uuid.UUID(project_id)
    project = db.query(AoProject).filter_by(id=pid).first()
    if not project:
        raise ValueError("Projet introuvable")

    # ── Charger TOUTES les analyses ──────────────────────────────────────
    def _get(result_type):
        r = db.query(ExtractionResult).filter_by(
            project_id=pid, result_type=result_type
        ).order_by(ExtractionResult.version.desc()).first()
        return r.payload if r and r.payload else {}

    summary = _get("summary")
    criteria = _get("criteria")
    timeline = _get("timeline")
    gonogo = _get("gonogo")
    ccap = _get("ccap_risks") or _get("ccap")
    rc = _get("rc_analysis") or _get("rc")
    cctp = _get("cctp_analysis") or _get("cctp")
    questions_data = _get("questions")
    scoring = _get("scoring")
    cashflow = _get("cashflow_simulation") or _get("cashflow")
    subcontracting = _get("subcontracting")
    ae = _get("ae_analysis")
    dc_check = _get("dc_check")
    conflicts_data = _get("conflicts")

    # DPGF Pricing
    dpgf_pricing = None
    try:
        dpgf_result = db.query(ExtractionResult).filter_by(
            project_id=pid, result_type="dpgf_pricing"
        ).order_by(ExtractionResult.version.desc()).first()
        if dpgf_result and dpgf_result.payload:
            dpgf_pricing = dpgf_result.payload if isinstance(dpgf_result.payload, list) else dpgf_result.payload.get("lines", [])
        if not dpgf_pricing:
            dpgf_extract = db.query(ExtractionResult).filter_by(
                project_id=pid, result_type="dpgf_extraction"
            ).order_by(ExtractionResult.version.desc()).first()
            if dpgf_extract and dpgf_extract.payload:
                rows = dpgf_extract.payload.get("lines") or dpgf_extract.payload.get("rows") or []
                if rows:
                    from app.services.btp_pricing import check_dpgf_pricing
                    dpgf_pricing = check_dpgf_pricing(rows)
    except Exception:
        dpgf_pricing = None

    # Glossaire BTP
    try:
        from app.services.btp_knowledge import BTP_GLOSSARY
        priority_terms = [
            "RC", "CCTP", "CCAP", "DPGF", "BPU", "AE", "DCE", "CCAG",
            "SOPAQ", "DOE", "DIUO", "PPSPS", "OPR", "DQE", "SOGED",
            "MOA", "MOE", "AMO", "CSPS", "BET", "DTU", "NF",
            "Retenue de garantie", "Avance forfaitaire", "Pénalités de retard",
            "Sous-traitance", "DC1", "DC2", "DUME", "Allotissement",
        ]
        glossaire_btp = [(t, BTP_GLOSSARY[t]) for t in priority_terms if t in BTP_GLOSSARY]
    except Exception:
        glossaire_btp = []

    checklist_items = db.query(ChecklistItem).filter_by(
        project_id=pid
    ).order_by(ChecklistItem.criticality, ChecklistItem.category).all()

    documents = db.query(AoDocument).filter_by(project_id=pid).all()

    po = summary.get("project_overview", {})

    doc = Document()

    # ── Couleurs PDF ─────────────────────────────────────────────────────
    DARK_BLUE = RGBColor(0x0F, 0x1B, 0x4C)
    ACCENT_BLUE = RGBColor(0x25, 0x63, 0xEB)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    RED = RGBColor(0xB9, 0x1C, 0x1C)
    GREEN_DARK = RGBColor(0x04, 0x78, 0x57)
    ORANGE = RGBColor(0xB4, 0x53, 0x09)
    GRAY = RGBColor(0x64, 0x74, 0x8B)
    GRAY_LIGHT = RGBColor(0x9C, 0xA3, 0xAF)

    # ── Style global ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)  # 11pt — BTP audit: lisibilité comités (50+ ans)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = DARK_BLUE
        hs.font.name = "Calibri"

    # ── Helper: cell shading ─────────────────────────────────────────────
    def _shade_cell(cell, hex_color):
        """Apply background color to a table cell."""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _shade_row(row, hex_color):
        for cell in row.cells:
            _shade_cell(cell, hex_color)

    def _set_cell_text(cell, text, bold=False, color=None, size=None):
        """Set cell text with formatting — clears default and adds a run."""
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = bold
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = size
        run.font.name = "Calibri"
        return run

    # ── Helper: styled table with dark blue header row ────────────────
    def _styled_table(headers, col_widths=None):
        """Create table with dark blue (#0F1B4C) header row, white text."""
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            _shade_cell(cell, "0F1B4C")
            _set_cell_text(cell, h, bold=True, color=WHITE, size=Pt(9))
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        return t

    # ── Helper: format date ──────────────────────────────────────────────
    def _fmt_date(val):
        if not val:
            return "—"
        s = str(val)
        if "T" in s:
            s = s.split("T")[0]
        try:
            from datetime import datetime as dt
            d = dt.strptime(s, "%Y-%m-%d")
            return d.strftime("%d/%m/%Y")
        except Exception:
            return s

    # ── Helper: colored encadré (info-box, warning-box, etc.) ────────
    def _add_encadre(text, bg_hex="EFF6FF", border_hex="2563EB", text_color=None):
        """Add a bordered paragraph mimicking PDF encadrés."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        # Background shading
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_hex)
        pPr.append(shd)
        # Left border
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "36")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), border_hex)
        pBdr.append(left)
        pPr.append(pBdr)
        # Indentation for padding effect
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "180")
        ind.set(qn("w:right"), "180")
        pPr.append(ind)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        if text_color:
            run.font.color.rgb = text_color
        return p

    # ── Helper: severity badge text ──────────────────────────────────────
    def _severity_color(sev):
        s = (sev or "").lower()
        if s in ("high", "fort", "haut", "éliminatoire", "eliminatoire"):
            return RED
        elif s in ("medium", "moyen", "moyenne"):
            return ORANGE
        elif s in ("low", "bas", "faible"):
            return GREEN_DARK
        return GRAY

    def _severity_bg(sev):
        s = (sev or "").lower()
        if s in ("high", "fort", "haut", "éliminatoire", "eliminatoire"):
            return "FEF2F2"
        elif s in ("medium", "moyen", "moyenne"):
            return "FFFBEB"
        return None

    def _status_color(status):
        s = (status or "").upper()
        if s == "OK":
            return GREEN_DARK
        elif s in ("MANQUANT", "NON", "KO"):
            return RED
        elif s in ("PARTIEL", "EN COURS"):
            return ORANGE
        return GRAY

    def _add_part_header(title, subtitle=""):
        """Add PARTIE visual separator (dark blue background, white text)."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0F1B4C")
        pPr.append(shd)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:right"), "0")
        pPr.append(ind)
        run_t = p.add_run(title)
        run_t.bold = True
        run_t.font.size = Pt(13)
        run_t.font.color.rgb = WHITE
        run_t.font.name = "Calibri"
        if subtitle:
            run_s = p.add_run(f"\n{subtitle}")
            run_s.font.size = Pt(9)
            run_s.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            run_s.font.name = "Calibri"

    # ══════════════════════════════════════════════════════════════════════
    #                        PAGE DE GARDE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    # Brand
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_b = brand.add_run("AO COPILOT")
    run_b.font.color.rgb = ACCENT_BLUE
    run_b.font.size = Pt(14)
    run_b.bold = True
    run_b.font.name = "Calibri"

    # Separator line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = line.add_run("━━━━━━━━━━")
    run_l.font.color.rgb = ACCENT_BLUE
    run_l.font.size = Pt(12)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sub.add_run("RAPPORT D'ANALYSE DCE")
    run_s.font.color.rgb = GRAY
    run_s.font.size = Pt(11)
    run_s.font.name = "Calibri"

    doc.add_paragraph()

    # Title
    title = doc.add_heading(project.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(24)

    # Buyer
    buyer_p = doc.add_paragraph()
    buyer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_buyer = buyer_p.add_run(po.get("buyer") or project.buyer or "Acheteur public")
    run_buyer.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    run_buyer.font.size = Pt(14)

    doc.add_paragraph()

    # Info table (cover)
    cover_info = _styled_table(["Information", "Détail"], [7, 10])
    for label, value in [
        ("Référence", project.reference or "N/A"),
        ("Lieu", po.get("location", "—")),
        ("Type de marché", po.get("market_type", "—")),
        ("Date limite", _fmt_date(po.get("deadline_submission"))),
        ("Budget estimé", po.get("estimated_budget") or "Non précisé"),
        ("Documents analysés", f"{len(documents)} documents"),
        ("Généré le", datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]:
        row = cover_info.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=Pt(9))
        _set_cell_text(row[1], str(value), size=Pt(9))

    doc.add_paragraph()

    # ── Go/No-Go badge on cover ──────────────────────────────────────────
    if gonogo:
        score = gonogo.get("score", 0)
        rec = (gonogo.get("recommendation") or "ATTENTION").upper()
        gonogo_p = doc.add_paragraph()
        gonogo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_rec = gonogo_p.add_run(f"RECOMMANDATION : {rec}  —  {score}/100")
        run_rec.bold = True
        run_rec.font.size = Pt(14)
        if rec == "GO":
            run_rec.font.color.rgb = GREEN_DARK
        elif rec == "NO-GO":
            run_rec.font.color.rgb = RED
        else:
            run_rec.font.color.rgb = ORANGE

    # Cover footer
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = meta_p.add_run(
        f"Généré le {datetime.now().strftime('%d/%m/%Y %H:%M')} | Rapport confidentiel"
    )
    run_meta.font.color.rgb = GRAY_LIGHT
    run_meta.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #                          SOMMAIRE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Sommaire", 1)

    ccag_derogations = ccap.get("ccag_derogations") or ccap.get("derogations", [])
    ccap_clauses_risquees = ccap.get("clauses_risquees", [])
    questions_list = questions_data.get("questions", [])

    # Sommaire structuré en 7 parties (miroir PDF)
    toc_items = [
        ("", "PARTIE I — DÉCISION STRATÉGIQUE"),
        ("1.", "Synthèse décisionnelle — Recommandation Go/No-Go et indicateurs clés"),
    ]
    if rc:
        toc_items.append(("2.", "Fiche signalétique du marché — Procédure, allotissement, groupement"))
    toc_items.append(("3.", "Résumé exécutif — Objet du marché, points clés extraits"))
    toc_items.append(("", "PARTIE II — CONFORMITÉ JURIDIQUE"))
    if dc_check:
        toc_items.append(("4.", "Vérification administrative DC — Attestations, certifications"))
    if checklist_items:
        toc_items.append(("5.", f"Checklist de conformité — {len(checklist_items)} exigences"))
    if ccag_derogations:
        toc_items.append(("6.", f"Dérogations CCAG-Travaux 2021 — {len(ccag_derogations)} détectées"))
    if ccap_clauses_risquees:
        toc_items.append(("7.", f"Clauses risquées CCAP — {len(ccap_clauses_risquees)} clauses"))
    toc_items.append(("", "PARTIE III — ANALYSE FINANCIÈRE"))
    if ae:
        toc_items.append(("8.", "Analyse Acte d'Engagement — Prix, pénalités, garanties"))
    toc_items.append(("9.", "Synthèse financière — Montants, avance, pénalités, révision"))
    toc_items.append(("10.", "Benchmark tarifaire DPGF — Prix unitaires vs références marché"))
    if cashflow:
        toc_items.append(("11.", "Simulation trésorerie et BFR — Impact financier prévisionnel"))
    toc_items.append(("", "PARTIE IV — ANALYSE TECHNIQUE"))
    if cctp:
        toc_items.append(("12.", "Analyse technique CCTP — Exigences, normes DTU, contradictions"))
    if conflicts_data:
        toc_items.append(("13.", f"Contradictions inter-documents — {conflicts_data.get('nb_total', 0)} incohérences"))
    if subcontracting:
        toc_items.append(("14.", "Analyse sous-traitance — Obligations, risques, conformité"))
    toc_items.append(("", "PARTIE V — SCORING ET STRATÉGIE"))
    if criteria:
        toc_items.append(("15.", "Critères d'attribution — Conditions d'éligibilité et grille de notation"))
    if scoring:
        toc_items.append(("16.", "Simulation note acheteur — Note estimée et leviers"))
    if gonogo and (gonogo.get("dimension_scores") or gonogo.get("breakdown")):
        toc_items.append(("17.", "Go/No-Go détaillé — 9 dimensions d'analyse"))
    toc_items.append(("", "PARTIE VI — PLAN D'ACTION"))
    toc_items.append(("18.", "Analyse des risques — Risques identifiés et plan d'actions 48h"))
    if questions_list:
        toc_items.append(("19.", f"Questions pour l'acheteur — {len(questions_list)} questions"))
    if timeline:
        toc_items.append(("20.", "Calendrier et dates clés — Échéances de soumission et d'exécution"))
    toc_items.append(("", "ANNEXES"))
    if documents:
        toc_items.append(("A1.", f"Inventaire des documents — {len(documents)} pièces DCE analysées"))
    toc_items.append(("A2.", "Glossaire BTP — Termes clés utilisés dans ce rapport"))
    toc_items.append(("A3.", "Avertissement IA et mentions légales"))

    for num, label in toc_items:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        if not num:
            # PARTIE header — dark blue background, white text
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "0F1B4C")
            pPr.append(shd)
            run_text = p.add_run(label)
            run_text.bold = True
            run_text.font.size = Pt(10)
            run_text.font.color.rgb = WHITE
            run_text.font.name = "Calibri"
        else:
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "dotted")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CBD5E1")
            pBdr.append(bottom)
            pPr.append(pBdr)
            run_num = p.add_run(f"{num}  ")
            run_num.bold = True
            run_num.font.color.rgb = ACCENT_BLUE
            run_num.font.size = Pt(10)
            run_text = p.add_run(label)
            run_text.font.size = Pt(10)
            run_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph()
    _add_encadre(
        "Guide de lecture rapide :\n"
        "5 min — Couverture + Synthèse décisionnelle (pages 1-2)\n"
        "15 min — + Résumé exécutif + Risques + Calendrier\n"
        "30 min — Rapport complet avec checklist et critères détaillés",
        bg_hex="EFF6FF", border_hex="3B82F6",
        text_color=RGBColor(0x1E, 0x3A, 0x8A)
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE I — DÉCISION STRATÉGIQUE
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE I — DÉCISION STRATÉGIQUE",
                     "Recommandation Go/No-Go, fiche marché et résumé exécutif")

    # ── 1. SYNTHÈSE DÉCISIONNELLE ──
    doc.add_heading("1. Synthèse décisionnelle", 1)
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("Vue d'ensemble en 1 page pour les décideurs (DG, Directeur commercial, Responsable AO)")
    run_sub.font.color.rgb = GRAY
    run_sub.font.size = Pt(9)

    if gonogo:
        score = gonogo.get("score", 0)
        rec = (gonogo.get("recommendation") or "ATTENTION").upper()

        # Go/No-Go box
        if rec == "GO":
            bg, border_c = "D1FAE5", "059669"
            txt_color = GREEN_DARK
        elif rec == "NO-GO":
            bg, border_c = "FDE8E8", "DC2626"
            txt_color = RED
        else:
            bg, border_c = "FEF3C7", "D97706"
            txt_color = ORANGE

        gonogo_box = doc.add_paragraph()
        gonogo_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = gonogo_box._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg)
        pPr.append(shd)
        pBdr = OxmlElement("w:pBdr")
        for side in ["top", "left", "bottom", "right"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "18")
            el.set(qn("w:space"), "4")
            el.set(qn("w:color"), border_c)
            pBdr.append(el)
        pPr.append(pBdr)

        run_score = gonogo_box.add_run(f"\n{score}/100\n")
        run_score.bold = True
        run_score.font.size = Pt(28)
        run_score.font.color.rgb = txt_color
        run_label = gonogo_box.add_run(f"{rec}\n")
        run_label.bold = True
        run_label.font.size = Pt(16)
        run_label.font.color.rgb = txt_color

        if gonogo.get("summary"):
            _add_encadre(gonogo["summary"], bg_hex="EFF6FF", border_hex="2563EB")

        # KPI stats (4 cells — matching PDF stat-grid)
        elim_count = sum(1 for i in checklist_items if i.criticality and "liminatoire" in i.criticality.lower()) if checklist_items else 0
        imp_count = sum(1 for i in checklist_items if (i.criticality or "").lower() == "important") if checklist_items else 0
        risk_count = len(summary.get("risks", []))
        actions_count = len(summary.get("actions_next_48h", []))

        stat_table = doc.add_table(rows=1, cols=4)
        stat_table.style = "Table Grid"
        stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        stats = [
            (str(elim_count), "Éliminatoires"),
            (str(imp_count), "Importants"),
            (str(risk_count), "Risques identifiés"),
            (str(actions_count), "Actions à mener"),
        ]
        for i, (val, label) in enumerate(stats):
            cell = stat_table.rows[0].cells[i]
            _shade_cell(cell, "F8FAFC")
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_v = p.add_run(f"{val}\n")
            run_v.bold = True
            run_v.font.size = Pt(20)
            run_v.font.color.rgb = DARK_BLUE
            run_lb = p.add_run(label)
            run_lb.font.size = Pt(8)
            run_lb.font.color.rgb = GRAY

    # ── Top 3 risques critiques (matching PDF) ──
    risks = summary.get("risks", [])
    if risks:
        doc.add_heading("Top 3 risques critiques", 3)
        top3_table = _styled_table(["Risque", "Sévérité", "Impact"])
        for r in risks[:3]:
            row = top3_table.add_row().cells
            _set_cell_text(row[0], (r.get("risk") or "")[:50], bold=True, size=Pt(9))
            sev = r.get("severity") or ""
            sev_label = {"high": "FORT", "medium": "MOYEN", "low": "BAS"}.get(sev.lower(), sev.upper())
            _set_cell_text(row[1], sev_label, bold=True, color=_severity_color(sev), size=Pt(9))
            bg_r = _severity_bg(sev)
            if bg_r:
                _shade_cell(row[1], bg_r)
            _set_cell_text(row[2], (r.get("why") or "")[:120], size=Pt(8))

    # ── Forces / Points de vigilance (2-column table — matching PDF) ──
    if gonogo and (gonogo.get("strengths") or gonogo.get("risks")):
        sf_table = doc.add_table(rows=1, cols=2)
        sf_table.style = "Table Grid"
        sf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header: Forces | Points de vigilance
        _shade_cell(sf_table.rows[0].cells[0], "D1FAE5")
        _set_cell_text(sf_table.rows[0].cells[0], "Forces", bold=True,
                      color=GREEN_DARK, size=Pt(10))
        _shade_cell(sf_table.rows[0].cells[1], "FEF2F2")
        _set_cell_text(sf_table.rows[0].cells[1], "Points de vigilance", bold=True,
                      color=RED, size=Pt(10))
        strengths = (gonogo.get("strengths") or [])[:3]
        gng_risks = (gonogo.get("risks") or [])[:3]
        max_rows = max(len(strengths), len(gng_risks))
        for idx in range(max_rows):
            row = sf_table.add_row().cells
            if idx < len(strengths):
                _set_cell_text(row[0], f"+ {strengths[idx]}", size=Pt(9),
                              color=GREEN_DARK)
            if idx < len(gng_risks):
                _set_cell_text(row[1], f"- {gng_risks[idx]}", size=Pt(9),
                              color=RED)

    # ── Dimensions Go/No-Go breakdown (matching PDF) ──
    if gonogo and gonogo.get("breakdown"):
        doc.add_heading("Scores par dimension", 3)
        dim_table = _styled_table(["Dimension", "Score", "Niveau"])
        breakdown = gonogo["breakdown"]
        dim_mapping = [
            ("Adéquation technique", breakdown.get("technical_fit")),
            ("Capacité financière", breakdown.get("financial_capacity")),
            ("Faisabilité planning", breakdown.get("timeline_feasibility")),
            ("Position concurrentielle", breakdown.get("competitive_position")),
        ]
        for dim_name, dim_score in dim_mapping:
            if dim_score is not None:
                row = dim_table.add_row().cells
                _set_cell_text(row[0], dim_name, size=Pt(9))
                _set_cell_text(row[1], f"{dim_score}/100", bold=True, size=Pt(9))
                if dim_score >= 70:
                    level_label, level_color, level_bg = "BON", GREEN_DARK, "D1FAE5"
                elif dim_score >= 50:
                    level_label, level_color, level_bg = "MOYEN", ORANGE, "FEF3C7"
                else:
                    level_label, level_color, level_bg = "FAIBLE", RED, "FEF2F2"
                _set_cell_text(row[2], level_label, bold=True, color=level_color, size=Pt(9))
                _shade_cell(row[2], level_bg)

    # ── Actions prioritaires P0 (filtered — matching PDF) ──
    actions_p0 = [a for a in summary.get("actions_next_48h", []) if a.get("priority") == "P0"]
    if actions_p0:
        doc.add_heading("Actions prioritaires P0", 3)
        p0_table = _styled_table(["Action", "Responsable"])
        for a in actions_p0:
            row = p0_table.add_row().cells
            _set_cell_text(row[0], (a.get("action") or "")[:100], size=Pt(8))
            _set_cell_text(row[1], (a.get("owner_role") or "")[:30], size=Pt(8))

    # ── Confiance IA (matching PDF) ──
    confidence = summary.get("confidence") or summary.get("avg_confidence")
    if confidence:
        conf_p = doc.add_paragraph()
        run_cl = conf_p.add_run(f"Indice de confiance IA : {confidence * 100:.0f}%")
        run_cl.font.size = Pt(10)
        run_cl.font.color.rgb = GRAY
        # Visual confidence bar as a 1-row table
        bar_table = doc.add_table(rows=1, cols=1)
        bar_table.style = "Table Grid"
        cell = bar_table.rows[0].cells[0]
        if confidence >= 0.8:
            bar_color = "059669"
        elif confidence >= 0.6:
            bar_color = "D97706"
        else:
            bar_color = "DC2626"
        _shade_cell(cell, bar_color)
        pct_text = f"{confidence * 100:.0f}%"
        _set_cell_text(cell, pct_text, bold=True, color=WHITE, size=Pt(9))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set width proportionally
        cell.width = Cm(confidence * 17)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §2. FICHE SIGNALÉTIQUE DU MARCHÉ (RC)
    # ══════════════════════════════════════════════════════════════════════
    if rc:
        doc.add_heading("2. Fiche signalétique du marché", 1)
        p_sub_rc = doc.add_paragraph()
        run_sub_rc = p_sub_rc.add_run("Données extraites du Règlement de Consultation (RC) et des pièces contractuelles")
        run_sub_rc.font.color.rgb = GRAY
        run_sub_rc.font.size = Pt(9)
        rc_table = _styled_table(["Rubrique", "Détail"])
        rc_fields = [
            ("Procédure", rc.get("procedure_type")),
            ("Allotissement", rc.get("allotissement")),
            ("Groupement", rc.get("groupement")),
            ("Variantes", rc.get("variantes")),
            ("Sous-traitance", "Autorisée" if rc.get("subcontracting_allowed") else "Non autorisée"),
            ("CCAG de référence", rc.get("ccag_reference")),
            ("Visite obligatoire", "Oui" if rc.get("visite_obligatoire") else "Non"),
            ("DUME requis", "Oui" if rc.get("dume_required") else "Non"),
        ]
        for label, value in rc_fields:
            if value:
                row = rc_table.add_row().cells
                _set_cell_text(row[0], label, bold=True, size=Pt(9))
                _set_cell_text(row[1], str(value), size=Pt(9))

        lots = rc.get("lots", [])
        if lots:
            doc.add_heading("Lots", 2)
            lots_table = _styled_table(["N°", "Intitulé", "Montant estimé"])
            for lot in lots:
                row = lots_table.add_row().cells
                _set_cell_text(row[0], str(lot.get("number") or lot.get("lot_number") or "—"), size=Pt(9))
                _set_cell_text(row[1], lot.get("title") or lot.get("label") or "—", size=Pt(9))
                _set_cell_text(row[2], str(lot.get("estimated_amount") or "—"), size=Pt(9))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §3. RÉSUMÉ EXÉCUTIF
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Résumé exécutif", 1)

    # Summary box — structured like PDF (Objet/Acheteur/Lieu/Date/Budget)
    summary_fields = []
    if po.get("scope") or po.get("object"):
        summary_fields.append(("Objet", po.get("scope") or po.get("object")))
    if po.get("buyer"):
        summary_fields.append(("Acheteur", po["buyer"]))
    if po.get("location"):
        summary_fields.append(("Lieu", po["location"]))
    if po.get("deadline_submission"):
        summary_fields.append(("Date limite", _fmt_date(po["deadline_submission"])))
    if po.get("estimated_budget"):
        summary_fields.append(("Budget estimé", po["estimated_budget"]))
    if po.get("market_type"):
        summary_fields.append(("Type de marché", po["market_type"]))

    if summary_fields:
        box_text = "\n".join(f"{label} : {value}" for label, value in summary_fields)
        _add_encadre(box_text, bg_hex="EFF6FF", border_hex="2563EB")

    # Points clés
    key_points = summary.get("key_points", [])
    if key_points:
        doc.add_heading("Points clés du DCE", 2)
        kp_table = _styled_table(["Priorité", "Point clé"])
        for kp in key_points:
            row = kp_table.add_row().cells
            importance = (kp.get("importance") or "medium").lower()
            label = {"high": "HAUT", "medium": "MOYEN", "low": "BAS"}.get(importance, importance.upper())
            _set_cell_text(row[0], label, bold=True, color=_severity_color(importance), size=Pt(9))
            bg = _severity_bg(importance)
            if bg:
                _shade_cell(row[0], bg)
            _set_cell_text(row[1], kp.get("point", "") or kp.get("label", ""), size=Pt(9))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE II — CONFORMITÉ JURIDIQUE
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE II — CONFORMITÉ JURIDIQUE",
                     "Vérification administrative, conformité, dérogations CCAG")

    # ══════════════════════════════════════════════════════════════════════
    #  §4. VÉRIFICATION ADMINISTRATIVE DC
    # ══════════════════════════════════════════════════════════════════════
    if dc_check:
        doc.add_heading("4. Vérification administrative DC", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Vérification des pièces administratives — DC1, DC2, attestations, certifications")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        dc_docs = dc_check.get("documents", [])
        if dc_docs:
            dc_table = _styled_table(["Document", "Statut", "Date validité", "Alerte"])
            for d in dc_docs:
                row = dc_table.add_row().cells
                _set_cell_text(row[0], d.get("name") or d.get("document") or "—", bold=True, size=Pt(9))
                status = d.get("status") or "—"
                _set_cell_text(row[1], status, bold=True, color=_status_color(status), size=Pt(9))
                bg = "D1FAE5" if status == "OK" else ("FEF2F2" if status in ("MANQUANT", "KO") else None)
                if bg:
                    _shade_cell(row[1], bg)
                _set_cell_text(row[2], _fmt_date(d.get("expiry_date") or d.get("validity_date")), size=Pt(9))
                alert = d.get("alert") or ""
                _set_cell_text(row[3], alert[:80], color=RED if alert else GRAY, size=Pt(8))
                if status in ("MANQUANT", "KO"):
                    _shade_row(dc_table.rows[-1], "FEF2F2")

        if dc_check.get("summary") or dc_check.get("resume"):
            _add_encadre(dc_check.get("summary") or dc_check.get("resume"), bg_hex="EFF6FF", border_hex="2563EB")

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §5. CHECKLIST DE CONFORMITÉ
    # ══════════════════════════════════════════════════════════════════════
    if checklist_items:
        elim_count = sum(1 for i in checklist_items if i.criticality and "liminatoire" in i.criticality.lower())
        manq_count = sum(1 for i in checklist_items if i.status == "MANQUANT")
        ok_count = sum(1 for i in checklist_items if i.status == "OK")

        doc.add_heading(f"5. Checklist de conformité ({len(checklist_items)} exigences)", 1)

        # Stats summary box
        _add_encadre(
            f"{elim_count} éliminatoires  |  {manq_count} manquants  |  {ok_count} conformes",
            bg_hex="EFF6FF", border_hex="2563EB"
        )

        cl_table = _styled_table(["#", "Exigence", "Catégorie", "Criticité", "Statut", "À fournir"])
        for idx, item in enumerate(checklist_items, 1):
            row = cl_table.add_row().cells
            _set_cell_text(row[0], str(idx), size=Pt(8))
            _set_cell_text(row[1], item.requirement or "", size=Pt(8))
            _set_cell_text(row[2], item.category or "—", size=Pt(8))
            crit = item.criticality or "—"
            _set_cell_text(row[3], crit, bold=True, color=_severity_color(crit), size=Pt(8))
            status = item.status or "—"
            _set_cell_text(row[4], status, bold=True, color=_status_color(status), size=Pt(8))
            _set_cell_text(row[5], item.what_to_provide or "—", size=Pt(8))
            # Row background for eliminatoire manquant
            if crit and "liminatoire" in crit.lower() and status == "MANQUANT":
                _shade_row(cl_table.rows[-1], "FEF2F2")

        # ── Section 5b: Documents prioritaires éliminatoires MANQUANTS ──
        elim_manquants = [i for i in checklist_items
                          if i.criticality and "liminatoire" in i.criticality.lower()
                          and i.status == "MANQUANT"]
        if elim_manquants:
            doc.add_page_break()
            h_5b = doc.add_heading(f"5b. Documents prioritaires à préparer ({len(elim_manquants)} éliminatoires)", 2)
            for run in h_5b.runs:
                run.font.color.rgb = RED

            _add_encadre(
                f"Attention : Ces {len(elim_manquants)} documents sont éliminatoires. "
                "Leur absence entraînera le rejet automatique de votre candidature. "
                "Préparez-les en priorité absolue.",
                bg_hex="FEF2F2", border_hex="EF4444",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

            elim_table = _styled_table(["#", "Document / Justificatif requis", "Détail à fournir"])
            for idx_e, item in enumerate(elim_manquants, 1):
                row = elim_table.add_row().cells
                _set_cell_text(row[0], str(idx_e), size=Pt(8))
                _set_cell_text(row[1], (item.requirement or "")[:80], bold=True, size=Pt(8))
                _set_cell_text(row[2], (item.what_to_provide or "Voir exigences du RC/CCAP")[:100], size=Pt(8))
                _shade_row(elim_table.rows[-1], "FEF2F2")

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §6. DÉROGATIONS CCAG-TRAVAUX 2021
    # ══════════════════════════════════════════════════════════════════════
    if ccag_derogations:
        doc.add_heading("6. Dérogations CCAG-Travaux 2021", 1)
        _add_encadre(
            f"{len(ccag_derogations)} dérogation(s) au CCAG-Travaux 2021 détectée(s). "
            "Certaines peuvent avoir un impact significatif sur votre prix et vos risques contractuels.",
            bg_hex="FEF2F2", border_hex="EF4444",
            text_color=RGBColor(0x99, 0x1B, 0x1B)
        )

        d_table = _styled_table(["Article CCAG", "Dérogation CCAP", "Gravité", "Évaluation"])
        for d in ccag_derogations:
            row = d_table.add_row().cells
            _set_cell_text(row[0], d.get("article_ccag") or d.get("article") or "—",
                          bold=True, size=Pt(9))
            _set_cell_text(row[1], d.get("derogation", "—"), size=Pt(9))
            sev = d.get("severity") or d.get("impact") or "—"
            sev_label = {"high": "FORT", "medium": "MOYEN", "low": "BAS",
                        "fort": "FORT", "moyen": "MOYEN"}.get(sev.lower() if isinstance(sev, str) else sev, sev.upper() if isinstance(sev, str) else str(sev))
            _set_cell_text(row[2], sev_label, bold=True,
                          color=_severity_color(sev), size=Pt(9))
            bg = _severity_bg(sev)
            if bg:
                _shade_cell(row[2], bg)
            _set_cell_text(row[3], d.get("evaluation") or d.get("risk_comment") or "—", size=Pt(9))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §7. CLAUSES RISQUÉES CCAP
    # ══════════════════════════════════════════════════════════════════════
    clauses_risquees = ccap.get("clauses_risquees", [])
    if clauses_risquees:
        doc.add_heading("7. Clauses risquées CCAP", 1)
        cr_table = _styled_table(["Clause", "Risque", "Sévérité", "Recommandation"])
        for cl in clauses_risquees[:15]:
            row = cr_table.add_row().cells
            _set_cell_text(row[0], cl.get("clause") or cl.get("article") or "—", bold=True, size=Pt(9))
            _set_cell_text(row[1], cl.get("risk") or cl.get("risk_description") or "—", size=Pt(9))
            sev = cl.get("severity") or "—"
            _set_cell_text(row[2], sev.upper() if isinstance(sev, str) else str(sev),
                          bold=True, color=_severity_color(sev), size=Pt(9))
            _set_cell_text(row[3], cl.get("recommendation") or cl.get("mitigation") or "—", size=Pt(9))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE III — ANALYSE FINANCIÈRE
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE III — ANALYSE FINANCIÈRE",
                     "Acte d'engagement, synthèse financière, trésorerie")

    # ══════════════════════════════════════════════════════════════════════
    #  §8. ANALYSE ACTE D'ENGAGEMENT
    # ══════════════════════════════════════════════════════════════════════
    if ae:
        doc.add_heading("8. Analyse Acte d'Engagement", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Conditions contractuelles extraites de l'Acte d'Engagement")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        ae_fields = [
            ("Type de prix", ae.get("type_prix")),
            ("Forme du prix", ae.get("forme_prix")),
            ("Révision des prix", ae.get("revision_prix")),
            ("Indice de révision", ae.get("indice_revision")),
            ("Durée du marché", ae.get("duree_marche")),
            ("Reconduction", ae.get("reconduction")),
            ("Pénalités de retard", ae.get("penalites_retard")),
            ("Retenue de garantie", f"{ae.get('retenue_garantie_pct')}%" if ae.get("retenue_garantie_pct") else None),
            ("Avance", f"{ae.get('avance_pct')}%" if ae.get("avance_pct") else None),
            ("Délai de paiement", ae.get("delai_paiement")),
            ("Délai global d'exécution", ae.get("delai_global")),
        ]
        ae_table = _styled_table(["Rubrique", "Valeur"])
        for label, value in ae_fields:
            if value:
                row = ae_table.add_row().cells
                _set_cell_text(row[0], label, bold=True, size=Pt(9))
                _set_cell_text(row[1], str(value), size=Pt(9))

        ae_clauses = ae.get("clauses_risquees", [])
        if ae_clauses:
            doc.add_heading("Clauses risquées de l'AE", 2)
            acl_table = _styled_table(["Clause", "Risque", "Sévérité", "Conseil"])
            for cl in ae_clauses[:10]:
                row = acl_table.add_row().cells
                _set_cell_text(row[0], cl.get("clause") or cl.get("article") or "—", bold=True, size=Pt(9))
                _set_cell_text(row[1], cl.get("risk") or cl.get("description") or "—", size=Pt(9))
                sev = cl.get("severity") or "—"
                _set_cell_text(row[2], sev.upper() if isinstance(sev, str) else str(sev),
                              bold=True, color=_severity_color(sev), size=Pt(9))
                bg = _severity_bg(sev)
                if bg:
                    _shade_cell(row[2], bg)
                _set_cell_text(row[3], cl.get("conseil") or cl.get("recommendation") or "—", size=Pt(8))

        score_ae = ae.get("score_risque")
        if score_ae is not None:
            _add_encadre(
                f"Score de risque AE : {score_ae}/100",
                bg_hex="FEF2F2" if score_ae >= 60 else ("FEF3C7" if score_ae >= 30 else "D1FAE5"),
                border_hex="EF4444" if score_ae >= 60 else ("D97706" if score_ae >= 30 else "059669"),
                text_color=RGBColor(0x99, 0x1B, 0x1B) if score_ae >= 60 else (ORANGE if score_ae >= 30 else GREEN_DARK)
            )

        if ae.get("resume"):
            _add_encadre(ae["resume"], bg_hex="EFF6FF", border_hex="2563EB")

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §9. SYNTHÈSE FINANCIÈRE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Synthèse financière", 1)
    p_sub_fin = doc.add_paragraph()
    run_sub_fin = p_sub_fin.add_run("Éléments financiers clés extraits du DCE pour l'aide à la décision")
    run_sub_fin.font.color.rgb = GRAY
    run_sub_fin.font.size = Pt(9)

    _add_encadre(
        f"Budget global estimé : {po.get('estimated_budget') or 'Non précisé dans le DCE'}\n"
        f"Type de prix : {po.get('market_type') or 'Non précisé'}",
        bg_hex="ECFDF5", border_hex="10B981",
        text_color=GREEN_DARK
    )

    # Points clés financiers
    fin_points = [kp for kp in summary.get("key_points", [])
                  if any(w in (kp.get("point") or "").lower()
                         for w in ["prix", "avance", "retenue", "paiement", "penalite",
                                   "revision", "financ", "budget", "garantie", "caution"])]
    if fin_points:
        doc.add_heading("Éléments financiers extraits", 2)
        fin_table = _styled_table(["Point financier"])
        for kp in fin_points:
            row = fin_table.add_row().cells
            _set_cell_text(row[0], (kp.get("point") or "")[:200], size=Pt(9))

    # Risques financiers
    fin_risks = [r for r in summary.get("risks", [])
                 if any(w in (r.get("why") or r.get("risk") or "").lower()
                        for w in ["financ", "prix", "penalite", "cout", "tresorerie", "paiement"])]
    if fin_risks:
        doc.add_heading("Risques financiers identifiés", 2)
        fr_table = _styled_table(["Risque", "Sévérité", "Impact financier"])
        for r in fin_risks:
            row = fr_table.add_row().cells
            _set_cell_text(row[0], (r.get("risk") or "")[:50], bold=True, size=Pt(9))
            sev = r.get("severity") or ""
            _set_cell_text(row[1], sev.upper(), bold=True,
                          color=_severity_color(sev), size=Pt(9))
            _set_cell_text(row[2], (r.get("why") or "")[:150], size=Pt(8))

    # Recommandation warning box (matching PDF)
    _add_encadre(
        "Recommandation : Avant de chiffrer votre offre, vérifiez les éléments suivants : "
        "formule de révision des prix, montant de l'avance, conditions de paiement, pénalités de retard, "
        "retenue de garantie, et tout risque de surcoût identifié (pollution, aléas géotechniques, etc.).",
        bg_hex="FEF2F2", border_hex="EF4444",
        text_color=RGBColor(0x99, 0x1B, 0x1B)
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §10. BENCHMARK TARIFAIRE DPGF
    # ══════════════════════════════════════════════════════════════════════
    if dpgf_pricing:
        doc.add_heading("10. Benchmark tarifaire DPGF", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Comparaison des prix unitaires avec les références régionales du marché BTP")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        sous = sum(1 for l in dpgf_pricing if (l.get("status") or "") == "SOUS_EVALUE")
        sur = sum(1 for l in dpgf_pricing if (l.get("status") or "") == "SUR_EVALUE")
        normal = sum(1 for l in dpgf_pricing if (l.get("status") or "") == "NORMAL")

        stat_table = doc.add_table(rows=1, cols=4)
        stat_table.style = "Table Grid"
        stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (val, label, bg_color) in enumerate([
            (str(len(dpgf_pricing)), "Postes analysés", "EFF6FF"),
            (str(sous), "Sous-évalués", "FEF2F2"),
            (str(sur), "Sur-évalués", "FEF3C7"),
            (str(normal), "Normaux", "D1FAE5"),
        ]):
            cell = stat_table.rows[0].cells[i]
            _shade_cell(cell, bg_color)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_v = p.add_run(f"{val}\n")
            run_v.bold = True
            run_v.font.size = Pt(16)
            run_v.font.color.rgb = DARK_BLUE
            run_lb = p.add_run(label)
            run_lb.font.size = Pt(8)
            run_lb.font.color.rgb = GRAY

        doc.add_paragraph()

        dp_table = _styled_table(["Désignation", "Prix unit.", "Réf. min", "Réf. moy", "Réf. max", "Statut"])
        for line in dpgf_pricing[:20]:
            if isinstance(line, dict):
                row = dp_table.add_row().cells
                _set_cell_text(row[0], (line.get("designation") or "—")[:40], size=Pt(8))
                _set_cell_text(row[1], f"{line.get('prix_unitaire', '—')} €", size=Pt(8))
                ref = line.get("reference_match") or {}
                if isinstance(ref, dict):
                    _set_cell_text(row[2], f"{ref.get('prix_min', '—')} €", size=Pt(8))
                    _set_cell_text(row[3], f"{ref.get('prix_moyen', '—')} €", size=Pt(8))
                    _set_cell_text(row[4], f"{ref.get('prix_max', '—')} €", size=Pt(8))
                else:
                    for j in range(2, 5):
                        _set_cell_text(row[j], "—", size=Pt(8))
                status = line.get("status") or "INCONNU"
                status_colors = {
                    "SOUS_EVALUE": (RED, "FEF2F2"),
                    "SUR_EVALUE": (ORANGE, "FEF3C7"),
                    "NORMAL": (GREEN_DARK, "D1FAE5"),
                }
                s_color, s_bg = status_colors.get(status, (GRAY, None))
                _set_cell_text(row[5], status.replace("_", " "), bold=True, color=s_color, size=Pt(8))
                if s_bg:
                    _shade_cell(row[5], s_bg)

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §11. SIMULATION TRÉSORERIE ET BFR
    # ══════════════════════════════════════════════════════════════════════
    if cashflow:
        doc.add_heading("11. Simulation trésorerie et BFR", 1)
        p_sub_cf = doc.add_paragraph()
        run_sub_cf = p_sub_cf.add_run("Impact prévisionnel sur la trésorerie de l'entreprise — aide à la décision financière")
        run_sub_cf.font.color.rgb = GRAY
        run_sub_cf.font.size = Pt(9)

        # BFR peak financial box — support both key formats (bfr_peak or bfr_eur)
        bfr_peak = cashflow.get("bfr_peak") or cashflow.get("bfr_eur")
        if bfr_peak:
            bfr_val = float(bfr_peak)
            bfr_text = f"BFR maximal estimé : {abs(bfr_val):,.0f} €".replace(",", " ")
            bfr_month = cashflow.get("bfr_month")
            if bfr_month:
                bfr_text += f" — atteint au mois {bfr_month}"
            margin = cashflow.get("margin_estimate") or cashflow.get("marge_brute_pct")
            if margin:
                bfr_text += f"\nMarge estimée : {float(margin):.1f}%"
            montant = cashflow.get("montant_total_ht")
            if montant:
                bfr_text += f"\nMontant total HT : {float(montant):,.0f} €".replace(",", " ")
            _add_encadre(bfr_text, bg_hex="ECFDF5", border_hex="10B981", text_color=GREEN_DARK)

        # Résumé cashflow
        if cashflow.get("resume"):
            _add_encadre(cashflow["resume"], bg_hex="EFF6FF", border_hex="2563EB")

        # Monthly cashflow table — support both key formats
        monthly = cashflow.get("monthly_cashflow", [])
        if monthly:
            doc.add_heading("Trésorerie mensuelle prévisionnelle", 3)
            mcf_table = _styled_table(["Mois", "Dépenses HT", "Encaissements HT", "Solde cumulé"])
            for m in monthly[:15]:
                row = mcf_table.add_row().cells
                month_num = m.get("month") or m.get("mois", "?")
                month_label = m.get("label") or f"M{month_num}"
                _set_cell_text(row[0], month_label, size=Pt(8))
                expenses = m.get("expenses") or m.get("depenses_ht") or 0
                _set_cell_text(row[1], f"-{float(expenses):,.0f} €".replace(",", " "),
                              color=RED, size=Pt(8))
                income = m.get("income") or m.get("encaissement_ht") or 0
                _set_cell_text(row[2], f"+{float(income):,.0f} €".replace(",", " "),
                              color=GREEN_DARK, size=Pt(8))
                cumulative = m.get("cumulative") or m.get("solde_cumule") or 0
                cum_val = float(cumulative)
                cum_color = RED if cum_val < 0 else GREEN_DARK
                _set_cell_text(row[3], f"{cum_val:,.0f} €".replace(",", " "),
                              bold=True, color=cum_color, size=Pt(8))
                if cum_val < 0:
                    _shade_cell(row[3], "FEF2F2")

        # Risk level
        risk_level = cashflow.get("risk_level")
        if risk_level:
            _add_encadre(
                f"Niveau de risque trésorerie : {risk_level}",
                bg_hex="FEF2F2" if "LEV" in (risk_level or "").upper() else "FEF3C7",
                border_hex="EF4444" if "LEV" in (risk_level or "").upper() else "D97706",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

        # Tension months
        tension = cashflow.get("tension_months", [])
        if tension:
            _add_encadre(
                f"Mois en tension ({len(tension)}) : M" + ", M".join(str(t) for t in tension) +
                "\nPrévoir un financement de trésorerie ou ligne de crédit.",
                bg_hex="FEF2F2", border_hex="EF4444",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

        warnings = cashflow.get("warnings", [])
        if warnings:
            doc.add_heading("Alertes trésorerie", 2)
            for w in warnings[:5]:
                w_text = w.get("message") if isinstance(w, dict) else str(w)
                _add_encadre(w_text, bg_hex="FEF2F2", border_hex="EF4444",
                            text_color=RGBColor(0x99, 0x1B, 0x1B))

        # Hypothèses box
        avance_eur = cashflow.get("avance_impact_eur")
        retenue_eur = cashflow.get("retenue_impact_eur")
        duree = cashflow.get("duree_mois")
        hyp_parts = ["Hypothèses : Délai de paiement acheteur 30j"]
        if avance_eur:
            hyp_parts.append(f"avance forfaitaire {float(avance_eur):,.0f} €".replace(",", " "))
        if retenue_eur:
            hyp_parts.append(f"retenue de garantie {float(retenue_eur):,.0f} €".replace(",", " "))
        if duree:
            hyp_parts.append(f"durée {duree} mois")
        hyp_parts.append("Ajustez selon les termes du CCAP.")
        _add_encadre(
            ", ".join(hyp_parts),
            bg_hex="EFF6FF", border_hex="3B82F6",
            text_color=RGBColor(0x1E, 0x3A, 0x8A)
        )

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE IV — ANALYSE TECHNIQUE
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE IV — ANALYSE TECHNIQUE",
                     "CCTP, contradictions inter-documents, sous-traitance")

    # ══════════════════════════════════════════════════════════════════════
    #  §12. ANALYSE TECHNIQUE CCTP
    # ══════════════════════════════════════════════════════════════════════
    if cctp:
        doc.add_heading("12. Analyse technique CCTP", 1)
        p_sub_cctp = doc.add_paragraph()
        run_sub_cctp = p_sub_cctp.add_run("Synthèse des exigences techniques extraites du Cahier des Clauses Techniques Particulières")
        run_sub_cctp.font.color.rgb = GRAY
        run_sub_cctp.font.size = Pt(9)

        if cctp.get("technical_summary"):
            _add_encadre(f"Synthèse technique : {cctp['technical_summary']}", bg_hex="EFF6FF", border_hex="2563EB")

        categories = cctp.get("categories", [])
        if categories:
            for cat in categories:
                cat_name = cat.get("name") or cat.get("category") or "—"
                # Category heading with risk level badge (matching PDF)
                h_cat = doc.add_heading(cat_name, 3)
                risk_level = cat.get("risk_level")
                if risk_level:
                    run_rl = h_cat.add_run(f"  [{risk_level.upper()}]")
                    run_rl.font.size = Pt(8)
                    run_rl.font.color.rgb = _severity_color(risk_level)

                # Items as table with Norme/DTU and Risque columns (matching PDF)
                items = cat.get("items") or cat.get("requirements") or []
                if items:
                    cat_table = _styled_table(["Exigence", "Détail", "Norme/DTU", "Risque"])
                    for item in items[:8]:
                        if isinstance(item, dict):
                            row = cat_table.add_row().cells
                            _set_cell_text(row[0], (item.get("requirement") or item.get("label") or "")[:60],
                                          bold=True, size=Pt(8))
                            _set_cell_text(row[1], (item.get("detail") or item.get("value") or item.get("description") or "")[:100],
                                          size=Pt(8))
                            _set_cell_text(row[2], (item.get("norm") or item.get("standard") or "")[:30],
                                          size=Pt(8))
                            item_risk = item.get("risk", "")
                            risk_label = (item_risk.upper() if item_risk else "INFO")
                            _set_cell_text(row[3], risk_label, bold=True,
                                          color=_severity_color(item_risk), size=Pt(8))
                        else:
                            row = cat_table.add_row().cells
                            _set_cell_text(row[0], str(item)[:60], size=Pt(8))

        # Contradictions as 3-column table (matching PDF: Contradiction / Références / Recommandation)
        contradictions = cctp.get("contradictions", [])
        if contradictions:
            h_contra = doc.add_heading("Contradictions détectées dans le CCTP", 2)
            for run in h_contra.runs:
                run.font.color.rgb = RED

            _add_encadre(
                f"Attention : {len(contradictions)} contradiction(s) interne(s) détectée(s) dans le CCTP. "
                "Posez la question à l'acheteur avant la date limite.",
                bg_hex="FEF2F2", border_hex="EF4444",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

            contra_table = _styled_table(["Contradiction", "Références", "Recommandation"])
            for c in contradictions[:5]:
                if isinstance(c, dict):
                    row = contra_table.add_row().cells
                    _set_cell_text(row[0], (c.get("description") or c.get("issue") or c.get("contradiction") or "")[:120],
                                  size=Pt(8))
                    _set_cell_text(row[1], (c.get("references") or c.get("source") or "")[:60], size=Pt(8))
                    _set_cell_text(row[2], (c.get("recommendation") or "Demander clarification")[:80], size=Pt(8))
                    _shade_row(contra_table.rows[-1], "FEF2F2")
                else:
                    row = contra_table.add_row().cells
                    _set_cell_text(row[0], str(c)[:120], size=Pt(8), color=RED)

        env_req = cctp.get("environmental_requirements", [])
        if env_req:
            doc.add_heading("Exigences environnementales", 2)
            _add_encadre(
                "\n".join(f"• {(req.get('requirement') if isinstance(req, dict) else str(req))}"
                          for req in env_req[:5]),
                bg_hex="EFF6FF", border_hex="3B82F6",
                text_color=RGBColor(0x1E, 0x3A, 0x8A)
            )

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §13. CONTRADICTIONS INTER-DOCUMENTS
    # ══════════════════════════════════════════════════════════════════════
    if conflicts_data:
        conflicts_list = conflicts_data.get("conflicts", [])
        nb_total = conflicts_data.get("nb_total", len(conflicts_list))
        nb_critical = conflicts_data.get("nb_critical", 0)

        doc.add_heading("13. Contradictions inter-documents", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Incohérences détectées entre les pièces du DCE (RC, CCAP, CCTP, DPGF, AE)")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        _add_encadre(
            f"{nb_total} contradiction(s) détectée(s) dont {nb_critical} critique(s). "
            "Posez la question à l'acheteur via la plateforme de dématérialisation.",
            bg_hex="FEF2F2", border_hex="EF4444",
            text_color=RGBColor(0x99, 0x1B, 0x1B)
        )

        if conflicts_list:
            cf_table = _styled_table(["Type", "Doc A ↔ Doc B", "Description", "Sévérité", "Recommandation"])
            for c in conflicts_list[:15]:
                if isinstance(c, dict):
                    row = cf_table.add_row().cells
                    _set_cell_text(row[0], (c.get("type") or "—")[:20], bold=True, size=Pt(8))
                    docs = f"{c.get('doc_a') or '?'} ↔ {c.get('doc_b') or '?'}"
                    _set_cell_text(row[1], docs[:25], size=Pt(8))
                    _set_cell_text(row[2], (c.get("description") or "")[:120], size=Pt(8))
                    sev = c.get("severity") or "—"
                    _set_cell_text(row[3], sev.upper() if isinstance(sev, str) else str(sev),
                                  bold=True, color=_severity_color(sev), size=Pt(8))
                    bg = _severity_bg(sev)
                    if bg:
                        _shade_cell(row[3], bg)
                    _set_cell_text(row[4], (c.get("recommendation") or c.get("recommandation") or "Demander clarification")[:80],
                                  size=Pt(8))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §14. ANALYSE SOUS-TRAITANCE
    # ══════════════════════════════════════════════════════════════════════
    if subcontracting:
        doc.add_heading("14. Analyse sous-traitance", 1)
        p_sub_st = doc.add_paragraph()
        run_sub_st = p_sub_st.add_run("Obligations et risques liés à la sous-traitance — conformité loi du 31/12/1975")
        run_sub_st.font.color.rgb = GRAY
        run_sub_st.font.size = Pt(9)

        if subcontracting.get("summary"):
            _add_encadre(subcontracting["summary"], bg_hex="EFF6FF", border_hex="2563EB")

        # Conditions as structured table with risk level (matching PDF)
        conditions = subcontracting.get("conditions", [])
        if conditions:
            doc.add_heading("Conditions contractuelles", 2)
            cond_table = _styled_table(["Condition", "Détail", "Risque"])
            for c in conditions:
                if isinstance(c, dict):
                    row = cond_table.add_row().cells
                    _set_cell_text(row[0], (c.get("label") or c.get("condition") or "")[:50],
                                  bold=True, size=Pt(8))
                    _set_cell_text(row[1], (c.get("detail") or c.get("value") or "")[:120], size=Pt(8))
                    c_risk = c.get("risk", "")
                    risk_label = (c_risk.upper() if c_risk else "INFO")
                    _set_cell_text(row[2], risk_label, bold=True,
                                  color=_severity_color(c_risk), size=Pt(8))
                else:
                    row = cond_table.add_row().cells
                    _set_cell_text(row[0], str(c)[:50], size=Pt(8))

        # Risks as warning boxes with mitigation (matching PDF)
        sub_risks = subcontracting.get("risks", [])
        if sub_risks:
            doc.add_heading("Risques identifiés", 2)
            for r in sub_risks[:5]:
                if isinstance(r, dict):
                    risk_text = f"{r.get('risk') or r.get('title') or ''} : {r.get('detail') or r.get('description') or ''}"
                    mitigation = r.get("mitigation")
                    if mitigation:
                        risk_text += f"\nMitigation : {mitigation}"
                    _add_encadre(risk_text, bg_hex="FEF2F2", border_hex="EF4444",
                                text_color=RGBColor(0x99, 0x1B, 0x1B))
                else:
                    _add_encadre(str(r), bg_hex="FEF2F2", border_hex="EF4444",
                                text_color=RGBColor(0x99, 0x1B, 0x1B))

        # Legal obligations (matching PDF)
        legal_obligations = subcontracting.get("legal_obligations", [])
        if legal_obligations:
            doc.add_heading("Obligations légales", 2)
            obligations_text = "\n".join(f"• {o}" for o in legal_obligations[:6])
            _add_encadre(obligations_text, bg_hex="EFF6FF", border_hex="3B82F6",
                        text_color=RGBColor(0x1E, 0x3A, 0x8A))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE V — SCORING ET STRATÉGIE
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE V — SCORING ET STRATÉGIE",
                     "Critères d'attribution, simulation note, Go/No-Go détaillé")

    # ══════════════════════════════════════════════════════════════════════
    #  §15. CRITÈRES D'ATTRIBUTION
    # ══════════════════════════════════════════════════════════════════════
    if criteria:
        doc.add_heading("15. Critères d'attribution", 1)
        ev = criteria.get("evaluation", {})

        elig = ev.get("eligibility_conditions", [])
        if elig:
            doc.add_heading("Conditions d'éligibilité", 2)
            e_table = _styled_table(["Condition", "Type"])
            for c in elig:
                row = e_table.add_row().cells
                _set_cell_text(row[0], c.get("condition", ""), size=Pt(9))
                ctype = (c.get("type") or "").upper()
                _set_cell_text(row[1], ctype, bold=True,
                              color=RED if ctype == "ELIMINATOIRE" else GRAY, size=Pt(9))

        scoring_crit = ev.get("scoring_criteria", [])
        if scoring_crit:
            doc.add_heading("Critères de notation", 2)
            sc_table = _styled_table(["Critère", "Pondération", "Notes"])
            for c in scoring_crit:
                row = sc_table.add_row().cells
                _set_cell_text(row[0], c.get("criterion", ""), size=Pt(9))
                w = c.get("weight_percent")
                _set_cell_text(row[1], f"{w}%" if w is not None else "N/S",
                              bold=True, color=ACCENT_BLUE, size=Pt(9))
                _set_cell_text(row[2], c.get("notes") or "—", size=Pt(9))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §16. SIMULATION NOTE ACHETEUR
    # ══════════════════════════════════════════════════════════════════════
    if scoring:
        doc.add_heading("16. Simulation note acheteur", 1)
        total = scoring.get("total_score", 0)
        max_s = scoring.get("max_score", 20)

        score_p = doc.add_paragraph()
        score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sc = score_p.add_run(f"{total:.1f} / {max_s}")
        run_sc.bold = True
        run_sc.font.size = Pt(24)
        run_sc.font.color.rgb = DARK_BLUE

        if scoring.get("rank_estimate"):
            rank_p = doc.add_paragraph()
            rank_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_rk = rank_p.add_run(f"Classement estimé : {scoring['rank_estimate']}")
            run_rk.font.size = Pt(11)
            run_rk.font.color.rgb = GRAY

        crit_scores = scoring.get("criteria_scores", [])
        if crit_scores:
            doc.add_heading("Détail par critère", 2)
            cs_table = _styled_table(["Critère", "Note", "Max", "Commentaire"])
            for cs in crit_scores:
                row = cs_table.add_row().cells
                _set_cell_text(row[0], cs.get("criterion", ""), size=Pt(9))
                _set_cell_text(row[1], str(cs.get("score", 0)), bold=True,
                              color=ACCENT_BLUE, size=Pt(9))
                _set_cell_text(row[2], str(cs.get("max_score", 20)), size=Pt(9))
                _set_cell_text(row[3], cs.get("comment") or cs.get("justification") or "—", size=Pt(9))

        levers = scoring.get("improvement_levers", [])
        if levers:
            doc.add_heading("Leviers d'amélioration", 2)
            _add_encadre(
                "Actions concrètes pour améliorer votre note :",
                bg_hex="ECFDF5", border_hex="10B981",
                text_color=GREEN_DARK
            )
            for lev in levers:
                p = doc.add_paragraph(style="List Bullet")
                title_text = lev.get("lever") or lev.get("title", "")
                impact = lev.get("potential_gain") or lev.get("impact", "")
                run_t = p.add_run(title_text)
                run_t.bold = True
                if impact:
                    run_i = p.add_run(f" — {impact}")
                    run_i.font.color.rgb = GREEN_DARK

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §17. GO/NO-GO DÉTAILLÉ — 9 DIMENSIONS
    # ══════════════════════════════════════════════════════════════════════
    if gonogo and (gonogo.get("dimension_scores") or gonogo.get("breakdown")):
        doc.add_heading("17. Go/No-Go détaillé — 9 dimensions", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Analyse multidimensionnelle de l'opportunité — scores par axe d'évaluation")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        score = gonogo.get("score", 0)
        rec = (gonogo.get("recommendation") or "ATTENTION").upper()

        # Score box
        if rec == "GO":
            bg, border_c = "D1FAE5", "059669"
            txt_color = GREEN_DARK
        elif rec == "NO-GO":
            bg, border_c = "FDE8E8", "DC2626"
            txt_color = RED
        else:
            bg, border_c = "FEF3C7", "D97706"
            txt_color = ORANGE
        _add_encadre(f"Score global : {score}/100 — Recommandation : {rec}", bg_hex=bg, border_hex=border_c, text_color=txt_color)

        # Dimension scores table
        dims = gonogo.get("dimension_scores", [])
        if dims:
            doc.add_heading("Détail par dimension", 2)
            dim_table = _styled_table(["Dimension", "Score", "Poids", "Confiance", "Explication"])
            for d in dims:
                row = dim_table.add_row().cells
                _set_cell_text(row[0], d.get("name") or d.get("dimension") or "—", bold=True, size=Pt(8))
                d_score = d.get("score", 0)
                _set_cell_text(row[1], f"{d_score}/100", bold=True, size=Pt(8))
                if d_score >= 70:
                    _shade_cell(row[1], "D1FAE5")
                elif d_score < 50:
                    _shade_cell(row[1], "FEF2F2")
                _set_cell_text(row[2], f"{d.get('weight', '—')}%", size=Pt(8))
                _set_cell_text(row[3], f"{d.get('confidence', 0)*100:.0f}%" if isinstance(d.get('confidence'), (int, float)) else "—", size=Pt(8))
                _set_cell_text(row[4], (d.get("explanation") or d.get("justification") or "")[:100], size=Pt(8))
        elif gonogo.get("breakdown"):
            # Fallback to old 4-dimension format
            doc.add_heading("Scores par dimension", 2)
            breakdown = gonogo["breakdown"]
            dim_table = _styled_table(["Dimension", "Score", "Niveau"])
            for dim_name, dim_key in [
                ("Adéquation technique", "technical_fit"),
                ("Capacité financière", "financial_capacity"),
                ("Faisabilité planning", "timeline_feasibility"),
                ("Position concurrentielle", "competitive_position"),
                ("Certifications", "certifications_match"),
                ("Zone géographique", "geographic_fit"),
                ("Assurances", "insurance_coverage"),
                ("Sous-traitance", "subcontracting_risk"),
                ("Taux de succès", "success_rate"),
            ]:
                dim_score = breakdown.get(dim_key)
                if dim_score is not None:
                    row = dim_table.add_row().cells
                    _set_cell_text(row[0], dim_name, size=Pt(9))
                    _set_cell_text(row[1], f"{dim_score}/100", bold=True, size=Pt(9))
                    if dim_score >= 70:
                        level_label, level_color, level_bg = "BON", GREEN_DARK, "D1FAE5"
                    elif dim_score >= 50:
                        level_label, level_color, level_bg = "MOYEN", ORANGE, "FEF3C7"
                    else:
                        level_label, level_color, level_bg = "FAIBLE", RED, "FEF2F2"
                    _set_cell_text(row[2], level_label, bold=True, color=level_color, size=Pt(9))
                    _shade_cell(row[2], level_bg)

        # Profile gaps
        gaps = gonogo.get("profile_gaps", [])
        if gaps:
            doc.add_heading("Écarts profil entreprise", 2)
            _add_encadre(
                "Points à renforcer :\n" + "\n".join(f"• {g}" for g in gaps[:8]),
                bg_hex="FEF2F2", border_hex="EF4444",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

        # Profile strengths
        strengths_list = gonogo.get("profile_strengths", [])
        if strengths_list:
            doc.add_heading("Points forts du profil", 2)
            _add_encadre(
                "Atouts de votre candidature :\n" + "\n".join(f"• {s}" for s in strengths_list[:8]),
                bg_hex="D1FAE5", border_hex="059669",
                text_color=GREEN_DARK
            )

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE VI — PLAN D'ACTION
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE VI — PLAN D'ACTION",
                     "Risques, questions, calendrier")

    # ══════════════════════════════════════════════════════════════════════
    #  §18. ANALYSE DES RISQUES
    # ══════════════════════════════════════════════════════════════════════
    risks = summary.get("risks", [])
    if risks:
        doc.add_heading("18. Analyse des risques", 1)
        risk_table = _styled_table(["Risque", "Sévérité", "Explication"])
        for r in risks:
            row = risk_table.add_row().cells
            _set_cell_text(row[0], r.get("risk", ""), size=Pt(9))
            sev = r.get("severity") or ""
            sev_label = {"high": "FORT", "medium": "MOYEN", "low": "BAS"}.get(sev.lower(), sev.upper())
            _set_cell_text(row[1], sev_label, bold=True, color=_severity_color(sev), size=Pt(9))
            bg = _severity_bg(sev)
            if bg:
                _shade_cell(row[1], bg)
            _set_cell_text(row[2], r.get("why", ""), size=Pt(9))

    # Actions 48h
    actions = summary.get("actions_next_48h", [])
    if actions:
        doc.add_heading("Actions prioritaires sous 48h", 2)
        _add_encadre(
            "Ces actions sont à engager immédiatement pour sécuriser la réponse à l'appel d'offres.",
            bg_hex="FEF2F2", border_hex="EF4444",
            text_color=RGBColor(0x99, 0x1B, 0x1B)
        )
        act_table = _styled_table(["Action", "Responsable", "Priorité"])
        for a in actions:
            row = act_table.add_row().cells
            _set_cell_text(row[0], a.get("action", ""), size=Pt(9))
            _set_cell_text(row[1], a.get("owner_role", ""), size=Pt(9))
            prio = a.get("priority", "")
            _set_cell_text(row[2], prio.upper() if prio else "—", bold=True,
                          color=_severity_color(prio), size=Pt(9))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §19. QUESTIONS PRIORITAIRES POUR L'ACHETEUR
    # ══════════════════════════════════════════════════════════════════════
    if questions_list:
        doc.add_heading("19. Questions prioritaires pour l'acheteur", 1)
        _add_encadre(
            "Conseil : Posez ces questions via la plateforme de dématérialisation "
            "avant la date limite de questions. Les réponses seront diffusées à tous les candidats.",
            bg_hex="EFF6FF", border_hex="3B82F6",
            text_color=RGBColor(0x1E, 0x3A, 0x8A)
        )
        q_table = _styled_table(["#", "Priorité", "Question", "Justification"])
        for idx, q in enumerate(questions_list[:15], 1):
            row = q_table.add_row().cells
            _set_cell_text(row[0], str(idx), size=Pt(9))
            prio = (q.get("priority") or "—")
            _set_cell_text(row[1], prio.upper(), bold=True,
                          color=_severity_color(prio), size=Pt(9))
            _set_cell_text(row[2], (q.get("question") or str(q))[:140], size=Pt(8))
            _set_cell_text(row[3], (q.get("justification") or "")[:100], size=Pt(8))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §20. CALENDRIER ET DATES CLÉS
    # ══════════════════════════════════════════════════════════════════════
    if timeline:
        doc.add_heading("20. Calendrier et dates clés", 1)
        cal_table = _styled_table(["Échéance", "Date / Durée"])

        if timeline.get("submission_deadline"):
            row = cal_table.add_row().cells
            _shade_row(cal_table.rows[-1], "FEF2F2")
            cell0 = row[0]
            cell0.text = ""
            p = cell0.paragraphs[0]
            run_t = p.add_run("Date limite de remise ")
            run_t.bold = True
            run_t.font.size = Pt(9)
            run_imp = p.add_run("(impératif)")
            run_imp.font.color.rgb = RED
            run_imp.font.size = Pt(7)
            run_imp.bold = True
            _set_cell_text(row[1], _fmt_date(timeline["submission_deadline"]),
                          bold=True, color=RED, size=Pt(9))

        if timeline.get("execution_duration_months"):
            row = cal_table.add_row().cells
            _set_cell_text(row[0], "Durée d'exécution", size=Pt(9))
            _set_cell_text(row[1], f"{timeline['execution_duration_months']} mois", size=Pt(9))

        for kd in timeline.get("key_dates", []):
            label = kd.get("label", "")
            if "remise des offres" in label.lower():
                continue
            row = cal_table.add_row().cells
            cell0 = row[0]
            cell0.text = ""
            p = cell0.paragraphs[0]
            run_t = p.add_run(label[:55])
            run_t.font.size = Pt(9)
            if kd.get("mandatory"):
                run_m = p.add_run(" (oblig.)")
                run_m.font.color.rgb = RED
                run_m.font.size = Pt(7)
            _set_cell_text(row[1], _fmt_date(kd.get("date")), size=Pt(9))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  ANNEXES
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("ANNEXES",
                     "Documents analysés, glossaire, mentions légales")

    # ══════════════════════════════════════════════════════════════════════
    #  A1. DOCUMENTS ANALYSÉS
    # ══════════════════════════════════════════════════════════════════════
    if documents:
        doc.add_heading("A1. Inventaire des documents analysés", 1)
        p_sub_doc = doc.add_paragraph()
        run_sub_doc = p_sub_doc.add_run("Liste des pièces du DCE prises en compte dans cette analyse")
        run_sub_doc.font.color.rgb = GRAY
        run_sub_doc.font.size = Pt(9)

        doc_table = _styled_table(["#", "Document", "Type", "Pages", "Taille", "Qualité OCR"])
        for idx_d, d in enumerate(documents, 1):
            row = doc_table.add_row().cells
            _set_cell_text(row[0], str(idx_d), size=Pt(8))
            _set_cell_text(row[1], (d.original_name or "—")[:45], size=Pt(8))
            dtype = d.doc_type or "AUTRES"
            _set_cell_text(row[2], dtype, bold=True,
                          color=ACCENT_BLUE if dtype in ("RC", "CCTP", "CCAP") else GRAY,
                          size=Pt(8))
            _set_cell_text(row[3], str(d.page_count or "—"), size=Pt(8))
            # Size display
            file_size = getattr(d, "file_size", None)
            if file_size:
                if file_size >= 1_000_000:
                    size_disp = f"{file_size / 1_000_000:.1f} Mo"
                elif file_size >= 1_000:
                    size_disp = f"{file_size / 1_000:.0f} Ko"
                else:
                    size_disp = f"{file_size} o"
            else:
                size_disp = "—"
            _set_cell_text(row[4], size_disp, size=Pt(8))
            # OCR quality
            ocr_q = getattr(d, "ocr_confidence", None)
            if ocr_q is not None:
                if ocr_q >= 70:
                    ocr_color, ocr_bg = GREEN_DARK, "D1FAE5"
                elif ocr_q >= 40:
                    ocr_color, ocr_bg = ORANGE, "FEF3C7"
                else:
                    ocr_color, ocr_bg = RED, "FEF2F2"
                _set_cell_text(row[5], f"{ocr_q:.0f}%", bold=True, color=ocr_color, size=Pt(8))
                _shade_cell(row[5], ocr_bg)
            else:
                _set_cell_text(row[5], "—", size=Pt(8))

        # Total summary line
        total_pages = sum(d.page_count or 0 for d in documents)
        total_p = doc.add_paragraph()
        run_total = total_p.add_run(
            f"Total : {len(documents)} document{'s' if len(documents) > 1 else ''}"
            f"{f' — {total_pages} pages analysées' if total_pages else ''}"
        )
        run_total.font.size = Pt(9)
        run_total.font.color.rgb = GRAY

    # ══════════════════════════════════════════════════════════════════════
    #  A2. GLOSSAIRE BTP
    # ══════════════════════════════════════════════════════════════════════
    if glossaire_btp:
        doc.add_page_break()
        doc.add_heading("A2. Glossaire BTP", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Termes clés utilisés dans ce rapport — Référentiel marchés publics BTP")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        gl_table = _styled_table(["Terme", "Définition"])
        for term, definition in glossaire_btp:
            row = gl_table.add_row().cells
            _set_cell_text(row[0], term, bold=True, color=ACCENT_BLUE, size=Pt(9))
            _set_cell_text(row[1], str(definition)[:200], size=Pt(8))

    # ══════════════════════════════════════════════════════════════════════
    #  A3. AVERTISSEMENT IA ET MENTIONS LÉGALES
    # ══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("A3. Avertissement IA et mentions légales", 1)

    # Disclaimer box (yellow background like PDF)
    disclaimer_text = (
        "Avertissement : Ce rapport est généré par intelligence artificielle "
        "(Claude, Anthropic) à partir des documents du DCE fournis. "
        "Il constitue une aide à la décision et ne se substitue pas à "
        "l'analyse humaine d'un expert marchés publics. "
        "Les informations extraites doivent être systématiquement vérifiées "
        "avant toute soumission d'offre."
    )
    # Add confidence score like PDF
    conf_val = summary.get("confidence") or summary.get("avg_confidence")
    if conf_val:
        disclaimer_text += f"\n\nConfiance globale de l'analyse : {conf_val * 100:.0f}%"
    _add_encadre(
        disclaimer_text,
        bg_hex="FFFBEB", border_hex="FBBF24",
        text_color=RGBColor(0x78, 0x35, 0x0F)
    )

    doc.add_paragraph()

    _add_encadre(
        "Données utilisées : Seuls les documents uploadés dans le projet "
        "ont été analysés. L'IA ne dispose d'aucune information externe "
        "(pas d'accès internet, pas de base de données externe). "
        "La qualité de l'analyse dépend directement de la qualité "
        "des documents fournis (lisibilité OCR, exhaustivité du DCE).",
        bg_hex="EFF6FF", border_hex="3B82F6",
        text_color=RGBColor(0x1E, 0x3A, 0x8A)
    )

    # ══════════ FOOTER ══════════
    doc.add_paragraph()
    doc.add_paragraph()
    footer_brand = doc.add_paragraph()
    footer_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fb = footer_brand.add_run("AO COPILOT")
    run_fb.bold = True
    run_fb.font.color.rgb = ACCENT_BLUE
    run_fb.font.size = Pt(11)

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fp = footer_para.add_run(
        f"aocopilot.fr — {datetime.now().strftime('%d/%m/%Y %H:%M')} — "
        "Rapport confidentiel — Reproduction interdite"
    )
    run_fp.font.color.rgb = GRAY_LIGHT
    run_fp.font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_memo_technique(db: Session, project_id: str) -> bytes:
    """Génère une mémoire technique Word (.docx) complète — 10 sections standard BTP.
    Design identique au rapport PDF/DOCX. Système hybride :
    - Texte bleu = pré-rempli par IA (proposition intelligente basée sur le DCE)
    - Texte orange = à compléter par le client (données entreprise spécifiques)
    Conforme aux attentes acheteurs publics (marche-public.fr, Doaken 2026).
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from io import BytesIO
    except ImportError as e:
        raise RuntimeError("python-docx non installé. Ajoutez python-docx==1.1.0.") from e

    from app.models.company_profile import CompanyProfile
    from app.models.organization import Organization

    pid = uuid.UUID(project_id)
    project = db.query(AoProject).filter_by(id=pid).first()
    if not project:
        raise ValueError("Projet introuvable")

    # ── Charger toutes les données ────────────────────────────────────────
    def _get(result_type):
        r = db.query(ExtractionResult).filter_by(
            project_id=pid, result_type=result_type
        ).order_by(ExtractionResult.version.desc()).first()
        return r.payload if r and r.payload else {}

    summary = _get("summary")
    criteria = _get("criteria")
    timeline = _get("timeline")
    cctp = _get("cctp")
    ccap = _get("ccap")
    scoring = _get("scoring")
    rc = _get("rc")
    subcontracting = _get("subcontracting")
    gonogo = _get("gonogo")

    po = summary.get("project_overview", {})
    key_points = summary.get("key_points", [])
    risks = summary.get("risks", [])
    evaluation = criteria.get("evaluation", {})
    scoring_criteria = evaluation.get("scoring_criteria", [])

    # ── Génération des narratifs LLM (graceful degradation) ──────────────
    _memo_intro_text: str | None = None
    _memo_positioning_text: str | None = None
    _memo_action_plan_text: str | None = None
    try:
        from app.services.llm import llm_service
        from app.services.prompts import (
            build_memo_intro_prompt,
            build_memo_positioning_prompt,
            build_memo_action_plan_prompt,
        )
        _company_dict = {}
        if company:
            _company_dict = {
                "name": org_name,
                "activity_sector": getattr(company, "activity_sector", None),
                "annual_revenue_eur": getattr(company, "revenue_eur", None),
                "certifications": getattr(company, "certifications", None),
                "regions": getattr(company, "regions", None),
                "staff_count": getattr(company, "employee_count", None),
                "main_clients": getattr(company, "main_clients", None),
                "references_btp": getattr(company, "references_btp", None),
            }
        _gonogo_score = gonogo.get("score") or 0
        _gonogo_dims = gonogo.get("dimension_scores") or gonogo.get("breakdown") or {}
        _deadline_str = (
            project.submission_deadline.strftime("%d/%m/%Y")
            if project.submission_deadline else "—"
        )

        sys_p, usr_p = build_memo_intro_prompt(
            project_title=project.title or "AO",
            buyer=project.buyer or "Acheteur public",
            scope=po.get("scope") or po.get("object") or "Prestations à définir",
            go_nogo_score=_gonogo_score,
            top_risks=risks[:3],
            company_profile=_company_dict,
        )
        _memo_intro_text = llm_service.chat_text(sys_p, usr_p, max_tokens=400)

        sys_p, usr_p = build_memo_positioning_prompt(
            company_profile=_company_dict,
            gonogo_dimensions=_gonogo_dims if isinstance(_gonogo_dims, dict) else {},
            eligibility_gaps=gonogo.get("profile_gaps") or [],
        )
        _memo_positioning_text = llm_service.chat_text(sys_p, usr_p, max_tokens=400)

        _actions = summary.get("actions_next_48h") or []
        sys_p, usr_p = build_memo_action_plan_prompt(
            actions_48h=_actions,
            risks=risks,
            deadline_submission=_deadline_str,
        )
        _memo_action_plan_text = llm_service.chat_text(sys_p, usr_p, max_tokens=500)
    except Exception as _llm_err:
        logger.warning("memo_llm_generation_skipped", error=str(_llm_err))

    # ── Génération des graphiques (graceful degradation) ──────────────────
    _radar_buf = None
    _cashflow_buf = None
    try:
        from app.services.chart_generator import generate_gonogo_radar, generate_cashflow_chart
        if gonogo:
            _dims = gonogo.get("dimension_scores") or gonogo.get("breakdown") or {}
            _radar_buf = generate_gonogo_radar(
                _dims if isinstance(_dims, dict) else {},
                gonogo.get("score"),
                project.title or "AO",
            )
        cashflow_raw = _get("cashflow")
        if cashflow_raw:
            _cashflow_buf = generate_cashflow_chart(cashflow_raw, project.title or "AO")
    except Exception as _chart_err:
        logger.warning("memo_chart_generation_skipped", error=str(_chart_err))

    checklist_items = db.query(ChecklistItem).filter_by(
        project_id=pid
    ).order_by(ChecklistItem.criticality, ChecklistItem.category).all()

    documents = db.query(AoDocument).filter_by(project_id=pid).all()

    # Charger le profil entreprise
    org = db.query(Organization).filter_by(id=project.org_id).first()
    company = None
    if org:
        company = db.query(CompanyProfile).filter_by(org_id=org.id).first()

    org_name = (org.name if org else None) or "Notre Entreprise"

    doc = Document()

    # ── Couleurs (identiques au rapport DOCX) ────────────────────────────
    DARK_BLUE = RGBColor(0x0F, 0x1B, 0x4C)
    ACCENT_BLUE = RGBColor(0x25, 0x63, 0xEB)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    RED = RGBColor(0xB9, 0x1C, 0x1C)
    GREEN_DARK = RGBColor(0x04, 0x78, 0x57)
    ORANGE = RGBColor(0xD9, 0x77, 0x06)
    GRAY = RGBColor(0x64, 0x74, 0x8B)
    GRAY_LIGHT = RGBColor(0x9C, 0xA3, 0xAF)
    IA_BLUE = RGBColor(0x1E, 0x40, 0xAF)  # Texte proposition IA

    # ── Style global ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)  # 11pt — BTP audit: lisibilité comités (50+ ans)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = DARK_BLUE
        hs.font.name = "Calibri"

    # ── Helpers (identiques au rapport DOCX) ─────────────────────────────
    def _shade_cell(cell, hex_color):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _shade_row(row, hex_color):
        for cell in row.cells:
            _shade_cell(cell, hex_color)

    def _set_cell_text(cell, text, bold=False, color=None, size=None):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = bold
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = size
        run.font.name = "Calibri"
        return run

    def _styled_table(headers):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            _shade_cell(cell, "0F1B4C")
            _set_cell_text(cell, h, bold=True, color=WHITE, size=Pt(9))
        return t

    def _fmt_date(val):
        if not val:
            return "—"
        s = str(val)
        if "T" in s:
            s = s.split("T")[0]
        try:
            from datetime import datetime as dt
            d = dt.strptime(s, "%Y-%m-%d")
            return d.strftime("%d/%m/%Y")
        except Exception:
            return s

    def _add_encadre(text, bg_hex="EFF6FF", border_hex="2563EB", text_color=None):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_hex)
        pPr.append(shd)
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "36")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), border_hex)
        pBdr.append(left)
        pPr.append(pBdr)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "180")
        ind.set(qn("w:right"), "180")
        pPr.append(ind)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        if text_color:
            run.font.color.rgb = text_color
        return p

    def _ia_text(text):
        """Paragraphe pré-rempli par l'IA (bleu, fond bleu clair)."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "EFF6FF")
        pPr.append(shd)
        run = p.add_run(text)
        run.font.color.rgb = IA_BLUE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        return p

    def _client_text(text):
        """Paragraphe à compléter par le client (orange, fond orange clair)."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "FEF3C7")
        pPr.append(shd)
        run = p.add_run(text)
        run.font.color.rgb = ORANGE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        return p

    def _severity_color(sev):
        s = (sev or "").lower()
        if s in ("high", "fort", "haut", "eliminatoire"):
            return RED
        elif s in ("medium", "moyen"):
            return ORANGE
        elif s in ("low", "bas"):
            return GREEN_DARK
        return GRAY

    def _status_color(status):
        s = (status or "").upper()
        if s == "OK":
            return GREEN_DARK
        elif s in ("MANQUANT", "KO"):
            return RED
        elif s in ("PARTIEL",):
            return ORANGE
        return GRAY

    # ── En-tête & Pied de page ─────────────────────────────────────────
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = ""
    run_h = header_para.add_run(f"{org_name} — Mémoire Technique — {project.title[:50]}")
    run_h.font.size = Pt(8)
    run_h.font.color.rgb = GRAY
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Page ")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = GRAY
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run_xml = footer_para.add_run()
    run_xml.font.size = Pt(8)
    run_xml._r.append(fldChar_begin)
    run_xml._r.append(instrText)
    run_xml._r.append(fldChar_end)

    # ══════════════════════════════════════════════════════════════════════
    #  PAGE DE GARDE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_b = brand.add_run(org_name.upper())
    run_b.font.color.rgb = DARK_BLUE
    run_b.font.size = Pt(14)
    run_b.bold = True

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = line.add_run("━━━━━━━━━━")
    run_l.font.color.rgb = ACCENT_BLUE

    title_heading = doc.add_heading("MÉMOIRE TECHNIQUE", 0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_heading.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sub.add_run(project.title)
    run_s.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    run_s.font.size = Pt(14)

    doc.add_paragraph()

    # Info table
    cover_info = _styled_table(["Information", "Détail"])
    deadline_str = project.submission_deadline.strftime("%d/%m/%Y") if project.submission_deadline else "—"
    for label, value in [
        ("Marché", project.title),
        ("Acheteur public", project.buyer or "—"),
        ("Référence", project.reference or "—"),
        ("Candidat", org_name),
        ("Lot(s)", ", ".join(l.get("title", "") for l in rc.get("lots", [])) or "Lot unique"),
        ("Date de remise", deadline_str),
        ("Document établi le", datetime.now().strftime("%d/%m/%Y")),
    ]:
        row = cover_info.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=Pt(9))
        _set_cell_text(row[1], str(value), size=Pt(9))

    doc.add_paragraph()

    # Legend box
    _add_encadre(
        "Légende des couleurs dans ce document :\n"
        "Texte bleu sur fond bleu clair = Proposition IA (pré-rempli à partir du DCE — à valider)\n"
        "Texte orange sur fond jaune clair = À compléter par le candidat (données entreprise)",
        bg_hex="F8FAFC", border_hex="2563EB"
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SOMMAIRE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Sommaire", 1)

    sommaire = [
        ("1.", "Présentation de l'entreprise — Identité, certifications, spécialités"),
        ("2.", "Compréhension du projet — Analyse du DCE et enjeux identifiés"),
        ("3.", "Méthodologie d'exécution — Préparation, exécution, réception"),
        ("4.", "Moyens humains — Organigramme, qualifications, effectifs"),
        ("5.", "Moyens matériels et techniques — Équipements et outils"),
        ("6.", "Planning prévisionnel — Jalons DCE et phasage d'exécution"),
        ("7.", "Gestion de la qualité (SOPAQ) — Contrôles, traçabilité, DOE"),
        ("8.", "Sécurité et prévention — PPSPS, EPI, formations"),
        ("9.", "Environnement, RSE et gestion des déchets — SOGED, insertion"),
        ("10.", "Références et expériences similaires"),
        ("A.", "Checklist de conformité DCE"),
        ("B.", "Analyse des risques — Proposition IA + mesures client"),
    ]
    for num, label in sommaire:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "dotted")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CBD5E1")
        pBdr.append(bottom)
        pPr.append(pBdr)
        run_num = p.add_run(f"{num}  ")
        run_num.bold = True
        run_num.font.color.rgb = ACCENT_BLUE
        run_num.font.size = Pt(10)
        run_text = p.add_run(label)
        run_text.font.size = Pt(10)
        run_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Critères de notation rappel + Matrice de conformité (audit BTP: mapper chaque section aux critères RC)
    if scoring_criteria:
        doc.add_paragraph()
        _add_encadre(
            "Rappel des critères de notation (RC) :\n" +
            "\n".join(
                f"- {c.get('criterion', '')}: {c.get('weight_percent', '?')}%"
                for c in scoring_criteria
            ) +
            "\n\nCe document est structuré pour maximiser votre note technique.",
            bg_hex="ECFDF5", border_hex="10B981",
            text_color=GREEN_DARK
        )

        # Conformity mapping table — maps each section to RC criteria
        doc.add_paragraph()
        doc.add_heading("Matrice de conformité — Sections ↔ Critères RC", 3)
        # Build mapping: each section addresses certain criteria keywords
        _criteria_section_map = [
            ("1. Présentation entreprise", ["capacité", "moyen", "compétence", "expérience", "qualif"]),
            ("2. Compréhension du projet", ["technique", "méthodologie", "compréhension"]),
            ("3. Méthodologie d'exécution", ["technique", "méthodologie", "qualité"]),
            ("4. Moyens humains", ["moyen", "personnel", "humain", "compétence"]),
            ("5. Moyens matériels", ["moyen", "matériel", "technique"]),
            ("6. Planning", ["délai", "planning", "organisation"]),
            ("7. Qualité (SOPAQ)", ["qualité", "contrôle", "gestion"]),
            ("8. Sécurité", ["sécurité", "prévention", "santé"]),
            ("9. RSE / Environnement", ["environnement", "développement durable", "rse", "social", "insertion"]),
            ("10. Références", ["référence", "expérience", "capacité"]),
        ]
        conf_table = _styled_table(["Section Mémoire", "Critère(s) RC adressé(s)", "Pondération"])
        for sec_name, keywords in _criteria_section_map:
            matched = [c for c in scoring_criteria
                       if any(kw in (c.get("criterion") or "").lower() for kw in keywords)]
            if matched:
                row = conf_table.add_row().cells
                _set_cell_text(row[0], sec_name, bold=True, size=Pt(8))
                criteria_text = ", ".join(c.get("criterion", "") for c in matched)
                _set_cell_text(row[1], criteria_text[:80], size=Pt(8))
                total_weight = sum(c.get("weight_percent", 0) or 0 for c in matched)
                _set_cell_text(row[2], f"{total_weight}%" if total_weight else "—",
                              bold=True, color=ACCENT_BLUE, size=Pt(9))

    doc.add_page_break()

    # ── Radar chart Go/No-Go (si disponible) ──────────────────────────────
    if _radar_buf:
        try:
            _radar_buf.seek(0)
            doc.add_heading("Synthèse Go/No-Go — 9 dimensions", 2)
            doc.add_picture(_radar_buf, width=Inches(5.5))
            # doc.add_picture() adds the picture to the last paragraph
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        except Exception as _e:
            logger.warning("memo_radar_insert_failed", error=str(_e))

    # ══════════════════════════════════════════════════════════════════════
    #  1. PRÉSENTATION DE L'ENTREPRISE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Présentation de l'entreprise", 1)

    # LLM positionnement stratégique
    if _memo_positioning_text:
        _ia_text(_memo_positioning_text)
        doc.add_paragraph()

    doc.add_heading("1.1 Identité", 2)
    id_table = _styled_table(["Rubrique", "Proposition IA", "À compléter / vérifier"])
    id_fields = [
        ("Raison sociale", org_name, "Vérifier"),
        ("Chiffre d'affaires", f"{company.revenue_eur:,} €".replace(",", " ") if company and company.revenue_eur else "—", "CA dernier exercice clos"),
        ("Effectif", f"{company.employee_count} salariés" if company and company.employee_count else "—", "Effectif total CDI"),
        ("Forme juridique", "—", "SAS, SARL, SA, etc."),
        ("SIRET", "—", "N° SIRET à 14 chiffres"),
        ("Code APE / NAF", "—", "Code activité"),
        ("Adresse siège", "—", "Adresse complète"),
        ("Date de création", "—", "Année de création"),
    ]
    for label, ia_val, client_val in id_fields:
        row = id_table.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=Pt(9))
        if ia_val != "—":
            _set_cell_text(row[1], ia_val, color=IA_BLUE, size=Pt(9))
            _shade_cell(row[1], "EFF6FF")
        else:
            _set_cell_text(row[1], "—", color=GRAY, size=Pt(9))
        _set_cell_text(row[2], client_val, color=ORANGE, size=Pt(9))
        _shade_cell(row[2], "FEF3C7")

    # Certifications
    doc.add_heading("1.2 Certifications et qualifications", 2)
    if company and company.certifications:
        _ia_text("Certifications identifiées dans votre profil (à compléter avec dates de validité) :")
        for cert in company.certifications:
            p = doc.add_paragraph(style="List Bullet")
            run_c = p.add_run(str(cert))
            run_c.bold = True
            run_c.font.color.rgb = IA_BLUE
    else:
        _client_text("Listez vos certifications : Qualibat (n° et classification), ISO 9001, ISO 14001, MASE, RGE, etc. Indiquez les dates de validité.")

    # Vérifier exigences RC
    elim_certs = [item for item in checklist_items
                  if item.category and "certif" in item.category.lower()
                  and item.criticality and "liminatoire" in item.criticality.lower()]
    if elim_certs:
        _add_encadre(
            "ATTENTION — Certifications éliminatoires exigées par le RC :\n" +
            "\n".join(f"- {e.requirement}" for e in elim_certs),
            bg_hex="FEF2F2", border_hex="EF4444",
            text_color=RGBColor(0x99, 0x1B, 0x1B)
        )

    # Spécialités
    doc.add_heading("1.3 Domaines de spécialité", 2)
    if company and company.specialties:
        for spec in company.specialties:
            doc.add_paragraph(str(spec).replace("_", " ").title(), style="List Bullet")
    else:
        _client_text("Listez vos domaines : gros-oeuvre, électricité, plomberie, CVC, couverture, etc.")

    # Zones d'intervention
    doc.add_heading("1.4 Zones d'intervention", 2)
    if company and company.regions:
        doc.add_paragraph(", ".join(str(r) for r in company.regions))
    location = po.get("location") or ""
    if location:
        _ia_text(f"Le marché se situe à : {location}. Confirmez que cette zone est couverte.")

    # Assurances
    doc.add_heading("1.5 Assurances", 2)
    assur_table = _styled_table(["Type d'assurance", "Proposition IA", "À compléter"])
    assur_fields = [
        ("Assurance décennale", "Souscrite" if company and company.assurance_decennale else "—", "N° police + assureur + date validité"),
        ("RC Professionnelle", f"{company.assurance_rc_montant:,} €".replace(",", " ") if company and company.assurance_rc_montant else "—", "Montant couverture + assureur"),
        ("RC Exploitation", "—", "Montant + assureur"),
        ("Tous Risques Chantier", "—", "Si souscrite"),
    ]
    for label, ia_val, client_val in assur_fields:
        row = assur_table.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=Pt(9))
        if ia_val != "—":
            _set_cell_text(row[1], ia_val, color=IA_BLUE, size=Pt(9))
            _shade_cell(row[1], "EFF6FF")
        else:
            _set_cell_text(row[1], "—", color=GRAY, size=Pt(9))
        _set_cell_text(row[2], client_val, color=ORANGE, size=Pt(9))
        _shade_cell(row[2], "FEF3C7")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  2. COMPRÉHENSION DU PROJET
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Compréhension du projet", 1)

    scope_text = po.get("scope") or po.get("object") or "Le présent marché porte sur des prestations à définir."
    market_type = po.get("market_type") or "non précisé"

    # Utiliser le narratif LLM si disponible, sinon fallback template
    if _memo_intro_text:
        _ia_text(_memo_intro_text)
    else:
        _ia_text(
            f"Dans le cadre de la consultation lancée par {project.buyer or 'l acheteur public'}, "
            f"portant sur un marché de {market_type}, notre société {org_name} "
            f"a analysé en détail les documents du Dossier de Consultation des Entreprises "
            f"afin de formuler une réponse parfaitement adaptée aux besoins exprimés."
        )

    doc.add_heading("2.1 Objet et périmètre", 2)
    _ia_text(scope_text)
    if location:
        _ia_text(f"Lieu d'exécution : {location}")
    if po.get("estimated_budget"):
        _ia_text(f"Enveloppe budgétaire estimée : {po['estimated_budget']}")
    duration = timeline.get("execution_duration_months")
    if duration:
        _ia_text(f"Durée d'exécution : {duration} mois")

    # Points clés (IA)
    if key_points:
        doc.add_heading("2.2 Points clés identifiés dans le DCE", 2)
        _add_encadre(
            "Les points suivants ont été automatiquement extraits du DCE par l'IA. "
            "Reprenez-les dans votre rédaction pour montrer votre compréhension du projet.",
            bg_hex="EFF6FF", border_hex="2563EB"
        )
        for kp in key_points:
            point = kp.get("point") or kp.get("label") or ""
            if point:
                p = doc.add_paragraph(style="List Bullet")
                run_kp = p.add_run(str(point))
                run_kp.font.color.rgb = IA_BLUE

    # Enjeux spécifiques
    doc.add_heading("2.3 Enjeux spécifiques et contraintes", 2)
    _client_text(
        "Décrivez ici votre compréhension personnelle des enjeux du projet. "
        "Mentionnez : la visite de site (si effectuée), les contraintes d'accès, "
        "la coactivité avec d'autres lots, les interfaces techniques, "
        "les exigences particulières du maître d'ouvrage."
    )

    # Contraintes CCTP
    if cctp:
        if cctp.get("environmental_requirements"):
            doc.add_heading("2.4 Contraintes environnementales identifiées", 2)
            for req in cctp["environmental_requirements"][:5]:
                text = req.get("requirement") if isinstance(req, dict) else str(req)
                p = doc.add_paragraph(style="List Bullet")
                run_r = p.add_run(text)
                run_r.font.color.rgb = IA_BLUE

        if cctp.get("contradictions"):
            _add_encadre(
                "ATTENTION — Contradictions détectées dans le CCTP :\n" +
                "\n".join(
                    f"- {(c.get('description') or c.get('contradiction') or str(c))}"
                    for c in cctp["contradictions"][:3]
                    if isinstance(c, dict)
                ) +
                "\nPosez une question à l'acheteur pour lever ces ambiguïtés.",
                bg_hex="FEF2F2", border_hex="EF4444",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  3. MÉTHODOLOGIE D'EXÉCUTION
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Méthodologie d'exécution", 1)
    _ia_text(
        "Notre approche repose sur une méthodologie rigoureuse et éprouvée, "
        "structurée en phases distinctes pour garantir la maîtrise des délais, "
        "de la qualité et des coûts. Chaque étape intègre les exigences "
        "spécifiques identifiées dans le CCTP."
    )

    doc.add_heading("3.1 Préparation de chantier", 2)
    prep_items = [
        "Analyse approfondie du CCTP, du RC et des pièces contractuelles",
        "Visite de site et relevé des contraintes in situ",
        "Validation des interfaces avec les autres intervenants et lots",
        "Élaboration du PPSPS et du plan d'installation de chantier",
        "Commande anticipée des matériaux à délais longs",
        "Mise en place de la base vie et des accès sécurisés",
    ]
    for item in prep_items:
        p = doc.add_paragraph(style="List Bullet")
        run_i = p.add_run(item)
        run_i.font.color.rgb = IA_BLUE
    _client_text("Complétez avec vos spécificités : méthodes constructives retenues, phasage logistique, plan d'installation de chantier.")

    doc.add_heading("3.2 Phase d'exécution", 2)
    exec_items = [
        "Suivi quotidien de l'avancement par le conducteur de travaux dédié",
        "Réunions de chantier hebdomadaires avec compte-rendu formalisé sous 48h",
        "Autocontrôles systématiques à chaque point d'arrêt défini dans le SOPAQ",
        "Gestion documentaire centralisée (plans d'exécution, DOE, fiches techniques)",
        "Coordination avec le CSPS et les autres entreprises",
        "Reporting mensuel d'avancement au maître d'ouvrage",
    ]
    for item in exec_items:
        p = doc.add_paragraph(style="List Bullet")
        run_i = p.add_run(item)
        run_i.font.color.rgb = IA_BLUE
    _client_text("Détaillez vos process internes, outils de suivi (logiciel de gestion), méthodes de reporting, gestion des approvisionnements.")

    doc.add_heading("3.3 Réception et levée de réserves", 2)
    for item in [
        "Pré-réception interne systématique avant les OPR",
        "Levée des réserves dans les délais contractuels (30 jours max)",
        "Remise du DOE complet et des plans de récolement sous 30 jours",
        "Garantie de parfait achèvement : intervention sous 48h en cas de sinistre",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run_i = p.add_run(item)
        run_i.font.color.rgb = IA_BLUE

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  4. MOYENS HUMAINS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Moyens humains", 1)
    if company and company.employee_count:
        _ia_text(
            f"Notre entreprise compte {company.employee_count} collaborateurs. "
            "L'équipe dédiée au présent marché est composée de professionnels qualifiés "
            "et expérimentés, sous la responsabilité d'un interlocuteur unique."
        )
    else:
        _ia_text(
            "L'équipe dédiée au présent marché est composée de professionnels qualifiés "
            "et expérimentés, sous la responsabilité d'un interlocuteur unique."
        )

    doc.add_heading("4.1 Organigramme de l'équipe projet", 2)
    _add_encadre(
        "Le tableau ci-dessous doit être renseigné avec les personnes nommément "
        "affectées au marché. Les CV peuvent être joints en annexe si le RC l'exige.",
        bg_hex="FEF3C7", border_hex="D97706",
        text_color=ORANGE
    )
    team_table = _styled_table(["Fonction", "Proposition IA", "Nom (à compléter)", "Qualification", "Expérience"])
    team_roles = [
        ("Directeur de travaux", "Pilotage global, validation technique"),
        ("Conducteur de travaux", "Suivi quotidien, coordination lots"),
        ("Chef de chantier", "Encadrement terrain, qualité"),
        ("Chef d'équipe", "Exécution, management compagnons"),
        ("Responsable QSE", "Qualité, sécurité, environnement"),
    ]
    for role, ia_desc in team_roles:
        row = team_table.add_row().cells
        _set_cell_text(row[0], role, bold=True, size=Pt(9))
        _set_cell_text(row[1], ia_desc, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[1], "EFF6FF")
        _set_cell_text(row[2], "Nom, prénom", color=ORANGE, size=Pt(8))
        _shade_cell(row[2], "FEF3C7")
        _set_cell_text(row[3], "Diplôme / habilitation", color=ORANGE, size=Pt(8))
        _shade_cell(row[3], "FEF3C7")
        _set_cell_text(row[4], "X ans", color=ORANGE, size=Pt(8))
        _shade_cell(row[4], "FEF3C7")

    doc.add_heading("4.2 Effectif mobilisé par phase", 2)
    eff_table = _styled_table(["Phase", "Ouvriers", "Techniciens", "Encadrement"])
    for phase in ["Préparation", "Gros-oeuvre", "Second oeuvre", "Finitions / OPR"]:
        row = eff_table.add_row().cells
        _set_cell_text(row[0], phase, bold=True, size=Pt(9))
        for i in range(1, 4):
            _set_cell_text(row[i], "À préciser", color=ORANGE, size=Pt(8))
            _shade_cell(row[i], "FEF3C7")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  5. MOYENS MATÉRIELS ET TECHNIQUES
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Moyens matériels et techniques", 1)

    doc.add_heading("5.1 Matériels affectés au chantier", 2)
    mat_table = _styled_table(["Type", "Désignation (modèle)", "Qté", "Propriété / Location"])
    for mat_type, ia_sugg in [
        ("Gros matériel", "Grue, nacelle, mini-pelle — à préciser"),
        ("Véhicules", "Fourgons, camions — à préciser"),
        ("Outillage spécifique", "Selon nature des travaux"),
        ("Matériel de mesure", "Niveau laser, théodolite, etc."),
        ("Matériel de sécurité", "Garde-corps, filets, balisage"),
    ]:
        row = mat_table.add_row().cells
        _set_cell_text(row[0], mat_type, bold=True, size=Pt(9))
        _set_cell_text(row[1], ia_sugg, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[1], "EFF6FF")
        _set_cell_text(row[2], "?", color=ORANGE, size=Pt(8))
        _shade_cell(row[2], "FEF3C7")
        _set_cell_text(row[3], "Propre / Location", color=ORANGE, size=Pt(8))
        _shade_cell(row[3], "FEF3C7")

    doc.add_heading("5.2 Outils numériques", 2)
    _ia_text("Outils recommandés pour valoriser votre offre :")
    for tool in [
        "Logiciel BIM (Revit, ArchiCAD) pour la maquette numérique",
        "Plateforme de suivi chantier (BatiScript, Fieldwire, Finalcad)",
        "GMAO pour la maintenance préventive des équipements",
        "Outils de reporting photo/vidéo horodatés",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run_t = p.add_run(tool)
        run_t.font.color.rgb = IA_BLUE
    _client_text("Listez vos outils réellement utilisés. Ne mentionnez que ce que vous maîtrisez.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  6. PLANNING PRÉVISIONNEL
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Planning prévisionnel", 1)

    deadline_str = project.submission_deadline.strftime("%d/%m/%Y") if project.submission_deadline else "à définir"
    duration = timeline.get("execution_duration_months")
    _ia_text(
        f"Date limite de remise : {deadline_str}. "
        + (f"Durée d'exécution prévue au marché : {duration} mois. " if duration else "")
        + "Le planning ci-dessous intègre les jalons contractuels identifiés dans le DCE."
    )

    # Jalons DCE
    key_dates = timeline.get("key_dates", [])
    if key_dates or timeline.get("submission_deadline"):
        doc.add_heading("6.1 Jalons contractuels du DCE", 2)
        jal_table = _styled_table(["Échéance", "Date", "Caractère"])
        if timeline.get("submission_deadline"):
            row = jal_table.add_row().cells
            _set_cell_text(row[0], "Date limite de remise", bold=True, size=Pt(9))
            _set_cell_text(row[1], _fmt_date(timeline["submission_deadline"]),
                          bold=True, color=RED, size=Pt(9))
            _set_cell_text(row[2], "IMPÉRATIF", bold=True, color=RED, size=Pt(9))
            _shade_row(jal_table.rows[-1], "FEF2F2")
        for kd in key_dates:
            if "remise des offres" in (kd.get("label") or "").lower():
                continue
            row = jal_table.add_row().cells
            _set_cell_text(row[0], kd.get("label", "—"), size=Pt(9))
            _set_cell_text(row[1], _fmt_date(kd.get("date")), color=IA_BLUE, size=Pt(9))
            mandatory = "OBLIGATOIRE" if kd.get("mandatory") else "Indicatif"
            _set_cell_text(row[2], mandatory, bold=kd.get("mandatory", False),
                          color=RED if kd.get("mandatory") else GRAY, size=Pt(9))

    doc.add_heading("6.2 Phasage prévisionnel d'exécution", 2)
    phase_table = _styled_table(["Phase", "Description", "Durée (Proposition IA)", "Durée réelle"])
    phases = [
        ("1 — Installation", "Base vie, clôture, accès, panneau", "2 semaines"),
        ("2 — Terrassement", "Terrassements, fondations", "À estimer"),
        ("3 — Gros-oeuvre", "Élévations, planchers, charpente", "À estimer"),
        ("4 — Clos & couvert", "Couverture, étanchéité, menuiseries ext.", "À estimer"),
        ("5 — Second oeuvre", "Cloisons, revêtements, peinture", "À estimer"),
        ("6 — Lots techniques", "CVC, plomberie, électricité, essais", "À estimer"),
        ("7 — VRD / Ext.", "Voirie, réseaux, espaces verts", "À estimer"),
        ("8 — OPR / Réception", "Pré-réception, levée réserves, DOE", "4 semaines"),
    ]
    for phase, desc, ia_duree in phases:
        row = phase_table.add_row().cells
        _set_cell_text(row[0], phase, bold=True, size=Pt(9))
        _set_cell_text(row[1], desc, size=Pt(8))
        _set_cell_text(row[2], ia_duree, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[2], "EFF6FF")
        _set_cell_text(row[3], "À compléter", color=ORANGE, size=Pt(8))
        _shade_cell(row[3], "FEF3C7")

    _client_text("Adaptez les phases et durées au marché spécifique. Joignez un diagramme de Gantt en annexe si possible. Un planning réaliste (intégrant aléas météo et interfaces) sera mieux noté qu'un planning optimiste.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  7. GESTION DE LA QUALITÉ (SOPAQ)
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Gestion de la qualité (SOPAQ)", 1)
    _ia_text(
        "Notre Schéma Organisationnel du Plan d'Assurance Qualité (SOPAQ) garantit "
        "la conformité des prestations aux spécifications du CCTP et aux normes en vigueur. "
        "Il sera décliné en PAQ définitif en phase préparation."
    )

    doc.add_heading("7.1 Points d'arrêt et contrôles", 2)
    qa_table = _styled_table(["Point de contrôle", "Fréquence", "Responsable", "Document"])
    qa_items = [
        ("Réception matériaux", "Chaque livraison", "Chef de chantier", "Fiche de réception"),
        ("Contrôle béton (éprouvettes)", "Chaque coulage", "Labo agréé", "PV d'essai"),
        ("Contrôle soudures", "100% soudures", "Bureau de contrôle", "PV de contrôle"),
        ("Étanchéité", "Par zone", "Responsable QSE", "Fiche autocontrôle"),
        ("Conformité plans d'exécution", "Avant chaque phase", "Conducteur travaux", "Visa MOE"),
        ("Essais réseaux (air, eau)", "Avant réception", "Bureau de contrôle", "PV d'essai"),
    ]
    for controle, freq, resp, doc_type in qa_items:
        row = qa_table.add_row().cells
        _set_cell_text(row[0], controle, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[0], "EFF6FF")
        _set_cell_text(row[1], freq, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[1], "EFF6FF")
        _set_cell_text(row[2], resp, color=ORANGE, size=Pt(8))
        _shade_cell(row[2], "FEF3C7")
        _set_cell_text(row[3], doc_type, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[3], "EFF6FF")

    doc.add_heading("7.2 Gestion des non-conformités", 2)
    for item in [
        "Détection → Fiche de non-conformité → Analyse cause → Action corrective → Vérification",
        "Traçabilité complète des matériaux (fiches techniques, PV d'essais, certificats)",
        "Audits internes trimestriels par le responsable QSE",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.color.rgb = IA_BLUE

    doc.add_heading("7.3 DOE (Dossier des Ouvrages Exécutés)", 2)
    _ia_text("Le DOE sera constitué progressivement tout au long du chantier et comprendra : plans de récolement, notices d'entretien, fiches techniques, PV d'essais, certificats de conformité.")
    _client_text("Précisez votre process de constitution du DOE, les outils utilisés (GED, plateforme collaborative), le format de remise (papier + numérique).")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  8. SÉCURITÉ ET PRÉVENTION
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Sécurité et prévention des risques", 1)
    _ia_text(
        "La sécurité est une valeur fondamentale de notre entreprise. "
        "Notre politique de prévention vise le zéro accident et s'appuie sur "
        "un système de management de la sécurité structuré."
    )

    doc.add_heading("8.1 Organisation sécurité", 2)
    secu_table = _styled_table(["Mesure", "Proposition IA", "Précisions client"])
    secu_items = [
        ("PPSPS", "Établi avant démarrage, mis à jour à chaque évolution", "Joindre le PPSPS type"),
        ("Accueil sécurité", "Obligatoire pour tout nouvel intervenant", "Durée, contenu, support"),
        ("Causeries sécurité", "Hebdomadaires, thème adapté à la phase", "Fréquence réelle"),
        ("Registre AT/incidents", "Tenu à jour en permanence", "Taux de fréquence AT"),
        ("Audits sécurité terrain", "Mensuels par le responsable QSE", "Fréquence, grille d'audit"),
    ]
    for mesure, ia_val, client_val in secu_items:
        row = secu_table.add_row().cells
        _set_cell_text(row[0], mesure, bold=True, size=Pt(9))
        _set_cell_text(row[1], ia_val, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[1], "EFF6FF")
        _set_cell_text(row[2], client_val, color=ORANGE, size=Pt(8))
        _shade_cell(row[2], "FEF3C7")

    doc.add_heading("8.2 Équipements de protection", 2)
    for item in [
        "EPI : casque, chaussures S3, gilet HV, gants, lunettes, protections auditives",
        "Protections collectives : garde-corps, filets, balisage, signalisation, échafaudages",
        "Premiers secours : trousse, défibrillateur, SST formés (min. 2 par équipe)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.color.rgb = IA_BLUE

    doc.add_heading("8.3 Formations et habilitations", 2)
    _ia_text("Habilitations requises (à vérifier selon les lots) : électrique (B1V/B2V/BR/BC), CACES R386/R389/R482, travail en hauteur, AIPR, SST.")
    _client_text("Précisez votre taux de fréquence AT des 3 dernières années, engagement OPPBTP, certifications MASE le cas échéant.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  9. ENVIRONNEMENT, RSE ET DÉCHETS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Environnement, RSE et gestion des déchets", 1)

    doc.add_heading("9.1 Politique environnementale", 2)
    _ia_text("Nos engagements environnementaux pour ce chantier :")
    for item in [
        "Réduction de l'empreinte carbone : optimisation des rotations camions, matériaux locaux",
        "Matériaux biosourcés et circuits courts privilégiés quand techniquement possible",
        "Suivi des consommations eau/énergie avec objectifs chiffrés",
        "Limitation des nuisances : horaires respectés, arrosage poussières, bâchage",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.color.rgb = IA_BLUE

    doc.add_heading("9.2 SOGED — Gestion des déchets", 2)
    _ia_text("Schéma d'Organisation et de Gestion des Déchets de chantier :")
    dechets_table = _styled_table(["Type de déchet", "Filière (Proposition IA)", "Objectif", "Prestataire"])
    for dtype, filiere, obj, prestataire in [
        ("Inertes (béton, terre, gravats)", "Plateforme ISDI agréée", "> 70% valorisés", ""),
        ("Non dangereux (bois, plâtre, plastique)", "Centre de tri multi-flux", "> 50% valorisés", ""),
        ("Dangereux (amiante, peinture, solvants)", "Filière ISDD agréée", "100% traités", ""),
        ("Métaux (acier, cuivre, alu)", "Ferrailleur / recyclage", "> 90% recyclés", ""),
        ("Emballages", "Collecte sélective sur site", "> 80% recyclés", ""),
    ]:
        row = dechets_table.add_row().cells
        _set_cell_text(row[0], dtype, size=Pt(8))
        _set_cell_text(row[1], filiere, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[1], "EFF6FF")
        _set_cell_text(row[2], obj, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[2], "EFF6FF")
        _set_cell_text(row[3], "À préciser", color=ORANGE, size=Pt(8))
        _shade_cell(row[3], "FEF3C7")

    doc.add_heading("9.3 Considérations sociales et insertion", 2)
    for item in [
        "Clause d'insertion : heures réservées à des publics éloignés de l'emploi",
        "Recours privilégié aux entreprises locales pour la sous-traitance",
        "Conditions de travail conformes aux conventions collectives BTP",
        "Formation et montée en compétence des compagnons",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.color.rgb = IA_BLUE
    _client_text("Complétez avec votre politique RSE concrète, certifications ISO 14001, actions insertion réalisées, label Vert.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  10. RÉFÉRENCES ET EXPÉRIENCES SIMILAIRES
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Références et expériences similaires", 1)
    _ia_text(
        "Les références ci-dessous doivent être similaires au présent marché en termes de : "
        "nature des travaux, montant, complexité technique, type d'acheteur public. "
        "Privilégiez les références récentes (< 5 ans) avec attestations de bonne exécution."
    )

    if po.get("market_type"):
        _add_encadre(
            f"Type de marché : {po['market_type']}. "
            "Sélectionnez des références correspondant à ce type de travaux.",
            bg_hex="EFF6FF", border_hex="2563EB"
        )

    ref_table = _styled_table(["Maître d'ouvrage", "Intitulé du marché", "Montant € HT", "Année", "Attestation"])
    for i in range(5):
        row = ref_table.add_row().cells
        for j, hint in enumerate(["Nom MOA", "Intitulé", "Montant", "Année", "Oui/Non"]):
            _set_cell_text(row[j], hint, color=ORANGE, size=Pt(8))
            _shade_cell(row[j], "FEF3C7")

    doc.add_paragraph()
    _add_encadre(
        "Conseils pour maximiser votre note :\n"
        "- Choisissez 3 à 5 références avec attestation de bonne exécution\n"
        "- Incluez des photos de chantier si possible\n"
        "- Montrez la similitude : même type de bâtiment, même complexité\n"
        "- Indiquez un contact vérifiable chez le maître d'ouvrage",
        bg_hex="ECFDF5", border_hex="10B981",
        text_color=GREEN_DARK
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  ANNEXE A — CHECKLIST
    # ══════════════════════════════════════════════════════════════════════
    if checklist_items:
        doc.add_heading("Annexe A — Checklist de conformité DCE", 1)

        elim_count = sum(1 for i in checklist_items if i.criticality and "liminatoire" in i.criticality.lower())
        manq_count = sum(1 for i in checklist_items if i.status == "MANQUANT")

        _add_encadre(
            f"{len(checklist_items)} exigences identifiées | {elim_count} éliminatoires | {manq_count} manquants\n"
            "Vérifiez chaque point avant soumission. Les éléments éliminatoires manquants entraînent le rejet automatique de l'offre.",
            bg_hex="FEF2F2", border_hex="EF4444",
            text_color=RGBColor(0x99, 0x1B, 0x1B)
        )

        cl_table = _styled_table(["#", "Exigence", "Catégorie", "Criticité", "Statut", "Action"])
        for idx, item in enumerate(checklist_items, 1):
            row = cl_table.add_row().cells
            _set_cell_text(row[0], str(idx), size=Pt(8))
            _set_cell_text(row[1], item.requirement or "—", size=Pt(8))
            _set_cell_text(row[2], item.category or "—", size=Pt(8))
            crit = item.criticality or "—"
            _set_cell_text(row[3], crit, bold=True, color=_severity_color(crit), size=Pt(8))
            status = item.status or "—"
            _set_cell_text(row[4], status, bold=True, color=_status_color(status), size=Pt(8))
            # IA suggestion for action
            if status == "MANQUANT":
                action = item.what_to_provide or "Fournir le document requis"
                _set_cell_text(row[5], action, color=RED, size=Pt(8))
                _shade_row(cl_table.rows[-1], "FEF2F2")
            elif status == "OK":
                _set_cell_text(row[5], "Conforme", color=GREEN_DARK, size=Pt(8))
            else:
                _set_cell_text(row[5], item.what_to_provide or "À vérifier", color=ORANGE, size=Pt(8))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  ANNEXE B — ANALYSE DES RISQUES (2 colonnes : IA + Client)
    # ══════════════════════════════════════════════════════════════════════
    if risks:
        doc.add_heading("Annexe B — Analyse des risques et mesures de mitigation", 1)

        # LLM plan d'action narratif
        if _memo_action_plan_text:
            _ia_text(_memo_action_plan_text)
            doc.add_paragraph()

        # Cashflow chart (si disponible)
        if _cashflow_buf:
            try:
                _cashflow_buf.seek(0)
                doc.add_picture(_cashflow_buf, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            except Exception as _e:
                logger.warning("memo_cashflow_insert_failed", error=str(_e))

        _add_encadre(
            "Les risques suivants ont été identifiés par l'IA lors de l'analyse du DCE. "
            "La colonne 'Proposition IA' contient une suggestion de mitigation. "
            "Complétez la colonne 'Mesure entreprise' avec votre réponse concrète.",
            bg_hex="EFF6FF", border_hex="2563EB"
        )

        risk_table = _styled_table(["Risque", "Sévérité", "Proposition IA", "Mesure entreprise"])
        # IA mitigation suggestions based on risk type
        def _ia_mitigation(risk_text, severity):
            r = (risk_text or "").lower()
            if any(w in r for w in ["delai", "retard", "planning"]):
                return "Prévoir marge de 10-15% sur le planning, identifier chemin critique"
            elif any(w in r for w in ["penalite", "pénalité"]):
                return "Chiffrer l'impact financier, négocier le plafonnement"
            elif any(w in r for w in ["prix", "financ", "cout", "coût", "budget"]):
                return "Vérifier la formule de révision, prévoir clause d'imprévision"
            elif any(w in r for w in ["sous-trait", "sous_trait"]):
                return "DC4 pré-rempli, agréments anticipés, paiement direct prévu"
            elif any(w in r for w in ["technique", "cctp", "norme"]):
                return "Consulter bureau d'études, demander précision à l'acheteur"
            elif any(w in r for w in ["assurance", "garantie", "caution"]):
                return "Vérifier couverture avec assureur, prévoir caution bancaire"
            return "Analyser l'impact et proposer une mesure adaptée"

        for risk in risks:
            row = risk_table.add_row().cells
            _set_cell_text(row[0], risk.get("risk", "—"), size=Pt(8))
            sev = risk.get("severity") or "—"
            sev_label = {"high": "FORT", "medium": "MOYEN", "low": "BAS"}.get(
                sev.lower() if isinstance(sev, str) else "", sev.upper() if isinstance(sev, str) else str(sev))
            _set_cell_text(row[1], sev_label, bold=True,
                          color=_severity_color(sev), size=Pt(8))
            # IA suggestion
            ia_mit = _ia_mitigation(risk.get("risk"), sev)
            _set_cell_text(row[2], ia_mit, color=IA_BLUE, size=Pt(8))
            _shade_cell(row[2], "EFF6FF")
            # Client to fill
            _set_cell_text(row[3], "Votre mesure concrète", color=ORANGE, size=Pt(8))
            _shade_cell(row[3], "FEF3C7")

    # ══════════════════════════════════════════════════════════════════════
    #  AVERTISSEMENT IA + FOOTER
    # ══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("Avertissement", 1)
    _add_encadre(
        "Ce document a été pré-rempli par intelligence artificielle (AO Copilot) "
        "à partir de l'analyse automatique des documents du DCE et des données "
        "de votre profil entreprise.\n\n"
        "Les éléments en bleu sont des propositions IA à valider et personnaliser.\n"
        "Les éléments en orange doivent être complétés avec vos données réelles.\n\n"
        "Ce document ne se substitue pas à l'expertise d'un rédacteur marchés publics. "
        "Vérifiez systématiquement que le contenu final correspond à vos capacités "
        "réelles et aux exigences spécifiques du RC.",
        bg_hex="FFFBEB", border_hex="FBBF24",
        text_color=RGBColor(0x78, 0x35, 0x0F)
    )

    doc.add_paragraph()
    doc.add_paragraph()
    footer_brand = doc.add_paragraph()
    footer_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fb = footer_brand.add_run("AO COPILOT")
    run_fb.bold = True
    run_fb.font.color.rgb = ACCENT_BLUE
    run_fb.font.size = Pt(11)

    gen_para = doc.add_paragraph()
    gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_gp = gen_para.add_run(
        f"aocopilot.fr — {datetime.now().strftime('%d/%m/%Y %H:%M')} — "
        "Document confidentiel"
    )
    run_gp.font.color.rgb = GRAY_LIGHT
    run_gp.font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
