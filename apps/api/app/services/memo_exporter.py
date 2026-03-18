"""Génération de la mémoire technique Word (.docx) — extrait de exporter.py."""
import structlog
from datetime import datetime
from io import BytesIO
from sqlalchemy.orm import Session

from app.services.export_data import fetch_export_data

logger = structlog.get_logger(__name__)


def generate_memo_technique(db: Session, project_id: str) -> bytes:
    """Génère une mémoire technique Word (.docx) complète — 10 sections standard BTP.
    Design identique au rapport PDF/DOCX. Système hybride :
    - Texte bleu = pré-rempli par IA (proposition intelligente basée sur le DCE)
    - Texte orange = à compléter par le client (données entreprise spécifiques)
    Conforme aux attentes acheteurs publics (marche-public.fr, Doaken 2026).
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from io import BytesIO
    except ImportError as e:
        raise RuntimeError("python-docx non installé. Ajoutez python-docx==1.1.0.") from e

    from app.models.company_profile import CompanyProfile
    from app.models.organization import Organization

    # ── Charger toutes les données via fetch_export_data ──────────────────
    data = fetch_export_data(db, project_id)
    project = data.project

    summary = data.summary or {}
    criteria = data.criteria or {}
    timeline = data.timeline or {}
    cctp = data.cctp_analysis or {}
    ccap = data.ccap_analysis or {}
    scoring = data.scoring or {}
    rc = data.rc_analysis or {}
    subcontracting = data.subcontracting or {}
    gonogo = data.gonogo or {}
    checklist_items = data.checklist_items
    documents = data.documents
    cashflow_data = data.cashflow or {}

    po = summary.get("project_overview", {})
    key_points = summary.get("key_points", [])
    risks = summary.get("risks", [])
    evaluation = criteria.get("evaluation", {})
    scoring_criteria = evaluation.get("scoring_criteria", [])

    # Charger le profil entreprise
    org = db.query(Organization).filter_by(id=project.org_id).first()
    company = None
    if org:
        company = db.query(CompanyProfile).filter_by(org_id=org.id).first()
    org_name = (org.name if org else None) or "Notre Entreprise"

    # ── Génération des narratifs LLM (graceful degradation) ──────────────
    _memo_intro_text: str | None = None
    _memo_positioning_text: str | None = None
    _memo_action_plan_text: str | None = None
    try:
        from app.services.llm import llm_service
        from app.services.prompts import (
            build_memo_intro_prompt,
            build_memo_positioning_prompt,
            build_memo_action_plan_prompt,
        )
        _company_dict = {}
        if company:
            _company_dict = {
                "name": org_name,
                "activity_sector": getattr(company, "activity_sector", None),
                "annual_revenue_eur": getattr(company, "revenue_eur", None),
                "certifications": getattr(company, "certifications", None),
                "regions": getattr(company, "regions", None),
                "staff_count": getattr(company, "employee_count", None),
                "main_clients": getattr(company, "main_clients", None),
                "references_btp": getattr(company, "references_btp", None),
            }
        _gonogo_score = gonogo.get("score") or 0
        _gonogo_dims = gonogo.get("dimension_scores") or gonogo.get("breakdown") or {}
        _deadline_str = (
            project.submission_deadline.strftime("%d/%m/%Y")
            if project.submission_deadline else "—"
        )

        sys_p, usr_p = build_memo_intro_prompt(
            project_title=project.title or "AO",
            buyer=project.buyer or "Acheteur public",
            scope=po.get("scope") or po.get("object") or "Prestations à définir",
            go_nogo_score=_gonogo_score,
            top_risks=risks[:3],
            company_profile=_company_dict,
        )
        _memo_intro_text = llm_service.chat_text(sys_p, usr_p, max_tokens=400)

        sys_p, usr_p = build_memo_positioning_prompt(
            company_profile=_company_dict,
            gonogo_dimensions=_gonogo_dims if isinstance(_gonogo_dims, dict) else {},
            eligibility_gaps=gonogo.get("profile_gaps") or [],
        )
        _memo_positioning_text = llm_service.chat_text(sys_p, usr_p, max_tokens=400)

        _actions = summary.get("actions_next_48h") or []
        sys_p, usr_p = build_memo_action_plan_prompt(
            actions_48h=_actions,
            risks=risks,
            deadline_submission=_deadline_str,
        )
        _memo_action_plan_text = llm_service.chat_text(sys_p, usr_p, max_tokens=500)
    except Exception as _llm_err:
        logger.warning("memo_llm_generation_skipped", error=str(_llm_err))

    # ── Génération des graphiques (graceful degradation) ──────────────────
    _radar_buf = None
    _cashflow_buf = None
    try:
        from app.services.chart_generator import generate_gonogo_radar, generate_cashflow_chart
        if gonogo:
            _dims = gonogo.get("dimension_scores") or gonogo.get("breakdown") or {}
            _radar_buf = generate_gonogo_radar(
                _dims if isinstance(_dims, dict) else {},
                gonogo.get("score"),
                project.title or "AO",
            )
        if cashflow_data:
            _cashflow_buf = generate_cashflow_chart(cashflow_data, project.title or "AO")
    except Exception as _chart_err:
        logger.warning("memo_chart_generation_skipped", error=str(_chart_err))

    doc = Document()

    # ── Couleurs (identiques au rapport DOCX) ────────────────────────────
    DARK_BLUE = RGBColor(0x0F, 0x1B, 0x4C)
    ACCENT_BLUE = RGBColor(0x25, 0x63, 0xEB)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    RED = RGBColor(0xB9, 0x1C, 0x1C)
    GREEN_DARK = RGBColor(0x04, 0x78, 0x57)
    ORANGE = RGBColor(0xD9, 0x77, 0x06)
    GRAY = RGBColor(0x64, 0x74, 0x8B)
    GRAY_LIGHT = RGBColor(0x9C, 0xA3, 0xAF)
    IA_BLUE = RGBColor(0x1E, 0x40, 0xAF)  # Texte proposition IA

    # ── Style global ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)  # 11pt — BTP audit: lisibilité comités (50+ ans)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = DARK_BLUE
        hs.font.name = "Calibri"

    # ── Helpers (identiques au rapport DOCX) ─────────────────────────────
    def _shade_cell(cell, hex_color):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _shade_row(row, hex_color):
        for cell in row.cells:
            _shade_cell(cell, hex_color)

    def _set_cell_text(cell, text, bold=False, color=None, size=None):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = bold
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = size
        run.font.name = "Calibri"
        return run

    def _styled_table(headers):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            _shade_cell(cell, "0F1B4C")
            _set_cell_text(cell, h, bold=True, color=WHITE, size=Pt(9))
        return t

    def _fmt_date(val):
        if not val:
            return "—"
        s = str(val)
        if "T" in s:
            s = s.split("T")[0]
        try:
            from datetime import datetime as dt
            d = dt.strptime(s, "%Y-%m-%d")
            return d.strftime("%d/%m/%Y")
        except Exception:
            return s

    def _add_encadre(text, bg_hex="EFF6FF", border_hex="2563EB", text_color=None):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_hex)
        pPr.append(shd)
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "36")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), border_hex)
        pBdr.append(left)
        pPr.append(pBdr)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "180")
        ind.set(qn("w:right"), "180")
        pPr.append(ind)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        if text_color:
            run.font.color.rgb = text_color
        return p

    def _ia_text(text):
        """Paragraphe pré-rempli par l'IA (bleu, fond bleu clair)."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "EFF6FF")
        pPr.append(shd)
        run = p.add_run(text)
        run.font.color.rgb = IA_BLUE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        return p

    def _client_text(text):
        """Paragraphe à compléter par le client (orange, fond orange clair)."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "FEF3C7")
        pPr.append(shd)
        run = p.add_run(text)
        run.font.color.rgb = ORANGE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        return p

    def _severity_color(sev):
        s = (sev or "").lower()
        if s in ("high", "fort", "haut", "eliminatoire"):
            return RED
        elif s in ("medium", "moyen"):
            return ORANGE
        elif s in ("low", "bas"):
            return GREEN_DARK
        return GRAY

    def _status_color(status):
        s = (status or "").upper()
        if s == "OK":
            return GREEN_DARK
        elif s in ("MANQUANT", "KO"):
            return RED
        elif s in ("PARTIEL",):
            return ORANGE
        return GRAY

    # ── En-tête & Pied de page ─────────────────────────────────────────
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = ""
    run_h = header_para.add_run(f"{org_name} — Mémoire Technique — {project.title[:50]}")
    run_h.font.size = Pt(8)
    run_h.font.color.rgb = GRAY
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Page ")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = GRAY
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run_xml = footer_para.add_run()
    run_xml.font.size = Pt(8)
    run_xml._r.append(fldChar_begin)
    run_xml._r.append(instrText)
    run_xml._r.append(fldChar_end)

    # ══════════════════════════════════════════════════════════════════════
    #  PAGE DE GARDE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_b = brand.add_run(org_name.upper())
    run_b.font.color.rgb = DARK_BLUE
    run_b.font.size = Pt(14)
    run_b.bold = True

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = line.add_run("━━━━━━━━━━")
    run_l.font.color.rgb = ACCENT_BLUE

    title_heading = doc.add_heading("MÉMOIRE TECHNIQUE", 0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_heading.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sub.add_run(project.title)
    run_s.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    run_s.font.size = Pt(14)

    doc.add_paragraph()

    # Info table
    cover_info = _styled_table(["Information", "Détail"])
    deadline_str = project.submission_deadline.strftime("%d/%m/%Y") if project.submission_deadline else "—"
    for label, value in [
        ("Marché", project.title),
        ("Acheteur public", project.buyer or "—"),
        ("Référence", project.reference or "—"),
        ("Candidat", org_name),
        ("Lot(s)", ", ".join(l.get("title", "") for l in rc.get("lots", [])) or "Lot unique"),
        ("Date de remise", deadline_str),
        ("Document établi le", datetime.now().strftime("%d/%m/%Y")),
    ]:
        row = cover_info.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=Pt(9))
        _set_cell_text(row[1], str(value), size=Pt(9))

    doc.add_paragraph()

    # Legend box
    _add_encadre(
        "Légende des couleurs dans ce document :\n"
        "Texte bleu sur fond bleu clair = Proposition IA (pré-rempli à partir du DCE — à valider)\n"
        "Texte orange sur fond jaune clair = À compléter par le candidat (données entreprise)",
        bg_hex="F8FAFC", border_hex="2563EB"
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SOMMAIRE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Sommaire", 1)

    sommaire = [
        ("1.", "Présentation de l'entreprise — Identité, certifications, spécialités"),
        ("2.", "Compréhension du projet — Analyse du DCE et enjeux identifiés"),
        ("3.", "Méthodologie d'exécution — Préparation, exécution, réception"),
        ("4.", "Moyens humains — Organigramme, qualifications, effectifs"),
        ("5.", "Moyens matériels et techniques — Équipements et outils"),
        ("6.", "Planning prévisionnel — Jalons DCE et phasage d'exécution"),
        ("7.", "Gestion de la qualité (SOPAQ) — Contrôles, traçabilité, DOE"),
        ("8.", "Sécurité et prévention — PPSPS, EPI, formations"),
        ("9.", "Environnement, RSE et gestion des déchets — SOGED, insertion"),
        ("10.", "Références et expériences similaires"),
        ("A.", "Checklist de conformité DCE"),
        ("B.", "Analyse des risques — Proposition IA + mesures client"),
    ]
    for num, label in sommaire:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "dotted")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CBD5E1")
        pBdr.append(bottom)
        pPr.append(pBdr)
        run_num = p.add_run(f"{num}  ")
        run_num.bold = True
        run_num.font.color.rgb = ACCENT_BLUE
        run_num.font.size = Pt(10)
        run_text = p.add_run(label)
        run_text.font.size = Pt(10)
        run_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Critères de notation rappel + Matrice de conformité (audit BTP: mapper chaque section aux critères RC)
    if scoring_criteria:
        doc.add_paragraph()
        _add_encadre(
            "Rappel des critères de notation (RC) :\n" +
            "\n".join(
                f"- {c.get('criterion', '')}: {c.get('weight_percent', '?')}%"
                for c in scoring_criteria
            ) +
            "\n\nCe document est structuré pour maximiser votre note technique.",
            bg_hex="ECFDF5", border_hex="10B981",
            text_color=GREEN_DARK
        )

        # Conformity mapping table — maps each section to RC criteria
        doc.add_paragraph()
        doc.add_heading("Matrice de conformité — Sections ↔ Critères RC", 3)
        # Build mapping: each section addresses certain criteria keywords
        _criteria_section_map = [
            ("1. Présentation entreprise", ["capacité", "moyen", "compétence", "expérience", "qualif"]),
            ("2. Compréhension du projet", ["technique", "méthodologie", "compréhension"]),
            ("3. Méthodologie d'exécution", ["technique", "méthodologie", "qualité"]),
            ("4. Moyens humains", ["moyen", "personnel", "humain", "compétence"]),
            ("5. Moyens matériels", ["moyen", "matériel", "technique"]),
            ("6. Planning", ["délai", "planning", "organisation"]),
            ("7. Qualité (SOPAQ)", ["qualité", "contrôle", "gestion"]),
            ("8. Sécurité", ["sécurité", "prévention", "santé"]),
            ("9. RSE / Environnement", ["environnement", "développement durable", "rse", "social", "insertion"]),
            ("10. Références", ["référence", "expérience", "capacité"]),
        ]
        conf_table = _styled_table(["Section Mémoire", "Critère(s) RC adressé(s)", "Pondération"])
        for sec_name, keywords in _criteria_section_map:
            matched = [c for c in scoring_criteria
                       if any(kw in (c.get("criterion") or "").lower() for kw in keywords)]
            if matched:
                row = conf_table.add_row().cells
                _set_cell_text(row[0], sec_name, bold=True, size=Pt(8))
                criteria_text = ", ".join(c.get("criterion", "") for c in matched)
                _set_cell_text(row[1], criteria_text[:80], size=Pt(8))
                total_weight = sum(c.get("weight_percent", 0) or 0 for c in matched)
                _set_cell_text(row[2], f"{total_weight}%" if total_weight else "—",
                              bold=True, color=ACCENT_BLUE, size=Pt(9))

    doc.add_page_break()

    # ── Radar chart Go/No-Go (si disponible) ──────────────────────────────
    if _radar_buf:
        try:
            _radar_buf.seek(0)
            doc.add_heading("Synthèse Go/No-Go — 9 dimensions", 2)
            doc.add_picture(_radar_buf, width=Inches(5.5))
            # doc.add_picture() adds the picture to the last paragraph
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        except Exception as _e:
            logger.warning("memo_radar_insert_failed", error=str(_e))

    # ══════════════════════════════════════════════════════════════════════
    #  1. PRÉSENTATION DE L'ENTREPRISE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Présentation de l'entreprise", 1)

    # LLM positionnement stratégique
    if _memo_positioning_text:
        _ia_text(_memo_positioning_text)
        doc.add_paragraph()

    doc.add_heading("1.1 Identité", 2)
    id_table = _styled_table(["Rubrique", "Proposition IA", "À compléter / vérifier"])
    id_fields = [
        ("Raison sociale", org_name, "Vérifier"),
        ("Chiffre d'affaires", f"{company.revenue_eur:,} €".replace(",", " ") if company and company.revenue_eur else "—", "CA dernier exercice clos"),
        ("Effectif", f"{company.employee_count} salariés" if company and company.employee_count else "—", "Effectif total CDI"),
        ("Forme juridique", "—", "SAS, SARL, SA, etc."),
        ("SIRET", "—", "N° SIRET à 14 chiffres"),
        ("Code APE / NAF", "—", "Code activité"),
        ("Adresse siège", "—", "Adresse complète"),
        ("Date de création", "—", "Année de création"),
    ]
    for label, ia_val, client_val in id_fields:
        row = id_table.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=Pt(9))
        if ia_val != "—":
            _set_cell_text(row[1], ia_val, color=IA_BLUE, size=Pt(9))
            _shade_cell(row[1], "EFF6FF")
        else:
            _set_cell_text(row[1], "—", color=GRAY, size=Pt(9))
        _set_cell_text(row[2], client_val, color=ORANGE, size=Pt(9))
        _shade_cell(row[2], "FEF3C7")

    # Certifications
    doc.add_heading("1.2 Certifications et qualifications", 2)
    if company and company.certifications:
        _ia_text("Certifications identifiées dans votre profil (à compléter avec dates de validité) :")
        for cert in company.certifications:
            p = doc.add_paragraph(style="List Bullet")
            run_c = p.add_run(str(cert))
            run_c.bold = True
            run_c.font.color.rgb = IA_BLUE
    else:
        _client_text("Listez vos certifications : Qualibat (n° et classification), ISO 9001, ISO 14001, MASE, RGE, etc. Indiquez les dates de validité.")

    # Vérifier exigences RC
    elim_certs = [item for item in checklist_items
                  if item.category and "certif" in item.category.lower()
                  and item.criticality and "liminatoire" in item.criticality.lower()]
    if elim_certs:
        _add_encadre(
            "ATTENTION — Certifications éliminatoires exigées par le RC :\n" +
            "\n".join(f"- {e.requirement}" for e in elim_certs),
            bg_hex="FEF2F2", border_hex="EF4444",
            text_color=RGBColor(0x99, 0x1B, 0x1B)
        )

    # Spécialités
    doc.add_heading("1.3 Domaines de spécialité", 2)
    if company and company.specialties:
        for spec in company.specialties:
            doc.add_paragraph(str(spec).replace("_", " ").title(), style="List Bullet")
    else:
        _client_text("Listez vos domaines : gros-oeuvre, électricité, plomberie, CVC, couverture, etc.")

    # Zones d'intervention
    doc.add_heading("1.4 Zones d'intervention", 2)
    if company and company.regions:
        doc.add_paragraph(", ".join(str(r) for r in company.regions))
    location = po.get("location") or ""
    if location:
        _ia_text(f"Le marché se situe à : {location}. Confirmez que cette zone est couverte.")

    # Assurances
    doc.add_heading("1.5 Assurances", 2)
    assur_table = _styled_table(["Type d'assurance", "Proposition IA", "À compléter"])
    assur_fields = [
        ("Assurance décennale", "Souscrite" if company and company.assurance_decennale else "—", "N° police + assureur + date validité"),
        ("RC Professionnelle", f"{company.assurance_rc_montant:,} €".replace(",", " ") if company and company.assurance_rc_montant else "—", "Montant couverture + assureur"),
        ("RC Exploitation", "—", "Montant + assureur"),
        ("Tous Risques Chantier", "—", "Si souscrite"),
    ]
    for label, ia_val, client_val in assur_fields:
        row = assur_table.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=Pt(9))
        if ia_val != "—":
            _set_cell_text(row[1], ia_val, color=IA_BLUE, size=Pt(9))
            _shade_cell(row[1], "EFF6FF")
        else:
            _set_cell_text(row[1], "—", color=GRAY, size=Pt(9))
        _set_cell_text(row[2], client_val, color=ORANGE, size=Pt(9))
        _shade_cell(row[2], "FEF3C7")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  2. COMPRÉHENSION DU PROJET
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Compréhension du projet", 1)

    scope_text = po.get("scope") or po.get("object") or "Le présent marché porte sur des prestations à définir."
    market_type = po.get("market_type") or "non précisé"

    # Utiliser le narratif LLM si disponible, sinon fallback template
    if _memo_intro_text:
        _ia_text(_memo_intro_text)
    else:
        _ia_text(
            f"Dans le cadre de la consultation lancée par {project.buyer or 'l acheteur public'}, "
            f"portant sur un marché de {market_type}, notre société {org_name} "
            f"a analysé en détail les documents du Dossier de Consultation des Entreprises "
            f"afin de formuler une réponse parfaitement adaptée aux besoins exprimés."
        )

    doc.add_heading("2.1 Objet et périmètre", 2)
    _ia_text(scope_text)
    if location:
        _ia_text(f"Lieu d'exécution : {location}")
    if po.get("estimated_budget"):
        _ia_text(f"Enveloppe budgétaire estimée : {po['estimated_budget']}")
    duration = timeline.get("execution_duration_months")
    if duration:
        _ia_text(f"Durée d'exécution : {duration} mois")

    # Points clés (IA)
    if key_points:
        doc.add_heading("2.2 Points clés identifiés dans le DCE", 2)
        _add_encadre(
            "Les points suivants ont été automatiquement extraits du DCE par l'IA. "
            "Reprenez-les dans votre rédaction pour montrer votre compréhension du projet.",
            bg_hex="EFF6FF", border_hex="2563EB"
        )
        for kp in key_points:
            point = kp.get("point") or kp.get("label") or ""
            if point:
                p = doc.add_paragraph(style="List Bullet")
                run_kp = p.add_run(str(point))
                run_kp.font.color.rgb = IA_BLUE

    # Enjeux spécifiques
    doc.add_heading("2.3 Enjeux spécifiques et contraintes", 2)
    _client_text(
        "Décrivez ici votre compréhension personnelle des enjeux du projet. "
        "Mentionnez : la visite de site (si effectuée), les contraintes d'accès, "
        "la coactivité avec d'autres lots, les interfaces techniques, "
        "les exigences particulières du maître d'ouvrage."
    )

    # Contraintes CCTP
    if cctp:
        if cctp.get("environmental_requirements"):
            doc.add_heading("2.4 Contraintes environnementales identifiées", 2)
            for req in cctp["environmental_requirements"][:5]:
                text = req.get("requirement") if isinstance(req, dict) else str(req)
                p = doc.add_paragraph(style="List Bullet")
                run_r = p.add_run(text)
                run_r.font.color.rgb = IA_BLUE

        if cctp.get("contradictions"):
            _add_encadre(
                "ATTENTION — Contradictions détectées dans le CCTP :\n" +
                "\n".join(
                    f"- {(c.get('description') or c.get('contradiction') or str(c))}"
                    for c in cctp["contradictions"][:3]
                    if isinstance(c, dict)
                ) +
                "\nPosez une question à l'acheteur pour lever ces ambiguïtés.",
                bg_hex="FEF2F2", border_hex="EF4444",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  3. MÉTHODOLOGIE D'EXÉCUTION
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Méthodologie d'exécution", 1)
    _ia_text(
        "Notre approche repose sur une méthodologie rigoureuse et éprouvée, "
        "structurée en phases distinctes pour garantir la maîtrise des délais, "
        "de la qualité et des coûts. Chaque étape intègre les exigences "
        "spécifiques identifiées dans le CCTP."
    )

    doc.add_heading("3.1 Préparation de chantier", 2)
    prep_items = [
        "Analyse approfondie du CCTP, du RC et des pièces contractuelles",
        "Visite de site et relevé des contraintes in situ",
        "Validation des interfaces avec les autres intervenants et lots",
        "Élaboration du PPSPS et du plan d'installation de chantier",
        "Commande anticipée des matériaux à délais longs",
        "Mise en place de la base vie et des accès sécurisés",
    ]
    for item in prep_items:
        p = doc.add_paragraph(style="List Bullet")
        run_i = p.add_run(item)
        run_i.font.color.rgb = IA_BLUE
    _client_text("Complétez avec vos spécificités : méthodes constructives retenues, phasage logistique, plan d'installation de chantier.")

    doc.add_heading("3.2 Phase d'exécution", 2)
    exec_items = [
        "Suivi quotidien de l'avancement par le conducteur de travaux dédié",
        "Réunions de chantier hebdomadaires avec compte-rendu formalisé sous 48h",
        "Autocontrôles systématiques à chaque point d'arrêt défini dans le SOPAQ",
        "Gestion documentaire centralisée (plans d'exécution, DOE, fiches techniques)",
        "Coordination avec le CSPS et les autres entreprises",
        "Reporting mensuel d'avancement au maître d'ouvrage",
    ]
    for item in exec_items:
        p = doc.add_paragraph(style="List Bullet")
        run_i = p.add_run(item)
        run_i.font.color.rgb = IA_BLUE
    _client_text("Détaillez vos process internes, outils de suivi (logiciel de gestion), méthodes de reporting, gestion des approvisionnements.")

    doc.add_heading("3.3 Réception et levée de réserves", 2)
    for item in [
        "Pré-réception interne systématique avant les OPR",
        "Levée des réserves dans les délais contractuels (30 jours max)",
        "Remise du DOE complet et des plans de récolement sous 30 jours",
        "Garantie de parfait achèvement : intervention sous 48h en cas de sinistre",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run_i = p.add_run(item)
        run_i.font.color.rgb = IA_BLUE

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  4. MOYENS HUMAINS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Moyens humains", 1)
    if company and company.employee_count:
        _ia_text(
            f"Notre entreprise compte {company.employee_count} collaborateurs. "
            "L'équipe dédiée au présent marché est composée de professionnels qualifiés "
            "et expérimentés, sous la responsabilité d'un interlocuteur unique."
        )
    else:
        _ia_text(
            "L'équipe dédiée au présent marché est composée de professionnels qualifiés "
            "et expérimentés, sous la responsabilité d'un interlocuteur unique."
        )

    doc.add_heading("4.1 Organigramme de l'équipe projet", 2)
    _add_encadre(
        "Le tableau ci-dessous doit être renseigné avec les personnes nommément "
        "affectées au marché. Les CV peuvent être joints en annexe si le RC l'exige.",
        bg_hex="FEF3C7", border_hex="D97706",
        text_color=ORANGE
    )
    team_table = _styled_table(["Fonction", "Proposition IA", "Nom (à compléter)", "Qualification", "Expérience"])
    team_roles = [
        ("Directeur de travaux", "Pilotage global, validation technique"),
        ("Conducteur de travaux", "Suivi quotidien, coordination lots"),
        ("Chef de chantier", "Encadrement terrain, qualité"),
        ("Chef d'équipe", "Exécution, management compagnons"),
        ("Responsable QSE", "Qualité, sécurité, environnement"),
    ]
    for role, ia_desc in team_roles:
        row = team_table.add_row().cells
        _set_cell_text(row[0], role, bold=True, size=Pt(9))
        _set_cell_text(row[1], ia_desc, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[1], "EFF6FF")
        _set_cell_text(row[2], "Nom, prénom", color=ORANGE, size=Pt(8))
        _shade_cell(row[2], "FEF3C7")
        _set_cell_text(row[3], "Diplôme / habilitation", color=ORANGE, size=Pt(8))
        _shade_cell(row[3], "FEF3C7")
        _set_cell_text(row[4], "X ans", color=ORANGE, size=Pt(8))
        _shade_cell(row[4], "FEF3C7")

    doc.add_heading("4.2 Effectif mobilisé par phase", 2)
    eff_table = _styled_table(["Phase", "Ouvriers", "Techniciens", "Encadrement"])
    for phase in ["Préparation", "Gros-oeuvre", "Second oeuvre", "Finitions / OPR"]:
        row = eff_table.add_row().cells
        _set_cell_text(row[0], phase, bold=True, size=Pt(9))
        for i in range(1, 4):
            _set_cell_text(row[i], "À préciser", color=ORANGE, size=Pt(8))
            _shade_cell(row[i], "FEF3C7")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  5. MOYENS MATÉRIELS ET TECHNIQUES
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Moyens matériels et techniques", 1)

    doc.add_heading("5.1 Matériels affectés au chantier", 2)
    mat_table = _styled_table(["Type", "Désignation (modèle)", "Qté", "Propriété / Location"])
    for mat_type, ia_sugg in [
        ("Gros matériel", "Grue, nacelle, mini-pelle — à préciser"),
        ("Véhicules", "Fourgons, camions — à préciser"),
        ("Outillage spécifique", "Selon nature des travaux"),
        ("Matériel de mesure", "Niveau laser, théodolite, etc."),
        ("Matériel de sécurité", "Garde-corps, filets, balisage"),
    ]:
        row = mat_table.add_row().cells
        _set_cell_text(row[0], mat_type, bold=True, size=Pt(9))
        _set_cell_text(row[1], ia_sugg, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[1], "EFF6FF")
        _set_cell_text(row[2], "?", color=ORANGE, size=Pt(8))
        _shade_cell(row[2], "FEF3C7")
        _set_cell_text(row[3], "Propre / Location", color=ORANGE, size=Pt(8))
        _shade_cell(row[3], "FEF3C7")

    doc.add_heading("5.2 Outils numériques", 2)
    _ia_text("Outils recommandés pour valoriser votre offre :")
    for tool in [
        "Logiciel BIM (Revit, ArchiCAD) pour la maquette numérique",
        "Plateforme de suivi chantier (BatiScript, Fieldwire, Finalcad)",
        "GMAO pour la maintenance préventive des équipements",
        "Outils de reporting photo/vidéo horodatés",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run_t = p.add_run(tool)
        run_t.font.color.rgb = IA_BLUE
    _client_text("Listez vos outils réellement utilisés. Ne mentionnez que ce que vous maîtrisez.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  6. PLANNING PRÉVISIONNEL
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Planning prévisionnel", 1)

    deadline_str = project.submission_deadline.strftime("%d/%m/%Y") if project.submission_deadline else "à définir"
    duration = timeline.get("execution_duration_months")
    _ia_text(
        f"Date limite de remise : {deadline_str}. "
        + (f"Durée d'exécution prévue au marché : {duration} mois. " if duration else "")
        + "Le planning ci-dessous intègre les jalons contractuels identifiés dans le DCE."
    )

    # Jalons DCE
    key_dates = timeline.get("key_dates", [])
    if key_dates or timeline.get("submission_deadline"):
        doc.add_heading("6.1 Jalons contractuels du DCE", 2)
        jal_table = _styled_table(["Échéance", "Date", "Caractère"])
        if timeline.get("submission_deadline"):
            row = jal_table.add_row().cells
            _set_cell_text(row[0], "Date limite de remise", bold=True, size=Pt(9))
            _set_cell_text(row[1], _fmt_date(timeline["submission_deadline"]),
                          bold=True, color=RED, size=Pt(9))
            _set_cell_text(row[2], "IMPÉRATIF", bold=True, color=RED, size=Pt(9))
            _shade_row(jal_table.rows[-1], "FEF2F2")
        for kd in key_dates:
            if "remise des offres" in (kd.get("label") or "").lower():
                continue
            row = jal_table.add_row().cells
            _set_cell_text(row[0], kd.get("label", "—"), size=Pt(9))
            _set_cell_text(row[1], _fmt_date(kd.get("date")), color=IA_BLUE, size=Pt(9))
            mandatory = "OBLIGATOIRE" if kd.get("mandatory") else "Indicatif"
            _set_cell_text(row[2], mandatory, bold=kd.get("mandatory", False),
                          color=RED if kd.get("mandatory") else GRAY, size=Pt(9))

    doc.add_heading("6.2 Phasage prévisionnel d'exécution", 2)
    phase_table = _styled_table(["Phase", "Description", "Durée (Proposition IA)", "Durée réelle"])
    phases = [
        ("1 — Installation", "Base vie, clôture, accès, panneau", "2 semaines"),
        ("2 — Terrassement", "Terrassements, fondations", "À estimer"),
        ("3 — Gros-oeuvre", "Élévations, planchers, charpente", "À estimer"),
        ("4 — Clos & couvert", "Couverture, étanchéité, menuiseries ext.", "À estimer"),
        ("5 — Second oeuvre", "Cloisons, revêtements, peinture", "À estimer"),
        ("6 — Lots techniques", "CVC, plomberie, électricité, essais", "À estimer"),
        ("7 — VRD / Ext.", "Voirie, réseaux, espaces verts", "À estimer"),
        ("8 — OPR / Réception", "Pré-réception, levée réserves, DOE", "4 semaines"),
    ]
    for phase, desc, ia_duree in phases:
        row = phase_table.add_row().cells
        _set_cell_text(row[0], phase, bold=True, size=Pt(9))
        _set_cell_text(row[1], desc, size=Pt(8))
        _set_cell_text(row[2], ia_duree, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[2], "EFF6FF")
        _set_cell_text(row[3], "À compléter", color=ORANGE, size=Pt(8))
        _shade_cell(row[3], "FEF3C7")

    _client_text("Adaptez les phases et durées au marché spécifique. Joignez un diagramme de Gantt en annexe si possible. Un planning réaliste (intégrant aléas météo et interfaces) sera mieux noté qu'un planning optimiste.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  7. GESTION DE LA QUALITÉ (SOPAQ)
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Gestion de la qualité (SOPAQ)", 1)
    _ia_text(
        "Notre Schéma Organisationnel du Plan d'Assurance Qualité (SOPAQ) garantit "
        "la conformité des prestations aux spécifications du CCTP et aux normes en vigueur. "
        "Il sera décliné en PAQ définitif en phase préparation."
    )

    doc.add_heading("7.1 Points d'arrêt et contrôles", 2)
    qa_table = _styled_table(["Point de contrôle", "Fréquence", "Responsable", "Document"])
    qa_items = [
        ("Réception matériaux", "Chaque livraison", "Chef de chantier", "Fiche de réception"),
        ("Contrôle béton (éprouvettes)", "Chaque coulage", "Labo agréé", "PV d'essai"),
        ("Contrôle soudures", "100% soudures", "Bureau de contrôle", "PV de contrôle"),
        ("Étanchéité", "Par zone", "Responsable QSE", "Fiche autocontrôle"),
        ("Conformité plans d'exécution", "Avant chaque phase", "Conducteur travaux", "Visa MOE"),
        ("Essais réseaux (air, eau)", "Avant réception", "Bureau de contrôle", "PV d'essai"),
    ]
    for controle, freq, resp, doc_type in qa_items:
        row = qa_table.add_row().cells
        _set_cell_text(row[0], controle, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[0], "EFF6FF")
        _set_cell_text(row[1], freq, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[1], "EFF6FF")
        _set_cell_text(row[2], resp, color=ORANGE, size=Pt(8))
        _shade_cell(row[2], "FEF3C7")
        _set_cell_text(row[3], doc_type, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[3], "EFF6FF")

    doc.add_heading("7.2 Gestion des non-conformités", 2)
    for item in [
        "Détection → Fiche de non-conformité → Analyse cause → Action corrective → Vérification",
        "Traçabilité complète des matériaux (fiches techniques, PV d'essais, certificats)",
        "Audits internes trimestriels par le responsable QSE",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.color.rgb = IA_BLUE

    doc.add_heading("7.3 DOE (Dossier des Ouvrages Exécutés)", 2)
    _ia_text("Le DOE sera constitué progressivement tout au long du chantier et comprendra : plans de récolement, notices d'entretien, fiches techniques, PV d'essais, certificats de conformité.")
    _client_text("Précisez votre process de constitution du DOE, les outils utilisés (GED, plateforme collaborative), le format de remise (papier + numérique).")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  8. SÉCURITÉ ET PRÉVENTION
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Sécurité et prévention des risques", 1)
    _ia_text(
        "La sécurité est une valeur fondamentale de notre entreprise. "
        "Notre politique de prévention vise le zéro accident et s'appuie sur "
        "un système de management de la sécurité structuré."
    )

    doc.add_heading("8.1 Organisation sécurité", 2)
    secu_table = _styled_table(["Mesure", "Proposition IA", "Précisions client"])
    secu_items = [
        ("PPSPS", "Établi avant démarrage, mis à jour à chaque évolution", "Joindre le PPSPS type"),
        ("Accueil sécurité", "Obligatoire pour tout nouvel intervenant", "Durée, contenu, support"),
        ("Causeries sécurité", "Hebdomadaires, thème adapté à la phase", "Fréquence réelle"),
        ("Registre AT/incidents", "Tenu à jour en permanence", "Taux de fréquence AT"),
        ("Audits sécurité terrain", "Mensuels par le responsable QSE", "Fréquence, grille d'audit"),
    ]
    for mesure, ia_val, client_val in secu_items:
        row = secu_table.add_row().cells
        _set_cell_text(row[0], mesure, bold=True, size=Pt(9))
        _set_cell_text(row[1], ia_val, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[1], "EFF6FF")
        _set_cell_text(row[2], client_val, color=ORANGE, size=Pt(8))
        _shade_cell(row[2], "FEF3C7")

    doc.add_heading("8.2 Équipements de protection", 2)
    for item in [
        "EPI : casque, chaussures S3, gilet HV, gants, lunettes, protections auditives",
        "Protections collectives : garde-corps, filets, balisage, signalisation, échafaudages",
        "Premiers secours : trousse, défibrillateur, SST formés (min. 2 par équipe)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.color.rgb = IA_BLUE

    doc.add_heading("8.3 Formations et habilitations", 2)
    _ia_text("Habilitations requises (à vérifier selon les lots) : électrique (B1V/B2V/BR/BC), CACES R386/R389/R482, travail en hauteur, AIPR, SST.")
    _client_text("Précisez votre taux de fréquence AT des 3 dernières années, engagement OPPBTP, certifications MASE le cas échéant.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  9. ENVIRONNEMENT, RSE ET DÉCHETS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Environnement, RSE et gestion des déchets", 1)

    doc.add_heading("9.1 Politique environnementale", 2)
    _ia_text("Nos engagements environnementaux pour ce chantier :")
    for item in [
        "Réduction de l'empreinte carbone : optimisation des rotations camions, matériaux locaux",
        "Matériaux biosourcés et circuits courts privilégiés quand techniquement possible",
        "Suivi des consommations eau/énergie avec objectifs chiffrés",
        "Limitation des nuisances : horaires respectés, arrosage poussières, bâchage",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.color.rgb = IA_BLUE

    doc.add_heading("9.2 SOGED — Gestion des déchets", 2)
    _ia_text("Schéma d'Organisation et de Gestion des Déchets de chantier :")
    dechets_table = _styled_table(["Type de déchet", "Filière (Proposition IA)", "Objectif", "Prestataire"])
    for dtype, filiere, obj, prestataire in [
        ("Inertes (béton, terre, gravats)", "Plateforme ISDI agréée", "> 70% valorisés", ""),
        ("Non dangereux (bois, plâtre, plastique)", "Centre de tri multi-flux", "> 50% valorisés", ""),
        ("Dangereux (amiante, peinture, solvants)", "Filière ISDD agréée", "100% traités", ""),
        ("Métaux (acier, cuivre, alu)", "Ferrailleur / recyclage", "> 90% recyclés", ""),
        ("Emballages", "Collecte sélective sur site", "> 80% recyclés", ""),
    ]:
        row = dechets_table.add_row().cells
        _set_cell_text(row[0], dtype, size=Pt(8))
        _set_cell_text(row[1], filiere, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[1], "EFF6FF")
        _set_cell_text(row[2], obj, color=IA_BLUE, size=Pt(8))
        _shade_cell(row[2], "EFF6FF")
        _set_cell_text(row[3], "À préciser", color=ORANGE, size=Pt(8))
        _shade_cell(row[3], "FEF3C7")

    doc.add_heading("9.3 Considérations sociales et insertion", 2)
    for item in [
        "Clause d'insertion : heures réservées à des publics éloignés de l'emploi",
        "Recours privilégié aux entreprises locales pour la sous-traitance",
        "Conditions de travail conformes aux conventions collectives BTP",
        "Formation et montée en compétence des compagnons",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.color.rgb = IA_BLUE
    _client_text("Complétez avec votre politique RSE concrète, certifications ISO 14001, actions insertion réalisées, label Vert.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  10. RÉFÉRENCES ET EXPÉRIENCES SIMILAIRES
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Références et expériences similaires", 1)
    _ia_text(
        "Les références ci-dessous doivent être similaires au présent marché en termes de : "
        "nature des travaux, montant, complexité technique, type d'acheteur public. "
        "Privilégiez les références récentes (< 5 ans) avec attestations de bonne exécution."
    )

    if po.get("market_type"):
        _add_encadre(
            f"Type de marché : {po['market_type']}. "
            "Sélectionnez des références correspondant à ce type de travaux.",
            bg_hex="EFF6FF", border_hex="2563EB"
        )

    ref_table = _styled_table(["Maître d'ouvrage", "Intitulé du marché", "Montant € HT", "Année", "Attestation"])
    for i in range(5):
        row = ref_table.add_row().cells
        for j, hint in enumerate(["Nom MOA", "Intitulé", "Montant", "Année", "Oui/Non"]):
            _set_cell_text(row[j], hint, color=ORANGE, size=Pt(8))
            _shade_cell(row[j], "FEF3C7")

    doc.add_paragraph()
    _add_encadre(
        "Conseils pour maximiser votre note :\n"
        "- Choisissez 3 à 5 références avec attestation de bonne exécution\n"
        "- Incluez des photos de chantier si possible\n"
        "- Montrez la similitude : même type de bâtiment, même complexité\n"
        "- Indiquez un contact vérifiable chez le maître d'ouvrage",
        bg_hex="ECFDF5", border_hex="10B981",
        text_color=GREEN_DARK
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  ANNEXE A — CHECKLIST
    # ══════════════════════════════════════════════════════════════════════
    if checklist_items:
        doc.add_heading("Annexe A — Checklist de conformité DCE", 1)

        elim_count = sum(1 for i in checklist_items if i.criticality and "liminatoire" in i.criticality.lower())
        manq_count = sum(1 for i in checklist_items if i.status == "MANQUANT")

        _add_encadre(
            f"{len(checklist_items)} exigences identifiées | {elim_count} éliminatoires | {manq_count} manquants\n"
            "Vérifiez chaque point avant soumission. Les éléments éliminatoires manquants entraînent le rejet automatique de l'offre.",
            bg_hex="FEF2F2", border_hex="EF4444",
            text_color=RGBColor(0x99, 0x1B, 0x1B)
        )

        cl_table = _styled_table(["#", "Exigence", "Catégorie", "Criticité", "Statut", "Action"])
        for idx, item in enumerate(checklist_items, 1):
            row = cl_table.add_row().cells
            _set_cell_text(row[0], str(idx), size=Pt(8))
            _set_cell_text(row[1], item.requirement or "—", size=Pt(8))
            _set_cell_text(row[2], item.category or "—", size=Pt(8))
            crit = item.criticality or "—"
            _set_cell_text(row[3], crit, bold=True, color=_severity_color(crit), size=Pt(8))
            status = item.status or "—"
            _set_cell_text(row[4], status, bold=True, color=_status_color(status), size=Pt(8))
            # IA suggestion for action
            if status == "MANQUANT":
                action = item.what_to_provide or "Fournir le document requis"
                _set_cell_text(row[5], action, color=RED, size=Pt(8))
                _shade_row(cl_table.rows[-1], "FEF2F2")
            elif status == "OK":
                _set_cell_text(row[5], "Conforme", color=GREEN_DARK, size=Pt(8))
            else:
                _set_cell_text(row[5], item.what_to_provide or "À vérifier", color=ORANGE, size=Pt(8))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  ANNEXE B — ANALYSE DES RISQUES (2 colonnes : IA + Client)
    # ══════════════════════════════════════════════════════════════════════
    if risks:
        doc.add_heading("Annexe B — Analyse des risques et mesures de mitigation", 1)

        # LLM plan d'action narratif
        if _memo_action_plan_text:
            _ia_text(_memo_action_plan_text)
            doc.add_paragraph()

        # Cashflow chart (si disponible)
        if _cashflow_buf:
            try:
                _cashflow_buf.seek(0)
                doc.add_picture(_cashflow_buf, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            except Exception as _e:
                logger.warning("memo_cashflow_insert_failed", error=str(_e))

        _add_encadre(
            "Les risques suivants ont été identifiés par l'IA lors de l'analyse du DCE. "
            "La colonne 'Proposition IA' contient une suggestion de mitigation. "
            "Complétez la colonne 'Mesure entreprise' avec votre réponse concrète.",
            bg_hex="EFF6FF", border_hex="2563EB"
        )

        risk_table = _styled_table(["Risque", "Sévérité", "Proposition IA", "Mesure entreprise"])
        # IA mitigation suggestions based on risk type
        def _ia_mitigation(risk_text, severity):
            r = (risk_text or "").lower()
            if any(w in r for w in ["delai", "retard", "planning"]):
                return "Prévoir marge de 10-15% sur le planning, identifier chemin critique"
            elif any(w in r for w in ["penalite", "pénalité"]):
                return "Chiffrer l'impact financier, négocier le plafonnement"
            elif any(w in r for w in ["prix", "financ", "cout", "coût", "budget"]):
                return "Vérifier la formule de révision, prévoir clause d'imprévision"
            elif any(w in r for w in ["sous-trait", "sous_trait"]):
                return "DC4 pré-rempli, agréments anticipés, paiement direct prévu"
            elif any(w in r for w in ["technique", "cctp", "norme"]):
                return "Consulter bureau d'études, demander précision à l'acheteur"
            elif any(w in r for w in ["assurance", "garantie", "caution"]):
                return "Vérifier couverture avec assureur, prévoir caution bancaire"
            return "Analyser l'impact et proposer une mesure adaptée"

        for risk in risks:
            row = risk_table.add_row().cells
            _set_cell_text(row[0], risk.get("risk", "—"), size=Pt(8))
            sev = risk.get("severity") or "—"
            sev_label = {"high": "FORT", "medium": "MOYEN", "low": "BAS"}.get(
                sev.lower() if isinstance(sev, str) else "", sev.upper() if isinstance(sev, str) else str(sev))
            _set_cell_text(row[1], sev_label, bold=True,
                          color=_severity_color(sev), size=Pt(8))
            # IA suggestion
            ia_mit = _ia_mitigation(risk.get("risk"), sev)
            _set_cell_text(row[2], ia_mit, color=IA_BLUE, size=Pt(8))
            _shade_cell(row[2], "EFF6FF")
            # Client to fill
            _set_cell_text(row[3], "Votre mesure concrète", color=ORANGE, size=Pt(8))
            _shade_cell(row[3], "FEF3C7")

    # ══════════════════════════════════════════════════════════════════════
    #  AVERTISSEMENT IA + FOOTER
    # ══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("Avertissement", 1)
    _add_encadre(
        "Ce document a été pré-rempli par intelligence artificielle (AO Copilot) "
        "à partir de l'analyse automatique des documents du DCE et des données "
        "de votre profil entreprise.\n\n"
        "Les éléments en bleu sont des propositions IA à valider et personnaliser.\n"
        "Les éléments en orange doivent être complétés avec vos données réelles.\n\n"
        "Ce document ne se substitue pas à l'expertise d'un rédacteur marchés publics. "
        "Vérifiez systématiquement que le contenu final correspond à vos capacités "
        "réelles et aux exigences spécifiques du RC.",
        bg_hex="FFFBEB", border_hex="FBBF24",
        text_color=RGBColor(0x78, 0x35, 0x0F)
    )

    doc.add_paragraph()
    doc.add_paragraph()
    footer_brand = doc.add_paragraph()
    footer_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fb = footer_brand.add_run("AO COPILOT")
    run_fb.bold = True
    run_fb.font.color.rgb = ACCENT_BLUE
    run_fb.font.size = Pt(11)

    gen_para = doc.add_paragraph()
    gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_gp = gen_para.add_run(
        f"aocopilot.fr — {datetime.now().strftime('%d/%m/%Y %H:%M')} — "
        "Document confidentiel"
    )
    run_gp.font.color.rgb = GRAY_LIGHT
    run_gp.font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
