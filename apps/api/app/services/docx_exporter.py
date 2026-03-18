"""Génération du rapport Word (.docx) complet — extrait de exporter.py.

Utilise fetch_export_data() pour centraliser les queries DB.
"""
import structlog
from datetime import datetime
from io import BytesIO
from sqlalchemy.orm import Session

from app.services.export_data import fetch_export_data

logger = structlog.get_logger(__name__)


def generate_export_docx(db: Session, project_id: str) -> bytes:
    """Génère un rapport Word (.docx) complet — design identique au PDF.
    Headers bleu foncé, encadrés colorés, badges, Go/No-Go box, disclaimer IA.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from io import BytesIO
    except ImportError as e:
        raise RuntimeError("python-docx non installé. Ajoutez python-docx==1.1.0.") from e

    # ── Charger TOUTES les données via le helper centralisé ──────────────
    data = fetch_export_data(db, project_id)
    project = data.project

    # Mapper les champs ExportData vers les variables locales
    # (certaines analyses retournent {} au lieu de None pour compatibilité)
    summary = data.summary or {}
    criteria = data.criteria or {}
    timeline = data.timeline or {}
    gonogo = data.gonogo or {}
    ccap = data.ccap_analysis or {}
    rc = data.rc_analysis or {}
    cctp = data.cctp_analysis or {}
    questions_data = {"questions": data.questions} if data.questions else {}
    scoring = data.scoring or {}
    cashflow = data.cashflow or {}
    subcontracting = data.subcontracting or {}
    ae = data.ae_analysis or {}
    dc_check = data.dc_check or {}
    conflicts_data = data.conflicts or {}
    dpgf_pricing = data.dpgf_pricing
    glossaire_btp = data.glossaire_btp or []
    checklist_items = data.checklist_items
    documents = data.documents

    po = summary.get("project_overview", {})

    doc = Document()

    # ── Couleurs PDF ─────────────────────────────────────────────────────
    DARK_BLUE = RGBColor(0x0F, 0x1B, 0x4C)
    ACCENT_BLUE = RGBColor(0x25, 0x63, 0xEB)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    RED = RGBColor(0xB9, 0x1C, 0x1C)
    GREEN_DARK = RGBColor(0x04, 0x78, 0x57)
    ORANGE = RGBColor(0xB4, 0x53, 0x09)
    GRAY = RGBColor(0x64, 0x74, 0x8B)
    GRAY_LIGHT = RGBColor(0x9C, 0xA3, 0xAF)

    # ── Style global ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)  # 11pt — BTP audit: lisibilité comités (50+ ans)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = DARK_BLUE
        hs.font.name = "Calibri"

    # ── Helper: cell shading ─────────────────────────────────────────────
    def _shade_cell(cell, hex_color):
        """Apply background color to a table cell."""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _shade_row(row, hex_color):
        for cell in row.cells:
            _shade_cell(cell, hex_color)

    def _set_cell_text(cell, text, bold=False, color=None, size=None):
        """Set cell text with formatting — clears default and adds a run."""
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = bold
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = size
        run.font.name = "Calibri"
        return run

    # ── Helper: styled table with dark blue header row ────────────────
    def _styled_table(headers, col_widths=None):
        """Create table with dark blue (#0F1B4C) header row, white text."""
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            _shade_cell(cell, "0F1B4C")
            _set_cell_text(cell, h, bold=True, color=WHITE, size=Pt(9))
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        return t

    # ── Helper: format date ──────────────────────────────────────────────
    def _fmt_date(val):
        if not val:
            return "—"
        s = str(val)
        if "T" in s:
            s = s.split("T")[0]
        try:
            from datetime import datetime as dt
            d = dt.strptime(s, "%Y-%m-%d")
            return d.strftime("%d/%m/%Y")
        except Exception:
            return s

    # ── Helper: colored encadré (info-box, warning-box, etc.) ────────
    def _add_encadre(text, bg_hex="EFF6FF", border_hex="2563EB", text_color=None):
        """Add a bordered paragraph mimicking PDF encadrés."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        # Background shading
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_hex)
        pPr.append(shd)
        # Left border
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "36")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), border_hex)
        pBdr.append(left)
        pPr.append(pBdr)
        # Indentation for padding effect
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "180")
        ind.set(qn("w:right"), "180")
        pPr.append(ind)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        if text_color:
            run.font.color.rgb = text_color
        return p

    # ── Helper: severity badge text ──────────────────────────────────────
    def _severity_color(sev):
        s = (sev or "").lower()
        if s in ("high", "fort", "haut", "éliminatoire", "eliminatoire"):
            return RED
        elif s in ("medium", "moyen", "moyenne"):
            return ORANGE
        elif s in ("low", "bas", "faible"):
            return GREEN_DARK
        return GRAY

    def _severity_bg(sev):
        s = (sev or "").lower()
        if s in ("high", "fort", "haut", "éliminatoire", "eliminatoire"):
            return "FEF2F2"
        elif s in ("medium", "moyen", "moyenne"):
            return "FFFBEB"
        return None

    def _status_color(status):
        s = (status or "").upper()
        if s == "OK":
            return GREEN_DARK
        elif s in ("MANQUANT", "NON", "KO"):
            return RED
        elif s in ("PARTIEL", "EN COURS"):
            return ORANGE
        return GRAY

    def _add_part_header(title, subtitle=""):
        """Add PARTIE visual separator (dark blue background, white text)."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0F1B4C")
        pPr.append(shd)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:right"), "0")
        pPr.append(ind)
        run_t = p.add_run(title)
        run_t.bold = True
        run_t.font.size = Pt(13)
        run_t.font.color.rgb = WHITE
        run_t.font.name = "Calibri"
        if subtitle:
            run_s = p.add_run(f"\n{subtitle}")
            run_s.font.size = Pt(9)
            run_s.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            run_s.font.name = "Calibri"

    # ══════════════════════════════════════════════════════════════════════
    #                        PAGE DE GARDE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    # Brand
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_b = brand.add_run("AO COPILOT")
    run_b.font.color.rgb = ACCENT_BLUE
    run_b.font.size = Pt(14)
    run_b.bold = True
    run_b.font.name = "Calibri"

    # Separator line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = line.add_run("━━━━━━━━━━")
    run_l.font.color.rgb = ACCENT_BLUE
    run_l.font.size = Pt(12)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sub.add_run("RAPPORT D'ANALYSE DCE")
    run_s.font.color.rgb = GRAY
    run_s.font.size = Pt(11)
    run_s.font.name = "Calibri"

    doc.add_paragraph()

    # Title
    title = doc.add_heading(project.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(24)

    # Buyer
    buyer_p = doc.add_paragraph()
    buyer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_buyer = buyer_p.add_run(po.get("buyer") or project.buyer or "Acheteur public")
    run_buyer.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    run_buyer.font.size = Pt(14)

    doc.add_paragraph()

    # Info table (cover)
    cover_info = _styled_table(["Information", "Détail"], [7, 10])
    for label, value in [
        ("Référence", project.reference or "N/A"),
        ("Lieu", po.get("location", "—")),
        ("Type de marché", po.get("market_type", "—")),
        ("Date limite", _fmt_date(po.get("deadline_submission"))),
        ("Budget estimé", po.get("estimated_budget") or "Non précisé"),
        ("Documents analysés", f"{len(documents)} documents"),
        ("Généré le", datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]:
        row = cover_info.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=Pt(9))
        _set_cell_text(row[1], str(value), size=Pt(9))

    doc.add_paragraph()

    # ── Go/No-Go badge on cover ──────────────────────────────────────────
    if gonogo:
        score = gonogo.get("score", 0)
        rec = (gonogo.get("recommendation") or "ATTENTION").upper()
        gonogo_p = doc.add_paragraph()
        gonogo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_rec = gonogo_p.add_run(f"RECOMMANDATION : {rec}  —  {score}/100")
        run_rec.bold = True
        run_rec.font.size = Pt(14)
        if rec == "GO":
            run_rec.font.color.rgb = GREEN_DARK
        elif rec == "NO-GO":
            run_rec.font.color.rgb = RED
        else:
            run_rec.font.color.rgb = ORANGE

    # Cover footer
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = meta_p.add_run(
        f"Généré le {datetime.now().strftime('%d/%m/%Y %H:%M')} | Rapport confidentiel"
    )
    run_meta.font.color.rgb = GRAY_LIGHT
    run_meta.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #                          SOMMAIRE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Sommaire", 1)

    ccag_derogations = ccap.get("ccag_derogations") or ccap.get("derogations", [])
    ccap_clauses_risquees = ccap.get("clauses_risquees", [])
    questions_list = questions_data.get("questions", [])

    # Sommaire structuré en 7 parties (miroir PDF)
    toc_items = [
        ("", "PARTIE I — DÉCISION STRATÉGIQUE"),
        ("1.", "Synthèse décisionnelle — Recommandation Go/No-Go et indicateurs clés"),
    ]
    if rc:
        toc_items.append(("2.", "Fiche signalétique du marché — Procédure, allotissement, groupement"))
    toc_items.append(("3.", "Résumé exécutif — Objet du marché, points clés extraits"))
    toc_items.append(("", "PARTIE II — CONFORMITÉ JURIDIQUE"))
    if dc_check:
        toc_items.append(("4.", "Vérification administrative DC — Attestations, certifications"))
    if checklist_items:
        toc_items.append(("5.", f"Checklist de conformité — {len(checklist_items)} exigences"))
    if ccag_derogations:
        toc_items.append(("6.", f"Dérogations CCAG-Travaux 2021 — {len(ccag_derogations)} détectées"))
    if ccap_clauses_risquees:
        toc_items.append(("7.", f"Clauses risquées CCAP — {len(ccap_clauses_risquees)} clauses"))
    toc_items.append(("", "PARTIE III — ANALYSE FINANCIÈRE"))
    if ae:
        toc_items.append(("8.", "Analyse Acte d'Engagement — Prix, pénalités, garanties"))
    toc_items.append(("9.", "Synthèse financière — Montants, avance, pénalités, révision"))
    toc_items.append(("10.", "Benchmark tarifaire DPGF — Prix unitaires vs références marché"))
    if cashflow:
        toc_items.append(("11.", "Simulation trésorerie et BFR — Impact financier prévisionnel"))
    toc_items.append(("", "PARTIE IV — ANALYSE TECHNIQUE"))
    if cctp:
        toc_items.append(("12.", "Analyse technique CCTP — Exigences, normes DTU, contradictions"))
    if conflicts_data:
        toc_items.append(("13.", f"Contradictions inter-documents — {conflicts_data.get('nb_total', 0)} incohérences"))
    if subcontracting:
        toc_items.append(("14.", "Analyse sous-traitance — Obligations, risques, conformité"))
    toc_items.append(("", "PARTIE V — SCORING ET STRATÉGIE"))
    if criteria:
        toc_items.append(("15.", "Critères d'attribution — Conditions d'éligibilité et grille de notation"))
    if scoring:
        toc_items.append(("16.", "Simulation note acheteur — Note estimée et leviers"))
    if gonogo and (gonogo.get("dimension_scores") or gonogo.get("breakdown")):
        toc_items.append(("17.", "Go/No-Go détaillé — 9 dimensions d'analyse"))
    toc_items.append(("", "PARTIE VI — PLAN D'ACTION"))
    toc_items.append(("18.", "Analyse des risques — Risques identifiés et plan d'actions 48h"))
    if questions_list:
        toc_items.append(("19.", f"Questions pour l'acheteur — {len(questions_list)} questions"))
    if timeline:
        toc_items.append(("20.", "Calendrier et dates clés — Échéances de soumission et d'exécution"))
    toc_items.append(("", "ANNEXES"))
    if documents:
        toc_items.append(("A1.", f"Inventaire des documents — {len(documents)} pièces DCE analysées"))
    toc_items.append(("A2.", "Glossaire BTP — Termes clés utilisés dans ce rapport"))
    toc_items.append(("A3.", "Avertissement IA et mentions légales"))

    for num, label in toc_items:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        if not num:
            # PARTIE header — dark blue background, white text
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "0F1B4C")
            pPr.append(shd)
            run_text = p.add_run(label)
            run_text.bold = True
            run_text.font.size = Pt(10)
            run_text.font.color.rgb = WHITE
            run_text.font.name = "Calibri"
        else:
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "dotted")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CBD5E1")
            pBdr.append(bottom)
            pPr.append(pBdr)
            run_num = p.add_run(f"{num}  ")
            run_num.bold = True
            run_num.font.color.rgb = ACCENT_BLUE
            run_num.font.size = Pt(10)
            run_text = p.add_run(label)
            run_text.font.size = Pt(10)
            run_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph()
    _add_encadre(
        "Guide de lecture rapide :\n"
        "5 min — Couverture + Synthèse décisionnelle (pages 1-2)\n"
        "15 min — + Résumé exécutif + Risques + Calendrier\n"
        "30 min — Rapport complet avec checklist et critères détaillés",
        bg_hex="EFF6FF", border_hex="3B82F6",
        text_color=RGBColor(0x1E, 0x3A, 0x8A)
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE I — DÉCISION STRATÉGIQUE
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE I — DÉCISION STRATÉGIQUE",
                     "Recommandation Go/No-Go, fiche marché et résumé exécutif")

    # ── 1. SYNTHÈSE DÉCISIONNELLE ──
    doc.add_heading("1. Synthèse décisionnelle", 1)
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("Vue d'ensemble en 1 page pour les décideurs (DG, Directeur commercial, Responsable AO)")
    run_sub.font.color.rgb = GRAY
    run_sub.font.size = Pt(9)

    if gonogo:
        score = gonogo.get("score", 0)
        rec = (gonogo.get("recommendation") or "ATTENTION").upper()

        # Go/No-Go box
        if rec == "GO":
            bg, border_c = "D1FAE5", "059669"
            txt_color = GREEN_DARK
        elif rec == "NO-GO":
            bg, border_c = "FDE8E8", "DC2626"
            txt_color = RED
        else:
            bg, border_c = "FEF3C7", "D97706"
            txt_color = ORANGE

        gonogo_box = doc.add_paragraph()
        gonogo_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = gonogo_box._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg)
        pPr.append(shd)
        pBdr = OxmlElement("w:pBdr")
        for side in ["top", "left", "bottom", "right"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "18")
            el.set(qn("w:space"), "4")
            el.set(qn("w:color"), border_c)
            pBdr.append(el)
        pPr.append(pBdr)

        run_score = gonogo_box.add_run(f"\n{score}/100\n")
        run_score.bold = True
        run_score.font.size = Pt(28)
        run_score.font.color.rgb = txt_color
        run_label = gonogo_box.add_run(f"{rec}\n")
        run_label.bold = True
        run_label.font.size = Pt(16)
        run_label.font.color.rgb = txt_color

        if gonogo.get("summary"):
            _add_encadre(gonogo["summary"], bg_hex="EFF6FF", border_hex="2563EB")

        # KPI stats (4 cells — matching PDF stat-grid)
        elim_count = sum(1 for i in checklist_items if i.criticality and "liminatoire" in i.criticality.lower()) if checklist_items else 0
        imp_count = sum(1 for i in checklist_items if (i.criticality or "").lower() == "important") if checklist_items else 0
        risk_count = len(summary.get("risks", []))
        actions_count = len(summary.get("actions_next_48h", []))

        stat_table = doc.add_table(rows=1, cols=4)
        stat_table.style = "Table Grid"
        stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        stats = [
            (str(elim_count), "Éliminatoires"),
            (str(imp_count), "Importants"),
            (str(risk_count), "Risques identifiés"),
            (str(actions_count), "Actions à mener"),
        ]
        for i, (val, label) in enumerate(stats):
            cell = stat_table.rows[0].cells[i]
            _shade_cell(cell, "F8FAFC")
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_v = p.add_run(f"{val}\n")
            run_v.bold = True
            run_v.font.size = Pt(20)
            run_v.font.color.rgb = DARK_BLUE
            run_lb = p.add_run(label)
            run_lb.font.size = Pt(8)
            run_lb.font.color.rgb = GRAY

    # ── Top 3 risques critiques (matching PDF) ──
    risks = summary.get("risks", [])
    if risks:
        doc.add_heading("Top 3 risques critiques", 3)
        top3_table = _styled_table(["Risque", "Sévérité", "Impact"])
        for r in risks[:3]:
            row = top3_table.add_row().cells
            _set_cell_text(row[0], (r.get("risk") or "")[:50], bold=True, size=Pt(9))
            sev = r.get("severity") or ""
            sev_label = {"high": "FORT", "medium": "MOYEN", "low": "BAS"}.get(sev.lower(), sev.upper())
            _set_cell_text(row[1], sev_label, bold=True, color=_severity_color(sev), size=Pt(9))
            bg_r = _severity_bg(sev)
            if bg_r:
                _shade_cell(row[1], bg_r)
            _set_cell_text(row[2], (r.get("why") or "")[:120], size=Pt(8))

    # ── Forces / Points de vigilance (2-column table — matching PDF) ──
    if gonogo and (gonogo.get("strengths") or gonogo.get("risks")):
        sf_table = doc.add_table(rows=1, cols=2)
        sf_table.style = "Table Grid"
        sf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header: Forces | Points de vigilance
        _shade_cell(sf_table.rows[0].cells[0], "D1FAE5")
        _set_cell_text(sf_table.rows[0].cells[0], "Forces", bold=True,
                      color=GREEN_DARK, size=Pt(10))
        _shade_cell(sf_table.rows[0].cells[1], "FEF2F2")
        _set_cell_text(sf_table.rows[0].cells[1], "Points de vigilance", bold=True,
                      color=RED, size=Pt(10))
        strengths = (gonogo.get("strengths") or [])[:3]
        gng_risks = (gonogo.get("risks") or [])[:3]
        max_rows = max(len(strengths), len(gng_risks))
        for idx in range(max_rows):
            row = sf_table.add_row().cells
            if idx < len(strengths):
                _set_cell_text(row[0], f"+ {strengths[idx]}", size=Pt(9),
                              color=GREEN_DARK)
            if idx < len(gng_risks):
                _set_cell_text(row[1], f"- {gng_risks[idx]}", size=Pt(9),
                              color=RED)

    # ── Dimensions Go/No-Go breakdown (matching PDF) ──
    if gonogo and gonogo.get("breakdown"):
        doc.add_heading("Scores par dimension", 3)
        dim_table = _styled_table(["Dimension", "Score", "Niveau"])
        breakdown = gonogo["breakdown"]
        dim_mapping = [
            ("Adéquation technique", breakdown.get("technical_fit")),
            ("Capacité financière", breakdown.get("financial_capacity")),
            ("Faisabilité planning", breakdown.get("timeline_feasibility")),
            ("Position concurrentielle", breakdown.get("competitive_position")),
        ]
        for dim_name, dim_score in dim_mapping:
            if dim_score is not None:
                row = dim_table.add_row().cells
                _set_cell_text(row[0], dim_name, size=Pt(9))
                _set_cell_text(row[1], f"{dim_score}/100", bold=True, size=Pt(9))
                if dim_score >= 70:
                    level_label, level_color, level_bg = "BON", GREEN_DARK, "D1FAE5"
                elif dim_score >= 50:
                    level_label, level_color, level_bg = "MOYEN", ORANGE, "FEF3C7"
                else:
                    level_label, level_color, level_bg = "FAIBLE", RED, "FEF2F2"
                _set_cell_text(row[2], level_label, bold=True, color=level_color, size=Pt(9))
                _shade_cell(row[2], level_bg)

    # ── Actions prioritaires P0 (filtered — matching PDF) ──
    actions_p0 = [a for a in summary.get("actions_next_48h", []) if a.get("priority") == "P0"]
    if actions_p0:
        doc.add_heading("Actions prioritaires P0", 3)
        p0_table = _styled_table(["Action", "Responsable"])
        for a in actions_p0:
            row = p0_table.add_row().cells
            _set_cell_text(row[0], (a.get("action") or "")[:100], size=Pt(8))
            _set_cell_text(row[1], (a.get("owner_role") or "")[:30], size=Pt(8))

    # ── Confiance IA (matching PDF) ──
    confidence = summary.get("confidence") or summary.get("avg_confidence")
    if confidence:
        conf_p = doc.add_paragraph()
        run_cl = conf_p.add_run(f"Indice de confiance IA : {confidence * 100:.0f}%")
        run_cl.font.size = Pt(10)
        run_cl.font.color.rgb = GRAY
        # Visual confidence bar as a 1-row table
        bar_table = doc.add_table(rows=1, cols=1)
        bar_table.style = "Table Grid"
        cell = bar_table.rows[0].cells[0]
        if confidence >= 0.8:
            bar_color = "059669"
        elif confidence >= 0.6:
            bar_color = "D97706"
        else:
            bar_color = "DC2626"
        _shade_cell(cell, bar_color)
        pct_text = f"{confidence * 100:.0f}%"
        _set_cell_text(cell, pct_text, bold=True, color=WHITE, size=Pt(9))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set width proportionally
        cell.width = Cm(confidence * 17)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §2. FICHE SIGNALÉTIQUE DU MARCHÉ (RC)
    # ══════════════════════════════════════════════════════════════════════
    if rc:
        doc.add_heading("2. Fiche signalétique du marché", 1)
        p_sub_rc = doc.add_paragraph()
        run_sub_rc = p_sub_rc.add_run("Données extraites du Règlement de Consultation (RC) et des pièces contractuelles")
        run_sub_rc.font.color.rgb = GRAY
        run_sub_rc.font.size = Pt(9)
        rc_table = _styled_table(["Rubrique", "Détail"])
        rc_fields = [
            ("Procédure", rc.get("procedure_type")),
            ("Allotissement", rc.get("allotissement")),
            ("Groupement", rc.get("groupement")),
            ("Variantes", rc.get("variantes")),
            ("Sous-traitance", "Autorisée" if rc.get("subcontracting_allowed") else "Non autorisée"),
            ("CCAG de référence", rc.get("ccag_reference")),
            ("Visite obligatoire", "Oui" if rc.get("visite_obligatoire") else "Non"),
            ("DUME requis", "Oui" if rc.get("dume_required") else "Non"),
        ]
        for label, value in rc_fields:
            if value:
                row = rc_table.add_row().cells
                _set_cell_text(row[0], label, bold=True, size=Pt(9))
                _set_cell_text(row[1], str(value), size=Pt(9))

        lots = rc.get("lots", [])
        if lots:
            doc.add_heading("Lots", 2)
            lots_table = _styled_table(["N°", "Intitulé", "Montant estimé"])
            for lot in lots:
                row = lots_table.add_row().cells
                _set_cell_text(row[0], str(lot.get("number") or lot.get("lot_number") or "—"), size=Pt(9))
                _set_cell_text(row[1], lot.get("title") or lot.get("label") or "—", size=Pt(9))
                _set_cell_text(row[2], str(lot.get("estimated_amount") or "—"), size=Pt(9))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §3. RÉSUMÉ EXÉCUTIF
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Résumé exécutif", 1)

    # Summary box — structured like PDF (Objet/Acheteur/Lieu/Date/Budget)
    summary_fields = []
    if po.get("scope") or po.get("object"):
        summary_fields.append(("Objet", po.get("scope") or po.get("object")))
    if po.get("buyer"):
        summary_fields.append(("Acheteur", po["buyer"]))
    if po.get("location"):
        summary_fields.append(("Lieu", po["location"]))
    if po.get("deadline_submission"):
        summary_fields.append(("Date limite", _fmt_date(po["deadline_submission"])))
    if po.get("estimated_budget"):
        summary_fields.append(("Budget estimé", po["estimated_budget"]))
    if po.get("market_type"):
        summary_fields.append(("Type de marché", po["market_type"]))

    if summary_fields:
        box_text = "\n".join(f"{label} : {value}" for label, value in summary_fields)
        _add_encadre(box_text, bg_hex="EFF6FF", border_hex="2563EB")

    # Points clés
    key_points = summary.get("key_points", [])
    if key_points:
        doc.add_heading("Points clés du DCE", 2)
        kp_table = _styled_table(["Priorité", "Point clé"])
        for kp in key_points:
            row = kp_table.add_row().cells
            importance = (kp.get("importance") or "medium").lower()
            label = {"high": "HAUT", "medium": "MOYEN", "low": "BAS"}.get(importance, importance.upper())
            _set_cell_text(row[0], label, bold=True, color=_severity_color(importance), size=Pt(9))
            bg = _severity_bg(importance)
            if bg:
                _shade_cell(row[0], bg)
            _set_cell_text(row[1], kp.get("point", "") or kp.get("label", ""), size=Pt(9))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE II — CONFORMITÉ JURIDIQUE
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE II — CONFORMITÉ JURIDIQUE",
                     "Vérification administrative, conformité, dérogations CCAG")

    # ══════════════════════════════════════════════════════════════════════
    #  §4. VÉRIFICATION ADMINISTRATIVE DC
    # ══════════════════════════════════════════════════════════════════════
    if dc_check:
        doc.add_heading("4. Vérification administrative DC", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Vérification des pièces administratives — DC1, DC2, attestations, certifications")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        dc_docs = dc_check.get("documents", [])
        if dc_docs:
            dc_table = _styled_table(["Document", "Statut", "Date validité", "Alerte"])
            for d in dc_docs:
                row = dc_table.add_row().cells
                _set_cell_text(row[0], d.get("name") or d.get("document") or "—", bold=True, size=Pt(9))
                status = d.get("status") or "—"
                _set_cell_text(row[1], status, bold=True, color=_status_color(status), size=Pt(9))
                bg = "D1FAE5" if status == "OK" else ("FEF2F2" if status in ("MANQUANT", "KO") else None)
                if bg:
                    _shade_cell(row[1], bg)
                _set_cell_text(row[2], _fmt_date(d.get("expiry_date") or d.get("validity_date")), size=Pt(9))
                alert = d.get("alert") or ""
                _set_cell_text(row[3], alert[:80], color=RED if alert else GRAY, size=Pt(8))
                if status in ("MANQUANT", "KO"):
                    _shade_row(dc_table.rows[-1], "FEF2F2")

        if dc_check.get("summary") or dc_check.get("resume"):
            _add_encadre(dc_check.get("summary") or dc_check.get("resume"), bg_hex="EFF6FF", border_hex="2563EB")

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §5. CHECKLIST DE CONFORMITÉ
    # ══════════════════════════════════════════════════════════════════════
    if checklist_items:
        elim_count = sum(1 for i in checklist_items if i.criticality and "liminatoire" in i.criticality.lower())
        manq_count = sum(1 for i in checklist_items if i.status == "MANQUANT")
        ok_count = sum(1 for i in checklist_items if i.status == "OK")

        doc.add_heading(f"5. Checklist de conformité ({len(checklist_items)} exigences)", 1)

        # Stats summary box
        _add_encadre(
            f"{elim_count} éliminatoires  |  {manq_count} manquants  |  {ok_count} conformes",
            bg_hex="EFF6FF", border_hex="2563EB"
        )

        cl_table = _styled_table(["#", "Exigence", "Catégorie", "Criticité", "Statut", "À fournir"])
        for idx, item in enumerate(checklist_items, 1):
            row = cl_table.add_row().cells
            _set_cell_text(row[0], str(idx), size=Pt(8))
            _set_cell_text(row[1], item.requirement or "", size=Pt(8))
            _set_cell_text(row[2], item.category or "—", size=Pt(8))
            crit = item.criticality or "—"
            _set_cell_text(row[3], crit, bold=True, color=_severity_color(crit), size=Pt(8))
            status = item.status or "—"
            _set_cell_text(row[4], status, bold=True, color=_status_color(status), size=Pt(8))
            _set_cell_text(row[5], item.what_to_provide or "—", size=Pt(8))
            # Row background for eliminatoire manquant
            if crit and "liminatoire" in crit.lower() and status == "MANQUANT":
                _shade_row(cl_table.rows[-1], "FEF2F2")

        # ── Section 5b: Documents prioritaires éliminatoires MANQUANTS ──
        elim_manquants = [i for i in checklist_items
                          if i.criticality and "liminatoire" in i.criticality.lower()
                          and i.status == "MANQUANT"]
        if elim_manquants:
            doc.add_page_break()
            h_5b = doc.add_heading(f"5b. Documents prioritaires à préparer ({len(elim_manquants)} éliminatoires)", 2)
            for run in h_5b.runs:
                run.font.color.rgb = RED

            _add_encadre(
                f"Attention : Ces {len(elim_manquants)} documents sont éliminatoires. "
                "Leur absence entraînera le rejet automatique de votre candidature. "
                "Préparez-les en priorité absolue.",
                bg_hex="FEF2F2", border_hex="EF4444",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

            elim_table = _styled_table(["#", "Document / Justificatif requis", "Détail à fournir"])
            for idx_e, item in enumerate(elim_manquants, 1):
                row = elim_table.add_row().cells
                _set_cell_text(row[0], str(idx_e), size=Pt(8))
                _set_cell_text(row[1], (item.requirement or "")[:80], bold=True, size=Pt(8))
                _set_cell_text(row[2], (item.what_to_provide or "Voir exigences du RC/CCAP")[:100], size=Pt(8))
                _shade_row(elim_table.rows[-1], "FEF2F2")

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §6. DÉROGATIONS CCAG-TRAVAUX 2021
    # ══════════════════════════════════════════════════════════════════════
    if ccag_derogations:
        doc.add_heading("6. Dérogations CCAG-Travaux 2021", 1)
        _add_encadre(
            f"{len(ccag_derogations)} dérogation(s) au CCAG-Travaux 2021 détectée(s). "
            "Certaines peuvent avoir un impact significatif sur votre prix et vos risques contractuels.",
            bg_hex="FEF2F2", border_hex="EF4444",
            text_color=RGBColor(0x99, 0x1B, 0x1B)
        )

        d_table = _styled_table(["Article CCAG", "Dérogation CCAP", "Gravité", "Évaluation"])
        for d in ccag_derogations:
            row = d_table.add_row().cells
            _set_cell_text(row[0], d.get("article_ccag") or d.get("article") or "—",
                          bold=True, size=Pt(9))
            _set_cell_text(row[1], d.get("derogation", "—"), size=Pt(9))
            sev = d.get("severity") or d.get("impact") or "—"
            sev_label = {"high": "FORT", "medium": "MOYEN", "low": "BAS",
                        "fort": "FORT", "moyen": "MOYEN"}.get(sev.lower() if isinstance(sev, str) else sev, sev.upper() if isinstance(sev, str) else str(sev))
            _set_cell_text(row[2], sev_label, bold=True,
                          color=_severity_color(sev), size=Pt(9))
            bg = _severity_bg(sev)
            if bg:
                _shade_cell(row[2], bg)
            _set_cell_text(row[3], d.get("evaluation") or d.get("risk_comment") or "—", size=Pt(9))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §7. CLAUSES RISQUÉES CCAP
    # ══════════════════════════════════════════════════════════════════════
    clauses_risquees = ccap.get("clauses_risquees", [])
    if clauses_risquees:
        doc.add_heading("7. Clauses risquées CCAP", 1)
        cr_table = _styled_table(["Clause", "Risque", "Sévérité", "Recommandation"])
        for cl in clauses_risquees[:15]:
            row = cr_table.add_row().cells
            _set_cell_text(row[0], cl.get("clause") or cl.get("article") or "—", bold=True, size=Pt(9))
            _set_cell_text(row[1], cl.get("risk") or cl.get("risk_description") or "—", size=Pt(9))
            sev = cl.get("severity") or "—"
            _set_cell_text(row[2], sev.upper() if isinstance(sev, str) else str(sev),
                          bold=True, color=_severity_color(sev), size=Pt(9))
            _set_cell_text(row[3], cl.get("recommendation") or cl.get("mitigation") or "—", size=Pt(9))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE III — ANALYSE FINANCIÈRE
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE III — ANALYSE FINANCIÈRE",
                     "Acte d'engagement, synthèse financière, trésorerie")

    # ══════════════════════════════════════════════════════════════════════
    #  §8. ANALYSE ACTE D'ENGAGEMENT
    # ══════════════════════════════════════════════════════════════════════
    if ae:
        doc.add_heading("8. Analyse Acte d'Engagement", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Conditions contractuelles extraites de l'Acte d'Engagement")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        ae_fields = [
            ("Type de prix", ae.get("type_prix")),
            ("Forme du prix", ae.get("forme_prix")),
            ("Révision des prix", ae.get("revision_prix")),
            ("Indice de révision", ae.get("indice_revision")),
            ("Durée du marché", ae.get("duree_marche")),
            ("Reconduction", ae.get("reconduction")),
            ("Pénalités de retard", ae.get("penalites_retard")),
            ("Retenue de garantie", f"{ae.get('retenue_garantie_pct')}%" if ae.get("retenue_garantie_pct") else None),
            ("Avance", f"{ae.get('avance_pct')}%" if ae.get("avance_pct") else None),
            ("Délai de paiement", ae.get("delai_paiement")),
            ("Délai global d'exécution", ae.get("delai_global")),
        ]
        ae_table = _styled_table(["Rubrique", "Valeur"])
        for label, value in ae_fields:
            if value:
                row = ae_table.add_row().cells
                _set_cell_text(row[0], label, bold=True, size=Pt(9))
                _set_cell_text(row[1], str(value), size=Pt(9))

        ae_clauses = ae.get("clauses_risquees", [])
        if ae_clauses:
            doc.add_heading("Clauses risquées de l'AE", 2)
            acl_table = _styled_table(["Clause", "Risque", "Sévérité", "Conseil"])
            for cl in ae_clauses[:10]:
                row = acl_table.add_row().cells
                _set_cell_text(row[0], cl.get("clause") or cl.get("article") or "—", bold=True, size=Pt(9))
                _set_cell_text(row[1], cl.get("risk") or cl.get("description") or "—", size=Pt(9))
                sev = cl.get("severity") or "—"
                _set_cell_text(row[2], sev.upper() if isinstance(sev, str) else str(sev),
                              bold=True, color=_severity_color(sev), size=Pt(9))
                bg = _severity_bg(sev)
                if bg:
                    _shade_cell(row[2], bg)
                _set_cell_text(row[3], cl.get("conseil") or cl.get("recommendation") or "—", size=Pt(8))

        score_ae = ae.get("score_risque")
        if score_ae is not None:
            _add_encadre(
                f"Score de risque AE : {score_ae}/100",
                bg_hex="FEF2F2" if score_ae >= 60 else ("FEF3C7" if score_ae >= 30 else "D1FAE5"),
                border_hex="EF4444" if score_ae >= 60 else ("D97706" if score_ae >= 30 else "059669"),
                text_color=RGBColor(0x99, 0x1B, 0x1B) if score_ae >= 60 else (ORANGE if score_ae >= 30 else GREEN_DARK)
            )

        if ae.get("resume"):
            _add_encadre(ae["resume"], bg_hex="EFF6FF", border_hex="2563EB")

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §9. SYNTHÈSE FINANCIÈRE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Synthèse financière", 1)
    p_sub_fin = doc.add_paragraph()
    run_sub_fin = p_sub_fin.add_run("Éléments financiers clés extraits du DCE pour l'aide à la décision")
    run_sub_fin.font.color.rgb = GRAY
    run_sub_fin.font.size = Pt(9)

    _add_encadre(
        f"Budget global estimé : {po.get('estimated_budget') or 'Non précisé dans le DCE'}\n"
        f"Type de prix : {po.get('market_type') or 'Non précisé'}",
        bg_hex="ECFDF5", border_hex="10B981",
        text_color=GREEN_DARK
    )

    # Points clés financiers
    fin_points = [kp for kp in summary.get("key_points", [])
                  if any(w in (kp.get("point") or "").lower()
                         for w in ["prix", "avance", "retenue", "paiement", "penalite",
                                   "revision", "financ", "budget", "garantie", "caution"])]
    if fin_points:
        doc.add_heading("Éléments financiers extraits", 2)
        fin_table = _styled_table(["Point financier"])
        for kp in fin_points:
            row = fin_table.add_row().cells
            _set_cell_text(row[0], (kp.get("point") or "")[:200], size=Pt(9))

    # Risques financiers
    fin_risks = [r for r in summary.get("risks", [])
                 if any(w in (r.get("why") or r.get("risk") or "").lower()
                        for w in ["financ", "prix", "penalite", "cout", "tresorerie", "paiement"])]
    if fin_risks:
        doc.add_heading("Risques financiers identifiés", 2)
        fr_table = _styled_table(["Risque", "Sévérité", "Impact financier"])
        for r in fin_risks:
            row = fr_table.add_row().cells
            _set_cell_text(row[0], (r.get("risk") or "")[:50], bold=True, size=Pt(9))
            sev = r.get("severity") or ""
            _set_cell_text(row[1], sev.upper(), bold=True,
                          color=_severity_color(sev), size=Pt(9))
            _set_cell_text(row[2], (r.get("why") or "")[:150], size=Pt(8))

    # Recommandation warning box (matching PDF)
    _add_encadre(
        "Recommandation : Avant de chiffrer votre offre, vérifiez les éléments suivants : "
        "formule de révision des prix, montant de l'avance, conditions de paiement, pénalités de retard, "
        "retenue de garantie, et tout risque de surcoût identifié (pollution, aléas géotechniques, etc.).",
        bg_hex="FEF2F2", border_hex="EF4444",
        text_color=RGBColor(0x99, 0x1B, 0x1B)
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §10. BENCHMARK TARIFAIRE DPGF
    # ══════════════════════════════════════════════════════════════════════
    if dpgf_pricing:
        doc.add_heading("10. Benchmark tarifaire DPGF", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Comparaison des prix unitaires avec les références régionales du marché BTP")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        sous = sum(1 for l in dpgf_pricing if (l.get("status") or "") == "SOUS_EVALUE")
        sur = sum(1 for l in dpgf_pricing if (l.get("status") or "") == "SUR_EVALUE")
        normal = sum(1 for l in dpgf_pricing if (l.get("status") or "") == "NORMAL")

        stat_table = doc.add_table(rows=1, cols=4)
        stat_table.style = "Table Grid"
        stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (val, label, bg_color) in enumerate([
            (str(len(dpgf_pricing)), "Postes analysés", "EFF6FF"),
            (str(sous), "Sous-évalués", "FEF2F2"),
            (str(sur), "Sur-évalués", "FEF3C7"),
            (str(normal), "Normaux", "D1FAE5"),
        ]):
            cell = stat_table.rows[0].cells[i]
            _shade_cell(cell, bg_color)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_v = p.add_run(f"{val}\n")
            run_v.bold = True
            run_v.font.size = Pt(16)
            run_v.font.color.rgb = DARK_BLUE
            run_lb = p.add_run(label)
            run_lb.font.size = Pt(8)
            run_lb.font.color.rgb = GRAY

        doc.add_paragraph()

        dp_table = _styled_table(["Désignation", "Prix unit.", "Réf. min", "Réf. moy", "Réf. max", "Statut"])
        for line in dpgf_pricing[:20]:
            if isinstance(line, dict):
                row = dp_table.add_row().cells
                _set_cell_text(row[0], (line.get("designation") or "—")[:40], size=Pt(8))
                _set_cell_text(row[1], f"{line.get('prix_unitaire', '—')} €", size=Pt(8))
                ref = line.get("reference_match") or {}
                if isinstance(ref, dict):
                    _set_cell_text(row[2], f"{ref.get('prix_min', '—')} €", size=Pt(8))
                    _set_cell_text(row[3], f"{ref.get('prix_moyen', '—')} €", size=Pt(8))
                    _set_cell_text(row[4], f"{ref.get('prix_max', '—')} €", size=Pt(8))
                else:
                    for j in range(2, 5):
                        _set_cell_text(row[j], "—", size=Pt(8))
                status = line.get("status") or "INCONNU"
                status_colors = {
                    "SOUS_EVALUE": (RED, "FEF2F2"),
                    "SUR_EVALUE": (ORANGE, "FEF3C7"),
                    "NORMAL": (GREEN_DARK, "D1FAE5"),
                }
                s_color, s_bg = status_colors.get(status, (GRAY, None))
                _set_cell_text(row[5], status.replace("_", " "), bold=True, color=s_color, size=Pt(8))
                if s_bg:
                    _shade_cell(row[5], s_bg)

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §11. SIMULATION TRÉSORERIE ET BFR
    # ══════════════════════════════════════════════════════════════════════
    if cashflow:
        doc.add_heading("11. Simulation trésorerie et BFR", 1)
        p_sub_cf = doc.add_paragraph()
        run_sub_cf = p_sub_cf.add_run("Impact prévisionnel sur la trésorerie de l'entreprise — aide à la décision financière")
        run_sub_cf.font.color.rgb = GRAY
        run_sub_cf.font.size = Pt(9)

        # BFR peak financial box — support both key formats (bfr_peak or bfr_eur)
        bfr_peak = cashflow.get("bfr_peak") or cashflow.get("bfr_eur")
        if bfr_peak:
            bfr_val = float(bfr_peak)
            bfr_text = f"BFR maximal estimé : {abs(bfr_val):,.0f} €".replace(",", " ")
            bfr_month = cashflow.get("bfr_month")
            if bfr_month:
                bfr_text += f" — atteint au mois {bfr_month}"
            margin = cashflow.get("margin_estimate") or cashflow.get("marge_brute_pct")
            if margin:
                bfr_text += f"\nMarge estimée : {float(margin):.1f}%"
            montant = cashflow.get("montant_total_ht")
            if montant:
                bfr_text += f"\nMontant total HT : {float(montant):,.0f} €".replace(",", " ")
            _add_encadre(bfr_text, bg_hex="ECFDF5", border_hex="10B981", text_color=GREEN_DARK)

        # Résumé cashflow
        if cashflow.get("resume"):
            _add_encadre(cashflow["resume"], bg_hex="EFF6FF", border_hex="2563EB")

        # Monthly cashflow table — support both key formats
        monthly = cashflow.get("monthly_cashflow", [])
        if monthly:
            doc.add_heading("Trésorerie mensuelle prévisionnelle", 3)
            mcf_table = _styled_table(["Mois", "Dépenses HT", "Encaissements HT", "Solde cumulé"])
            for m in monthly[:15]:
                row = mcf_table.add_row().cells
                month_num = m.get("month") or m.get("mois", "?")
                month_label = m.get("label") or f"M{month_num}"
                _set_cell_text(row[0], month_label, size=Pt(8))
                expenses = m.get("expenses") or m.get("depenses_ht") or 0
                _set_cell_text(row[1], f"-{float(expenses):,.0f} €".replace(",", " "),
                              color=RED, size=Pt(8))
                income = m.get("income") or m.get("encaissement_ht") or 0
                _set_cell_text(row[2], f"+{float(income):,.0f} €".replace(",", " "),
                              color=GREEN_DARK, size=Pt(8))
                cumulative = m.get("cumulative") or m.get("solde_cumule") or 0
                cum_val = float(cumulative)
                cum_color = RED if cum_val < 0 else GREEN_DARK
                _set_cell_text(row[3], f"{cum_val:,.0f} €".replace(",", " "),
                              bold=True, color=cum_color, size=Pt(8))
                if cum_val < 0:
                    _shade_cell(row[3], "FEF2F2")

        # Risk level
        risk_level = cashflow.get("risk_level")
        if risk_level:
            _add_encadre(
                f"Niveau de risque trésorerie : {risk_level}",
                bg_hex="FEF2F2" if "LEV" in (risk_level or "").upper() else "FEF3C7",
                border_hex="EF4444" if "LEV" in (risk_level or "").upper() else "D97706",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

        # Tension months
        tension = cashflow.get("tension_months", [])
        if tension:
            _add_encadre(
                f"Mois en tension ({len(tension)}) : M" + ", M".join(str(t) for t in tension) +
                "\nPrévoir un financement de trésorerie ou ligne de crédit.",
                bg_hex="FEF2F2", border_hex="EF4444",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

        warnings = cashflow.get("warnings", [])
        if warnings:
            doc.add_heading("Alertes trésorerie", 2)
            for w in warnings[:5]:
                w_text = w.get("message") if isinstance(w, dict) else str(w)
                _add_encadre(w_text, bg_hex="FEF2F2", border_hex="EF4444",
                            text_color=RGBColor(0x99, 0x1B, 0x1B))

        # Hypothèses box
        avance_eur = cashflow.get("avance_impact_eur")
        retenue_eur = cashflow.get("retenue_impact_eur")
        duree = cashflow.get("duree_mois")
        hyp_parts = ["Hypothèses : Délai de paiement acheteur 30j"]
        if avance_eur:
            hyp_parts.append(f"avance forfaitaire {float(avance_eur):,.0f} €".replace(",", " "))
        if retenue_eur:
            hyp_parts.append(f"retenue de garantie {float(retenue_eur):,.0f} €".replace(",", " "))
        if duree:
            hyp_parts.append(f"durée {duree} mois")
        hyp_parts.append("Ajustez selon les termes du CCAP.")
        _add_encadre(
            ", ".join(hyp_parts),
            bg_hex="EFF6FF", border_hex="3B82F6",
            text_color=RGBColor(0x1E, 0x3A, 0x8A)
        )

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE IV — ANALYSE TECHNIQUE
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE IV — ANALYSE TECHNIQUE",
                     "CCTP, contradictions inter-documents, sous-traitance")

    # ══════════════════════════════════════════════════════════════════════
    #  §12. ANALYSE TECHNIQUE CCTP
    # ══════════════════════════════════════════════════════════════════════
    if cctp:
        doc.add_heading("12. Analyse technique CCTP", 1)
        p_sub_cctp = doc.add_paragraph()
        run_sub_cctp = p_sub_cctp.add_run("Synthèse des exigences techniques extraites du Cahier des Clauses Techniques Particulières")
        run_sub_cctp.font.color.rgb = GRAY
        run_sub_cctp.font.size = Pt(9)

        if cctp.get("technical_summary"):
            _add_encadre(f"Synthèse technique : {cctp['technical_summary']}", bg_hex="EFF6FF", border_hex="2563EB")

        categories = cctp.get("categories", [])
        if categories:
            for cat in categories:
                cat_name = cat.get("name") or cat.get("category") or "—"
                # Category heading with risk level badge (matching PDF)
                h_cat = doc.add_heading(cat_name, 3)
                risk_level = cat.get("risk_level")
                if risk_level:
                    run_rl = h_cat.add_run(f"  [{risk_level.upper()}]")
                    run_rl.font.size = Pt(8)
                    run_rl.font.color.rgb = _severity_color(risk_level)

                # Items as table with Norme/DTU and Risque columns (matching PDF)
                items = cat.get("items") or cat.get("requirements") or []
                if items:
                    cat_table = _styled_table(["Exigence", "Détail", "Norme/DTU", "Risque"])
                    for item in items[:8]:
                        if isinstance(item, dict):
                            row = cat_table.add_row().cells
                            _set_cell_text(row[0], (item.get("requirement") or item.get("label") or "")[:60],
                                          bold=True, size=Pt(8))
                            _set_cell_text(row[1], (item.get("detail") or item.get("value") or item.get("description") or "")[:100],
                                          size=Pt(8))
                            _set_cell_text(row[2], (item.get("norm") or item.get("standard") or "")[:30],
                                          size=Pt(8))
                            item_risk = item.get("risk", "")
                            risk_label = (item_risk.upper() if item_risk else "INFO")
                            _set_cell_text(row[3], risk_label, bold=True,
                                          color=_severity_color(item_risk), size=Pt(8))
                        else:
                            row = cat_table.add_row().cells
                            _set_cell_text(row[0], str(item)[:60], size=Pt(8))

        # Contradictions as 3-column table (matching PDF: Contradiction / Références / Recommandation)
        contradictions = cctp.get("contradictions", [])
        if contradictions:
            h_contra = doc.add_heading("Contradictions détectées dans le CCTP", 2)
            for run in h_contra.runs:
                run.font.color.rgb = RED

            _add_encadre(
                f"Attention : {len(contradictions)} contradiction(s) interne(s) détectée(s) dans le CCTP. "
                "Posez la question à l'acheteur avant la date limite.",
                bg_hex="FEF2F2", border_hex="EF4444",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

            contra_table = _styled_table(["Contradiction", "Références", "Recommandation"])
            for c in contradictions[:5]:
                if isinstance(c, dict):
                    row = contra_table.add_row().cells
                    _set_cell_text(row[0], (c.get("description") or c.get("issue") or c.get("contradiction") or "")[:120],
                                  size=Pt(8))
                    _set_cell_text(row[1], (c.get("references") or c.get("source") or "")[:60], size=Pt(8))
                    _set_cell_text(row[2], (c.get("recommendation") or "Demander clarification")[:80], size=Pt(8))
                    _shade_row(contra_table.rows[-1], "FEF2F2")
                else:
                    row = contra_table.add_row().cells
                    _set_cell_text(row[0], str(c)[:120], size=Pt(8), color=RED)

        env_req = cctp.get("environmental_requirements", [])
        if env_req:
            doc.add_heading("Exigences environnementales", 2)
            _add_encadre(
                "\n".join(f"• {(req.get('requirement') if isinstance(req, dict) else str(req))}"
                          for req in env_req[:5]),
                bg_hex="EFF6FF", border_hex="3B82F6",
                text_color=RGBColor(0x1E, 0x3A, 0x8A)
            )

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §13. CONTRADICTIONS INTER-DOCUMENTS
    # ══════════════════════════════════════════════════════════════════════
    if conflicts_data:
        conflicts_list = conflicts_data.get("conflicts", [])
        nb_total = conflicts_data.get("nb_total", len(conflicts_list))
        nb_critical = conflicts_data.get("nb_critical", 0)

        doc.add_heading("13. Contradictions inter-documents", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Incohérences détectées entre les pièces du DCE (RC, CCAP, CCTP, DPGF, AE)")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        _add_encadre(
            f"{nb_total} contradiction(s) détectée(s) dont {nb_critical} critique(s). "
            "Posez la question à l'acheteur via la plateforme de dématérialisation.",
            bg_hex="FEF2F2", border_hex="EF4444",
            text_color=RGBColor(0x99, 0x1B, 0x1B)
        )

        if conflicts_list:
            cf_table = _styled_table(["Type", "Doc A ↔ Doc B", "Description", "Sévérité", "Recommandation"])
            for c in conflicts_list[:15]:
                if isinstance(c, dict):
                    row = cf_table.add_row().cells
                    _set_cell_text(row[0], (c.get("type") or "—")[:20], bold=True, size=Pt(8))
                    docs = f"{c.get('doc_a') or '?'} ↔ {c.get('doc_b') or '?'}"
                    _set_cell_text(row[1], docs[:25], size=Pt(8))
                    _set_cell_text(row[2], (c.get("description") or "")[:120], size=Pt(8))
                    sev = c.get("severity") or "—"
                    _set_cell_text(row[3], sev.upper() if isinstance(sev, str) else str(sev),
                                  bold=True, color=_severity_color(sev), size=Pt(8))
                    bg = _severity_bg(sev)
                    if bg:
                        _shade_cell(row[3], bg)
                    _set_cell_text(row[4], (c.get("recommendation") or c.get("recommandation") or "Demander clarification")[:80],
                                  size=Pt(8))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §14. ANALYSE SOUS-TRAITANCE
    # ══════════════════════════════════════════════════════════════════════
    if subcontracting:
        doc.add_heading("14. Analyse sous-traitance", 1)
        p_sub_st = doc.add_paragraph()
        run_sub_st = p_sub_st.add_run("Obligations et risques liés à la sous-traitance — conformité loi du 31/12/1975")
        run_sub_st.font.color.rgb = GRAY
        run_sub_st.font.size = Pt(9)

        if subcontracting.get("summary"):
            _add_encadre(subcontracting["summary"], bg_hex="EFF6FF", border_hex="2563EB")

        # Conditions as structured table with risk level (matching PDF)
        conditions = subcontracting.get("conditions", [])
        if conditions:
            doc.add_heading("Conditions contractuelles", 2)
            cond_table = _styled_table(["Condition", "Détail", "Risque"])
            for c in conditions:
                if isinstance(c, dict):
                    row = cond_table.add_row().cells
                    _set_cell_text(row[0], (c.get("label") or c.get("condition") or "")[:50],
                                  bold=True, size=Pt(8))
                    _set_cell_text(row[1], (c.get("detail") or c.get("value") or "")[:120], size=Pt(8))
                    c_risk = c.get("risk", "")
                    risk_label = (c_risk.upper() if c_risk else "INFO")
                    _set_cell_text(row[2], risk_label, bold=True,
                                  color=_severity_color(c_risk), size=Pt(8))
                else:
                    row = cond_table.add_row().cells
                    _set_cell_text(row[0], str(c)[:50], size=Pt(8))

        # Risks as warning boxes with mitigation (matching PDF)
        sub_risks = subcontracting.get("risks", [])
        if sub_risks:
            doc.add_heading("Risques identifiés", 2)
            for r in sub_risks[:5]:
                if isinstance(r, dict):
                    risk_text = f"{r.get('risk') or r.get('title') or ''} : {r.get('detail') or r.get('description') or ''}"
                    mitigation = r.get("mitigation")
                    if mitigation:
                        risk_text += f"\nMitigation : {mitigation}"
                    _add_encadre(risk_text, bg_hex="FEF2F2", border_hex="EF4444",
                                text_color=RGBColor(0x99, 0x1B, 0x1B))
                else:
                    _add_encadre(str(r), bg_hex="FEF2F2", border_hex="EF4444",
                                text_color=RGBColor(0x99, 0x1B, 0x1B))

        # Legal obligations (matching PDF)
        legal_obligations = subcontracting.get("legal_obligations", [])
        if legal_obligations:
            doc.add_heading("Obligations légales", 2)
            obligations_text = "\n".join(f"• {o}" for o in legal_obligations[:6])
            _add_encadre(obligations_text, bg_hex="EFF6FF", border_hex="3B82F6",
                        text_color=RGBColor(0x1E, 0x3A, 0x8A))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE V — SCORING ET STRATÉGIE
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE V — SCORING ET STRATÉGIE",
                     "Critères d'attribution, simulation note, Go/No-Go détaillé")

    # ══════════════════════════════════════════════════════════════════════
    #  §15. CRITÈRES D'ATTRIBUTION
    # ══════════════════════════════════════════════════════════════════════
    if criteria:
        doc.add_heading("15. Critères d'attribution", 1)
        ev = criteria.get("evaluation", {})

        elig = ev.get("eligibility_conditions", [])
        if elig:
            doc.add_heading("Conditions d'éligibilité", 2)
            e_table = _styled_table(["Condition", "Type"])
            for c in elig:
                row = e_table.add_row().cells
                _set_cell_text(row[0], c.get("condition", ""), size=Pt(9))
                ctype = (c.get("type") or "").upper()
                _set_cell_text(row[1], ctype, bold=True,
                              color=RED if ctype == "ELIMINATOIRE" else GRAY, size=Pt(9))

        scoring_crit = ev.get("scoring_criteria", [])
        if scoring_crit:
            doc.add_heading("Critères de notation", 2)
            sc_table = _styled_table(["Critère", "Pondération", "Notes"])
            for c in scoring_crit:
                row = sc_table.add_row().cells
                _set_cell_text(row[0], c.get("criterion", ""), size=Pt(9))
                w = c.get("weight_percent")
                _set_cell_text(row[1], f"{w}%" if w is not None else "N/S",
                              bold=True, color=ACCENT_BLUE, size=Pt(9))
                _set_cell_text(row[2], c.get("notes") or "—", size=Pt(9))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §16. SIMULATION NOTE ACHETEUR
    # ══════════════════════════════════════════════════════════════════════
    if scoring:
        doc.add_heading("16. Simulation note acheteur", 1)
        total = scoring.get("total_score", 0)
        max_s = scoring.get("max_score", 20)

        score_p = doc.add_paragraph()
        score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sc = score_p.add_run(f"{total:.1f} / {max_s}")
        run_sc.bold = True
        run_sc.font.size = Pt(24)
        run_sc.font.color.rgb = DARK_BLUE

        if scoring.get("rank_estimate"):
            rank_p = doc.add_paragraph()
            rank_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_rk = rank_p.add_run(f"Classement estimé : {scoring['rank_estimate']}")
            run_rk.font.size = Pt(11)
            run_rk.font.color.rgb = GRAY

        crit_scores = scoring.get("criteria_scores", [])
        if crit_scores:
            doc.add_heading("Détail par critère", 2)
            cs_table = _styled_table(["Critère", "Note", "Max", "Commentaire"])
            for cs in crit_scores:
                row = cs_table.add_row().cells
                _set_cell_text(row[0], cs.get("criterion", ""), size=Pt(9))
                _set_cell_text(row[1], str(cs.get("score", 0)), bold=True,
                              color=ACCENT_BLUE, size=Pt(9))
                _set_cell_text(row[2], str(cs.get("max_score", 20)), size=Pt(9))
                _set_cell_text(row[3], cs.get("comment") or cs.get("justification") or "—", size=Pt(9))

        levers = scoring.get("improvement_levers", [])
        if levers:
            doc.add_heading("Leviers d'amélioration", 2)
            _add_encadre(
                "Actions concrètes pour améliorer votre note :",
                bg_hex="ECFDF5", border_hex="10B981",
                text_color=GREEN_DARK
            )
            for lev in levers:
                p = doc.add_paragraph(style="List Bullet")
                title_text = lev.get("lever") or lev.get("title", "")
                impact = lev.get("potential_gain") or lev.get("impact", "")
                run_t = p.add_run(title_text)
                run_t.bold = True
                if impact:
                    run_i = p.add_run(f" — {impact}")
                    run_i.font.color.rgb = GREEN_DARK

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §17. GO/NO-GO DÉTAILLÉ — 9 DIMENSIONS
    # ══════════════════════════════════════════════════════════════════════
    if gonogo and (gonogo.get("dimension_scores") or gonogo.get("breakdown")):
        doc.add_heading("17. Go/No-Go détaillé — 9 dimensions", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Analyse multidimensionnelle de l'opportunité — scores par axe d'évaluation")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        score = gonogo.get("score", 0)
        rec = (gonogo.get("recommendation") or "ATTENTION").upper()

        # Score box
        if rec == "GO":
            bg, border_c = "D1FAE5", "059669"
            txt_color = GREEN_DARK
        elif rec == "NO-GO":
            bg, border_c = "FDE8E8", "DC2626"
            txt_color = RED
        else:
            bg, border_c = "FEF3C7", "D97706"
            txt_color = ORANGE
        _add_encadre(f"Score global : {score}/100 — Recommandation : {rec}", bg_hex=bg, border_hex=border_c, text_color=txt_color)

        # Dimension scores table
        dims = gonogo.get("dimension_scores", [])
        if dims:
            doc.add_heading("Détail par dimension", 2)
            dim_table = _styled_table(["Dimension", "Score", "Poids", "Confiance", "Explication"])
            for d in dims:
                row = dim_table.add_row().cells
                _set_cell_text(row[0], d.get("name") or d.get("dimension") or "—", bold=True, size=Pt(8))
                d_score = d.get("score", 0)
                _set_cell_text(row[1], f"{d_score}/100", bold=True, size=Pt(8))
                if d_score >= 70:
                    _shade_cell(row[1], "D1FAE5")
                elif d_score < 50:
                    _shade_cell(row[1], "FEF2F2")
                _set_cell_text(row[2], f"{d.get('weight', '—')}%", size=Pt(8))
                _set_cell_text(row[3], f"{d.get('confidence', 0)*100:.0f}%" if isinstance(d.get('confidence'), (int, float)) else "—", size=Pt(8))
                _set_cell_text(row[4], (d.get("explanation") or d.get("justification") or "")[:100], size=Pt(8))
        elif gonogo.get("breakdown"):
            # Fallback to old 4-dimension format
            doc.add_heading("Scores par dimension", 2)
            breakdown = gonogo["breakdown"]
            dim_table = _styled_table(["Dimension", "Score", "Niveau"])
            for dim_name, dim_key in [
                ("Adéquation technique", "technical_fit"),
                ("Capacité financière", "financial_capacity"),
                ("Faisabilité planning", "timeline_feasibility"),
                ("Position concurrentielle", "competitive_position"),
                ("Certifications", "certifications_match"),
                ("Zone géographique", "geographic_fit"),
                ("Assurances", "insurance_coverage"),
                ("Sous-traitance", "subcontracting_risk"),
                ("Taux de succès", "success_rate"),
            ]:
                dim_score = breakdown.get(dim_key)
                if dim_score is not None:
                    row = dim_table.add_row().cells
                    _set_cell_text(row[0], dim_name, size=Pt(9))
                    _set_cell_text(row[1], f"{dim_score}/100", bold=True, size=Pt(9))
                    if dim_score >= 70:
                        level_label, level_color, level_bg = "BON", GREEN_DARK, "D1FAE5"
                    elif dim_score >= 50:
                        level_label, level_color, level_bg = "MOYEN", ORANGE, "FEF3C7"
                    else:
                        level_label, level_color, level_bg = "FAIBLE", RED, "FEF2F2"
                    _set_cell_text(row[2], level_label, bold=True, color=level_color, size=Pt(9))
                    _shade_cell(row[2], level_bg)

        # Profile gaps
        gaps = gonogo.get("profile_gaps", [])
        if gaps:
            doc.add_heading("Écarts profil entreprise", 2)
            _add_encadre(
                "Points à renforcer :\n" + "\n".join(f"• {g}" for g in gaps[:8]),
                bg_hex="FEF2F2", border_hex="EF4444",
                text_color=RGBColor(0x99, 0x1B, 0x1B)
            )

        # Profile strengths
        strengths_list = gonogo.get("profile_strengths", [])
        if strengths_list:
            doc.add_heading("Points forts du profil", 2)
            _add_encadre(
                "Atouts de votre candidature :\n" + "\n".join(f"• {s}" for s in strengths_list[:8]),
                bg_hex="D1FAE5", border_hex="059669",
                text_color=GREEN_DARK
            )

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PARTIE VI — PLAN D'ACTION
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("PARTIE VI — PLAN D'ACTION",
                     "Risques, questions, calendrier")

    # ══════════════════════════════════════════════════════════════════════
    #  §18. ANALYSE DES RISQUES
    # ══════════════════════════════════════════════════════════════════════
    risks = summary.get("risks", [])
    if risks:
        doc.add_heading("18. Analyse des risques", 1)
        risk_table = _styled_table(["Risque", "Sévérité", "Explication"])
        for r in risks:
            row = risk_table.add_row().cells
            _set_cell_text(row[0], r.get("risk", ""), size=Pt(9))
            sev = r.get("severity") or ""
            sev_label = {"high": "FORT", "medium": "MOYEN", "low": "BAS"}.get(sev.lower(), sev.upper())
            _set_cell_text(row[1], sev_label, bold=True, color=_severity_color(sev), size=Pt(9))
            bg = _severity_bg(sev)
            if bg:
                _shade_cell(row[1], bg)
            _set_cell_text(row[2], r.get("why", ""), size=Pt(9))

    # Actions 48h
    actions = summary.get("actions_next_48h", [])
    if actions:
        doc.add_heading("Actions prioritaires sous 48h", 2)
        _add_encadre(
            "Ces actions sont à engager immédiatement pour sécuriser la réponse à l'appel d'offres.",
            bg_hex="FEF2F2", border_hex="EF4444",
            text_color=RGBColor(0x99, 0x1B, 0x1B)
        )
        act_table = _styled_table(["Action", "Responsable", "Priorité"])
        for a in actions:
            row = act_table.add_row().cells
            _set_cell_text(row[0], a.get("action", ""), size=Pt(9))
            _set_cell_text(row[1], a.get("owner_role", ""), size=Pt(9))
            prio = a.get("priority", "")
            _set_cell_text(row[2], prio.upper() if prio else "—", bold=True,
                          color=_severity_color(prio), size=Pt(9))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §19. QUESTIONS PRIORITAIRES POUR L'ACHETEUR
    # ══════════════════════════════════════════════════════════════════════
    if questions_list:
        doc.add_heading("19. Questions prioritaires pour l'acheteur", 1)
        _add_encadre(
            "Conseil : Posez ces questions via la plateforme de dématérialisation "
            "avant la date limite de questions. Les réponses seront diffusées à tous les candidats.",
            bg_hex="EFF6FF", border_hex="3B82F6",
            text_color=RGBColor(0x1E, 0x3A, 0x8A)
        )
        q_table = _styled_table(["#", "Priorité", "Question", "Justification"])
        for idx, q in enumerate(questions_list[:15], 1):
            row = q_table.add_row().cells
            _set_cell_text(row[0], str(idx), size=Pt(9))
            prio = (q.get("priority") or "—")
            _set_cell_text(row[1], prio.upper(), bold=True,
                          color=_severity_color(prio), size=Pt(9))
            _set_cell_text(row[2], (q.get("question") or str(q))[:140], size=Pt(8))
            _set_cell_text(row[3], (q.get("justification") or "")[:100], size=Pt(8))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  §20. CALENDRIER ET DATES CLÉS
    # ══════════════════════════════════════════════════════════════════════
    if timeline:
        doc.add_heading("20. Calendrier et dates clés", 1)
        cal_table = _styled_table(["Échéance", "Date / Durée"])

        if timeline.get("submission_deadline"):
            row = cal_table.add_row().cells
            _shade_row(cal_table.rows[-1], "FEF2F2")
            cell0 = row[0]
            cell0.text = ""
            p = cell0.paragraphs[0]
            run_t = p.add_run("Date limite de remise ")
            run_t.bold = True
            run_t.font.size = Pt(9)
            run_imp = p.add_run("(impératif)")
            run_imp.font.color.rgb = RED
            run_imp.font.size = Pt(7)
            run_imp.bold = True
            _set_cell_text(row[1], _fmt_date(timeline["submission_deadline"]),
                          bold=True, color=RED, size=Pt(9))

        if timeline.get("execution_duration_months"):
            row = cal_table.add_row().cells
            _set_cell_text(row[0], "Durée d'exécution", size=Pt(9))
            _set_cell_text(row[1], f"{timeline['execution_duration_months']} mois", size=Pt(9))

        for kd in timeline.get("key_dates", []):
            label = kd.get("label", "")
            if "remise des offres" in label.lower():
                continue
            row = cal_table.add_row().cells
            cell0 = row[0]
            cell0.text = ""
            p = cell0.paragraphs[0]
            run_t = p.add_run(label[:55])
            run_t.font.size = Pt(9)
            if kd.get("mandatory"):
                run_m = p.add_run(" (oblig.)")
                run_m.font.color.rgb = RED
                run_m.font.size = Pt(7)
            _set_cell_text(row[1], _fmt_date(kd.get("date")), size=Pt(9))

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  ANNEXES
    # ══════════════════════════════════════════════════════════════════════
    _add_part_header("ANNEXES",
                     "Documents analysés, glossaire, mentions légales")

    # ══════════════════════════════════════════════════════════════════════
    #  A1. DOCUMENTS ANALYSÉS
    # ══════════════════════════════════════════════════════════════════════
    if documents:
        doc.add_heading("A1. Inventaire des documents analysés", 1)
        p_sub_doc = doc.add_paragraph()
        run_sub_doc = p_sub_doc.add_run("Liste des pièces du DCE prises en compte dans cette analyse")
        run_sub_doc.font.color.rgb = GRAY
        run_sub_doc.font.size = Pt(9)

        doc_table = _styled_table(["#", "Document", "Type", "Pages", "Taille", "Qualité OCR"])
        for idx_d, d in enumerate(documents, 1):
            row = doc_table.add_row().cells
            _set_cell_text(row[0], str(idx_d), size=Pt(8))
            _set_cell_text(row[1], (d.original_name or "—")[:45], size=Pt(8))
            dtype = d.doc_type or "AUTRES"
            _set_cell_text(row[2], dtype, bold=True,
                          color=ACCENT_BLUE if dtype in ("RC", "CCTP", "CCAP") else GRAY,
                          size=Pt(8))
            _set_cell_text(row[3], str(d.page_count or "—"), size=Pt(8))
            # Size display
            file_size = getattr(d, "file_size", None)
            if file_size:
                if file_size >= 1_000_000:
                    size_disp = f"{file_size / 1_000_000:.1f} Mo"
                elif file_size >= 1_000:
                    size_disp = f"{file_size / 1_000:.0f} Ko"
                else:
                    size_disp = f"{file_size} o"
            else:
                size_disp = "—"
            _set_cell_text(row[4], size_disp, size=Pt(8))
            # OCR quality
            ocr_q = getattr(d, "ocr_confidence", None)
            if ocr_q is not None:
                if ocr_q >= 70:
                    ocr_color, ocr_bg = GREEN_DARK, "D1FAE5"
                elif ocr_q >= 40:
                    ocr_color, ocr_bg = ORANGE, "FEF3C7"
                else:
                    ocr_color, ocr_bg = RED, "FEF2F2"
                _set_cell_text(row[5], f"{ocr_q:.0f}%", bold=True, color=ocr_color, size=Pt(8))
                _shade_cell(row[5], ocr_bg)
            else:
                _set_cell_text(row[5], "—", size=Pt(8))

        # Total summary line
        total_pages = sum(d.page_count or 0 for d in documents)
        total_p = doc.add_paragraph()
        run_total = total_p.add_run(
            f"Total : {len(documents)} document{'s' if len(documents) > 1 else ''}"
            f"{f' — {total_pages} pages analysées' if total_pages else ''}"
        )
        run_total.font.size = Pt(9)
        run_total.font.color.rgb = GRAY

    # ══════════════════════════════════════════════════════════════════════
    #  A2. GLOSSAIRE BTP
    # ══════════════════════════════════════════════════════════════════════
    if glossaire_btp:
        doc.add_page_break()
        doc.add_heading("A2. Glossaire BTP", 1)
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("Termes clés utilisés dans ce rapport — Référentiel marchés publics BTP")
        run_sub.font.color.rgb = GRAY
        run_sub.font.size = Pt(9)

        gl_table = _styled_table(["Terme", "Définition"])
        for term, definition in glossaire_btp:
            row = gl_table.add_row().cells
            _set_cell_text(row[0], term, bold=True, color=ACCENT_BLUE, size=Pt(9))
            _set_cell_text(row[1], str(definition)[:200], size=Pt(8))

    # ══════════════════════════════════════════════════════════════════════
    #  A3. AVERTISSEMENT IA ET MENTIONS LÉGALES
    # ══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("A3. Avertissement IA et mentions légales", 1)

    # Disclaimer box (yellow background like PDF)
    disclaimer_text = (
        "Avertissement : Ce rapport est généré par intelligence artificielle "
        "(Claude, Anthropic) à partir des documents du DCE fournis. "
        "Il constitue une aide à la décision et ne se substitue pas à "
        "l'analyse humaine d'un expert marchés publics. "
        "Les informations extraites doivent être systématiquement vérifiées "
        "avant toute soumission d'offre."
    )
    # Add confidence score like PDF
    conf_val = summary.get("confidence") or summary.get("avg_confidence")
    if conf_val:
        disclaimer_text += f"\n\nConfiance globale de l'analyse : {conf_val * 100:.0f}%"
    _add_encadre(
        disclaimer_text,
        bg_hex="FFFBEB", border_hex="FBBF24",
        text_color=RGBColor(0x78, 0x35, 0x0F)
    )

    doc.add_paragraph()

    _add_encadre(
        "Données utilisées : Seuls les documents uploadés dans le projet "
        "ont été analysés. L'IA ne dispose d'aucune information externe "
        "(pas d'accès internet, pas de base de données externe). "
        "La qualité de l'analyse dépend directement de la qualité "
        "des documents fournis (lisibilité OCR, exhaustivité du DCE).",
        bg_hex="EFF6FF", border_hex="3B82F6",
        text_color=RGBColor(0x1E, 0x3A, 0x8A)
    )

    # ══════════ FOOTER ══════════
    doc.add_paragraph()
    doc.add_paragraph()
    footer_brand = doc.add_paragraph()
    footer_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fb = footer_brand.add_run("AO COPILOT")
    run_fb.bold = True
    run_fb.font.color.rgb = ACCENT_BLUE
    run_fb.font.size = Pt(11)

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fp = footer_para.add_run(
        f"aocopilot.fr — {datetime.now().strftime('%d/%m/%Y %H:%M')} — "
        "Rapport confidentiel — Reproduction interdite"
    )
    run_fp.font.color.rgb = GRAY_LIGHT
    run_fp.font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
